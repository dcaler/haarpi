"""The current draft is not always markdown.

A redline revise writes tracked changes into the docx and returns before the markdown is ever
written — the accepted text is the reviewer's to settle, so there is no `.md` of a draft whose
changes are still proposals. Everything downstream that reads "the current draft" therefore has
to read a docx, with insertions applied and deletions gone. `mindmap` did not, and died with
"render a draft first" standing in a directory that held the draft.
"""
import docx
from docx.oxml.ns import qn

from rabbithole import docxio, mindmap


def _tracked(para, text: str, author: str = "rabbitHole") -> None:
    """Move a paragraph's runs inside a `w:ins`, as a redline revise leaves them."""
    run = para.add_run(text)
    ins = para._p.makeelement(qn("w:ins"), {qn("w:id"): "1", qn("w:author"): author,
                                            qn("w:date"): "2026-09-05T09:38:28Z"})
    para._p.replace(run._r, ins)
    ins.append(run._r)


def _draft(path, *, insert_heading=False):
    d = docx.Document()
    d.add_heading("Literature Review: household sorting", level=1)
    d.add_heading("Narrative Review", level=2)
    d.add_heading("Structural convenience drives compliance", level=2)
    d.add_paragraph("Distance cuts missorting [@rousta2015].")
    if insert_heading:
        _tracked(d.add_heading("", level=2), "Norms amplify behaviour")
        _tracked(d.add_paragraph(), "Norms move high users [@allcott2011].")
    d.add_heading("Annotated Bibliography", level=2)
    d.add_paragraph("Rousta, K. (2015). Distance. [@rousta2015]")
    d.save(str(path))
    return path


def test_a_tracked_insertion_is_not_invisible(tmp_path):
    """python-docx's own `paragraph.text` reads only runs that are direct children of `w:p`, so
    a redline's every edit read as an empty string. Two headings in a real review came back ''."""
    p = _draft(tmp_path / "d.docx", insert_heading=True)
    body = docxio.read_body_text(p)
    assert "Norms amplify behaviour" in body
    assert "Norms move high users [@allcott2011]." in body


def test_a_deletion_stays_gone(tmp_path):
    d = docx.Document()
    para = d.add_paragraph()
    run = para.add_run("struck out")
    dl = para._p.makeelement(qn("w:del"), {qn("w:id"): "2", qn("w:author"): "DCR",
                                           qn("w:date"): "2026-09-05T09:38:28Z"})
    para._p.replace(run._r, dl)
    run._r.find(qn("w:t")).tag = qn("w:delText")      # deleted text moves to w:delText
    dl.append(run._r)
    para.add_run("kept")
    d.save(str(tmp_path / "d.docx"))
    body = docxio.read_body_text(tmp_path / "d.docx")
    assert "kept" in body and "struck out" not in body


def test_heading_styles_come_back_as_markdown_hashes(tmp_path):
    md = docxio.read_body_markdown(_draft(tmp_path / "d.docx"))
    assert "# Literature Review: household sorting" in md.splitlines()
    assert "## Structural convenience drives compliance" in md.splitlines()
    assert "Distance cuts missorting [@rousta2015]." in md.splitlines()   # body prose unprefixed


def test_the_reconstructed_markdown_parses_as_threads(tmp_path):
    """The point of the heading reconstruction: the same section parser, either format."""
    md = docxio.read_body_markdown(_draft(tmp_path / "d.docx", insert_heading=True))
    ts = mindmap.parse_threads(md)
    assert [t.theme for t in ts] == ["Structural convenience drives compliance",
                                     "Norms amplify behaviour"]     # the inserted one included
    assert ts[1].citekeys == ["allcott2011"]                        # and its citations with it


def test_find_review_takes_the_docx_when_a_redline_left_no_markdown(tmp_path):
    p = _draft(tmp_path / "260815_proj_litreview_ra_DCR_ra.docx")
    assert mindmap._find_review(tmp_path) == p


def test_find_review_prefers_the_markdown_of_the_same_draft(tmp_path):
    """The resynth path writes both, docx last — so newest-by-mtime would pick the round-trip."""
    md = tmp_path / "260815_proj_litreview_ra.md"
    md.write_text("## Theme\n\nText [@k].\n")
    _draft(tmp_path / "260815_proj_litreview_ra.docx")        # written after the markdown
    assert mindmap._find_review(tmp_path) == md


def test_find_review_ignores_a_word_lock_file(tmp_path):
    """`~$…docx` is Word's lock stub, not a draft — opening it raises."""
    p = _draft(tmp_path / "260815_proj_litreview_ra.docx")
    (tmp_path / "~$0815_proj_litreview_ra.docx").write_bytes(b"not a docx")
    assert mindmap._find_review(tmp_path) == p


def test_find_review_is_empty_handed_when_there_is_no_draft(tmp_path):
    (tmp_path / "refs.bib").write_text("@article{k,}\n")
    assert mindmap._find_review(tmp_path) is None


def test_read_review_dispatches_on_the_suffix(tmp_path):
    md = tmp_path / "a_litreview_ra.md"
    md.write_text("## Theme\n\nText [@k].\n")
    assert mindmap._read_review(md) == md.read_text()
    docx_draft = _draft(tmp_path / "b_litreview_ra.docx")
    assert mindmap._read_review(docx_draft).startswith("# Literature Review")
