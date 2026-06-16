"""rabbitHole revise — apply reviewer annotations from a _ra.docx to re-draft the narrative.

Pipeline (fast — no re-annotation, no re-location):
  1. Find *_ra.docx in output/ (or accept --file path)
  2. Extract tracked changes + comments from the docx
  3. Load existing notes from work/annotations/ and slim corpus from work/corpus.json
  4. Re-synthesise the narrative using the revision brief
  5. Rebuild bibliography from existing work/located/ files
  6. Write output/*_ra_r{N}.md and .docx
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

from . import config, render
from .brain import Brain
from .models import Candidate
from .summarize import (
    _make_citekeys, _digest, bibliography, citation_check,
    SYNTH_SYS, _critique_revise_synthesis,
)


def _require_docx() -> None:
    if not _DOCX_OK:
        print("[error] python-docx is required: pip install python-docx", file=sys.stderr)
        raise SystemExit(1)


# ── docx annotation extraction (mirrors raconteur/revise.py) ─────────────────

def _read_comments(path: Path) -> list[dict]:
    doc = Document(str(path))
    comments = []
    try:
        for rel in doc.part.rels.values():
            if "comments" not in rel.reltype.lower():
                continue
            for c in rel.target_part._element.findall(".//" + qn("w:comment")):
                author = c.get(qn("w:author"), "reviewer")
                texts = [t.text for t in c.findall(".//" + qn("w:t")) if t.text]
                if texts:
                    comments.append({"author": author, "text": " ".join(texts)})
            break
    except Exception:  # noqa: BLE001
        pass
    return comments


def _read_track_changes(path: Path) -> dict:
    doc = Document(str(path))
    body = doc.element.body
    insertions, deletions = [], []
    for ins in body.iter(qn("w:ins")):
        text = "".join(t.text or "" for t in ins.iter(qn("w:t")))
        if text.strip():
            insertions.append(text)
    for dele in body.iter(qn("w:del")):
        text = "".join(t.text or "" for t in dele.iter(qn("w:delText")))
        if text.strip():
            deletions.append(text)
    return {"insertions": insertions, "deletions": deletions}


def _read_body_text(path: Path) -> str:
    """Body text with tracked insertions visible, deletions gone (python-docx default)."""
    doc = Document(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _build_revision_context(path: Path) -> str:
    comments = _read_comments(path)
    changes = _read_track_changes(path)
    parts = []
    if changes["deletions"]:
        lines = "\n".join(f"  - {d}" for d in changes["deletions"])
        parts.append(f"DELETED TEXT (the reviewer removed these):\n{lines}")
    if changes["insertions"]:
        lines = "\n".join(f"  + {i}" for i in changes["insertions"])
        parts.append(f"INSERTED TEXT (the reviewer added these):\n{lines}")
    if comments:
        lines = "\n".join(f"  [{c['author']}]: {c['text']}" for c in comments)
        parts.append(f"REVIEWER COMMENTS:\n{lines}")
    return "\n\n".join(parts)


# ── file discovery ────────────────────────────────────────────────────────────

def _find_docx(paths) -> Path | None:
    """Find the most recently modified docx in output/ that was annotated by the user.

    Annotated files end in user initials (e.g. _DCR.docx); rabbitHole's own
    outputs end in _ra.docx. We exclude the latter so we never accidentally
    re-process our own output.
    """
    candidates = [p for p in paths.output.glob("*.docx")
                  if not p.stem.endswith("_ra")]
    if not candidates:
        return None
    return max(candidates, key=lambda p: p.stat().st_mtime)


def _revision_paths(paths, annotated_docx: Path) -> tuple[Path, Path]:
    """Output paths: append _ra to the annotated file's stem."""
    stem = f"{annotated_docx.stem}_ra"
    return paths.output / f"{stem}.md", paths.output / f"{stem}.docx"


# ── load existing intermediates ───────────────────────────────────────────────

def _load_corpus(paths) -> list[Candidate]:
    if not paths.corpus_json.exists():
        return []
    data = json.loads(paths.corpus_json.read_text(encoding="utf-8"))
    return [Candidate.from_dict(d) for d in data]


def _load_notes(paths, n: int) -> list[dict]:
    notes = []
    for i in range(n):
        fp = paths.annotations_dir / f"{i:03d}.json"
        notes.append(json.loads(fp.read_text(encoding="utf-8")) if fp.exists() else {})
    return notes


def _load_located(paths, n: int) -> list[list[dict]]:
    located_dir = paths.work / "located"
    located = []
    for i in range(n):
        fp = located_dir / f"{i:03d}.json"
        located.append(json.loads(fp.read_text(encoding="utf-8")) if fp.exists() else [])
    return located


# ── revision synthesis ────────────────────────────────────────────────────────

_REVISE_PROMPT = """\
Review topic: {topic}
Focus: {focus}

CURRENT NARRATIVE (the draft you are revising):
{narrative}

EVIDENCE DIGEST (ground truth — stay grounded in these sources):
{digest}

REVISION ANNOTATIONS (apply all of these):
{revision_context}

Produce the revised narrative. Apply every annotation faithfully. \
Maintain [@citekey] citation format throughout. \
Do not add a bibliography — that is generated separately."""


def _synthesize_revision(brain: Brain, cfg, corpus: list[Candidate],
                         notes: list[dict], citekeys: dict[int, str],
                         current_narrative: str, revision_context: str,
                         style_profile: str = "") -> str:
    digest = _digest(corpus, notes, citekeys)
    prompt = _REVISE_PROMPT.format(
        topic=cfg.topic, focus=cfg.focus,
        narrative=current_narrative.strip(),
        digest=digest,
        revision_context=revision_context,
    )
    sys_prompt = SYNTH_SYS
    if style_profile:
        sys_prompt = (sys_prompt.rstrip()
                      + f"\n\nWRITING STYLE\nMatch the following author's voice and "
                        f"prose style throughout:\n{style_profile}")
    print("  Re-synthesising narrative (coordinator)...", flush=True)
    narrative = brain.coordinator(prompt, sys_prompt, num_ctx=16384)
    return _critique_revise_synthesis(brain, narrative, digest, cfg.topic, cfg.focus or "")


# ── orchestration ─────────────────────────────────────────────────────────────

def run(directory: str = ".", brain_override: str | None = None,
        docx_path: str | None = None) -> int:
    _require_docx()

    cfg = config.load_project(directory)
    gc = config.load_global()
    paths = config.project_paths(directory)
    brain = Brain(cfg.brain, gc, backend_override=brain_override)

    print(f"rabbitHole revise — {cfg.project_name}")

    # 1. Find the docx to revise
    if docx_path:
        docx = Path(docx_path)
    else:
        docx = _find_docx(paths)
    if not docx or not docx.exists():
        print("[error] No *_ra.docx found in output/. "
              "Specify one with --file or run 'rabbitHole report' first.", file=sys.stderr)
        return 1
    print(f"  Annotated file: {docx.name}")

    # 2. Extract annotations
    revision_context = _build_revision_context(docx)
    if not revision_context:
        print("[warn] No tracked changes or comments found in the docx. Nothing to revise.")
        return 0
    n_comments = len(_read_comments(docx))
    tc = _read_track_changes(docx)
    print(f"  Found: {n_comments} comment(s), "
          f"{len(tc['deletions'])} deletion(s), {len(tc['insertions'])} insertion(s).")

    # 3. Read current narrative from the matching .md (same stem, without user initials)
    #    e.g. digipros_litreview_ra_DCR.docx → look for digipros_litreview_ra.md
    ra_stem = re.sub(r"_[^_]+$", "_ra", docx.stem)  # replace last suffix with _ra
    md_path = paths.output / f"{ra_stem}.md"
    if not md_path.exists():
        # Fall back to body text extracted from the docx itself
        current_narrative = _read_body_text(docx)
    else:
        current_narrative = md_path.read_text(encoding="utf-8")

    # 4. Load corpus + notes
    corpus = _load_corpus(paths)
    if not corpus:
        print("[error] work/corpus.json not found — run 'rabbitHole report' first.",
              file=sys.stderr)
        return 1
    notes = _load_notes(paths, len(corpus))
    citekeys = _make_citekeys(corpus)

    # 5. Style profile
    style_profile = ""
    if cfg.use_style:
        from .style import load_style_profile
        style_profile = load_style_profile()

    # 6. Re-synthesise
    print()
    narrative = _synthesize_revision(brain, cfg, corpus, notes, citekeys,
                                     current_narrative, revision_context, style_profile)

    # 7. Bibliography from existing located intermediates
    located_list = _load_located(paths, len(corpus))
    located = {i: items for i, items in enumerate(located_list) if items}
    biblio = bibliography(corpus, located)
    unmatched = citation_check(narrative, citekeys)
    if unmatched:
        print(f"\n[citation check] {len(unmatched)} unmatched citekey(s): "
              f"{', '.join(f'[@{k}]' for k in unmatched[:8])}"
              + (" ..." if len(unmatched) > 8 else ""))

    # 8. Write output
    out_md, out_docx = _revision_paths(paths, docx)
    from .render import build_markdown, pandoc_convert
    md_text = build_markdown(cfg, brain.backend, narrative, biblio, corpus, unmatched)
    out_md.write_text(md_text, encoding="utf-8")
    pandoc_convert(out_md, out_docx)

    print()
    print("=" * 60)
    print(" revise complete")
    print("=" * 60)
    print(f"  Review (md)  : {out_md}")
    if out_docx.exists():
        print(f"  Review (docx): {out_docx}")
    return 0
