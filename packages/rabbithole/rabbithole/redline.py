"""In-place, comment-preserving revision with tracked changes.

This is what `revise` does by default. The alternative (`revise --resynth`) re-synthesises
the whole narrative from markdown and renders a fresh .docx — which discards the reviewer's
Word comments and gives them no redline to read the tool's edits against. This module instead
edits a COPY of the annotated .docx in place: it answers each comment by rewriting only the
paragraph(s) that comment is anchored to, records every rewrite as a Word tracked change
attributed to `rabbitHole`, and leaves the comments anchored and every un-flagged paragraph
byte-for-byte untouched.

The reviewer opens the result and sees, per comment, their note beside the tool's
tracked-change answer — accept/reject, re-comment, repeat.

This file is the deterministic machinery only: XML surgery, GPU-free and unit-testable.
The LLM call that turns a comment + evidence into revised paragraph text lives in
`revise` (it is the only part that needs the brain).

OOXML notes:
  * A comment is anchored by ``<w:commentRangeStart w:id=N/>`` … ``<w:commentRangeEnd
    w:id=N/>`` markers bracketing a run range, plus a ``<w:commentReference w:id=N/>``
    run; the text lives in comments.xml. python-docx preserves all of these across an
    open/save, so we only manipulate the body XML.
  * A tracked deletion wraps the old run(s) in ``<w:del>`` and turns ``<w:t>`` into
    ``<w:delText>``; a tracked insertion wraps new run(s) in ``<w:ins>``. Both carry an
    author and date, and Word renders them as an accept/rejectable redline.

The annotated bibliography is regenerated against the post-edit narrative (see
`accepted_body_text` / `replace_bibliography`), so a newly-cited source still gets a
verifiable entry. A comment anchored to a heading is never answered by rewriting the heading
itself: when it NAMES sources already in the corpus ("cite Doblinger 2019 and Howell 2017"),
the caller cites them in the section's first body paragraph (see `first_body_paragraph_under`);
when it names nothing citeable it is a "find more" ask and routes to the corpus chain.

A paragraph is modelled as an ordered stream of TEXT and OPAQUE atoms (equations,
hyperlinks, footnote references), not as the text inside its ``w:r/w:t`` runs. That older
model was blind to everything else in the paragraph: an equation is a SIBLING of the text
runs, so the differ saw prose with holes where every number had been, no sentence could
match, and each rewrite collapsed to a whole-paragraph replacement — with the equations
left stranded at the paragraph tail, severed from the claims they verified.

Atoms serialize to sentinels (``⟦m:1⟧``) for the differ and for the LLM, and expand back to
their original elements on write. rabbitHole never authors an atom: an equation is re-laid
as accepted content between the redlined prose around it, never inside a ``w:ins``/``w:del``.

Known limitations (documented, not bugs):
  * Multiple comments on one paragraph are coarsened to bracket the whole revised
    paragraph — every comment stays valid and anchored, but loses sub-paragraph
    precision. (`comment_spans` recovers that precision on the way IN, which is what tells
    the minimal-edit guard which sentences a comment actually bears on.)
  * Assumes the annotated draft has no still-open tracked changes from a prior cycle
    (true for a freshly rendered _ra draft the reviewer annotated).
"""

from __future__ import annotations

import copy
import datetime
import difflib
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from . import guards

# ── the shared engine (haarpi.redline), re-exported under the old names ──────
from haarpi import redline as _engine
from haarpi.redline import (  # noqa: F401
    _MATH, _XML_SPACE, _SENTINEL_SPLIT, _OPAQUE_RUN_CHILDREN,
    _now, _Ids, _max_existing_id, ids_for, _rpr_clone, _text_run, _ins, _del,
    _is_text_run, _is_ref_run, _is_opaque, _sentinel_kind,
    serialize_paragraph, paragraph_text, atom_text, flatten_paragraph,
    _render, _segments, _redline_chunk, _relay,
    comments_by_id, comment_spans, anchored_sentences, is_heading_style,
    _accepted_para_text,
)
from haarpi.redline import (  # noqa: F401 — author/ids are always passed explicitly here
    tracked_replace, tracked_replace_sentencewise, tracked_insert_after,
)


# Threaded-comment namespaces: w14 carries paraId on each comment paragraph; w15
# (commentsExtended.xml) links a reply's paraId to its parent's via paraIdParent.
_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"


def comment_anchors(path: Path) -> list[dict]:
    """Paragraphs carrying a comment anchor, with the comment ids and current text.

    Returns a list of {para, ids, text, style, anchored} in document order. ``text`` is the
    serialized paragraph (atoms as sentinels) — the exact string the reviser is asked to
    revise. ``style`` lets callers tell a comment on a heading from one on a body paragraph
    (a heading comment must not be answered by rewriting the heading). ``anchored`` is the
    union of sentence indices the paragraph's comments bear on. Paragraphs with no comment
    are omitted.
    """
    doc = Document(str(path))
    out = []
    for i, p in enumerate(doc.paragraphs):
        ids = [s.get(qn("w:id")) for s in p._p.findall(qn("w:commentRangeStart"))]
        if not ids:
            continue
        text = paragraph_text(p._p)
        anchored: set[int] = set()
        for span in comment_spans(p._p).values():
            anchored |= anchored_sentences(text, span)
        out.append({"para": i, "ids": ids, "text": text,
                    "style": p.style.name if p.style is not None else "",
                    "anchored": sorted(anchored)})
    return out


def first_body_paragraph_under(path: Path, heading_para: int) -> dict | None:
    """The first non-empty body paragraph in the section opened by the heading at
    ``heading_para``.

    A comment left on a heading cannot be answered by rewriting the heading — but "cite these
    sources here" means "cite them in this section", so the answer belongs in the section's
    prose. This finds where: the first body paragraph after the heading, stopping at the next
    heading. Returns ``{para, text}`` (``text`` serialized as :func:`comment_anchors` does, so
    the reviser reads the same string), or ``None`` when the section has no body paragraph
    before the next heading or the document ends.
    """
    doc = Document(str(path))
    for i, p in enumerate(doc.paragraphs):
        if i <= heading_para:
            continue
        style = p.style.name if p.style is not None else ""
        if is_heading_style(style):
            return None
        text = paragraph_text(p._p)
        if text.strip():
            return {"para": i, "text": text}
    return None


# ── post-edit narrative + bibliography regeneration ──────────────────────────────
# After the body is redlined the cited set may have changed, so the annotated
# bibliography must be regenerated against the CURRENT text to stay verifiable. We
# read the "accepted" narrative (inserted + unchanged text, deletions dropped — w:t
# lives in normal and <w:ins> runs; deleted text is in <w:delText>), re-locate, and
# replace the bibliography section wholesale. The body keeps its tracked changes; the
# bibliography is rebuilt clean — 30 entries of tracked-change noise would be unreadable
# and the bibliography is a generated artifact, not something the reviewer redlines.

_BIB_HEADING = "annotated bibliography"


def accepted_body_text(path: Path, stop_heading: str = _BIB_HEADING) -> str:
    """Reconstruct the narrative (citekeys intact) from a redlined docx.

    Returns the body up to the bibliography heading, with tracked changes accepted, so
    callers can see which [@citekey] tags the revised draft now cites.
    """
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        txt = _accepted_para_text(p._p)
        if txt.strip().lower().startswith(stop_heading):
            break
        if txt.strip():
            parts.append(txt)
    return "\n\n".join(parts)


def _parse_bibliography_md(md: str) -> tuple[str, list[tuple[str, object]]]:
    """Parse bibliography markdown into (heading, items).

    Each item is ("sub", subheading_text) for a ``### `` tier heading, or
    ("entry", (citation, [claim_line, ...])) for a source entry."""
    heading = "Annotated Bibliography"
    items: list[tuple[str, object]] = []
    cur_claims: list[str] | None = None
    for line in md.splitlines():
        s = line.strip()
        if s.startswith("### "):
            items.append(("sub", s[4:].strip()))
            cur_claims = None
        elif s.startswith("## "):
            heading = s[3:].strip()
        elif s.startswith("**") and s.endswith("**") and len(s) > 4:
            cur_claims = []
            items.append(("entry", (s[2:-2].strip(), cur_claims)))
        elif s.startswith("- ") and cur_claims is not None:
            cur_claims.append(s[2:].strip())
    return heading, items


# XML 1.0 forbids most C0 control characters in element text (only \t \n \r and the printable
# ranges are legal). PDF-extracted fulltext — the source of the bibliography's located quotes —
# routinely carries NULLs, form-feeds, and other control bytes; python-docx (lxml) then raises
# "All strings must be XML compatible" and the WHOLE bibliography write fails, so the stale one is
# kept. Strip the illegal chars at the write chokepoint: one bad glyph is dropped, never the section.
_XML_ILLEGAL = re.compile(r"[^\x09\x0A\x0D\x20-퟿-�\U00010000-\U0010FFFF]")


def _xml_safe(text: str) -> str:
    return _XML_ILLEGAL.sub("", text)


def _strip_md(text: str) -> str:
    """Drop the light markdown emphasis the bibliography lines carry (*…*) and remove any
    XML-illegal control characters so the run is always writable (see :data:`_XML_ILLEGAL`)."""
    return _xml_safe(text.replace("*", ""))


def replace_bibliography(path: Path, biblio_md: str) -> dict:
    """Replace the annotated-bibliography section of ``path`` with freshly built entries.

    Deletes from the bibliography heading to the end of the body and rebuilds it from
    ``biblio_md``, reusing the heading's own style so it matches the document. The body
    (and its tracked changes + comments) above the heading is untouched. Returns a summary.
    """
    # Sanitize the whole markdown up front: PDF fulltext (the quotes' source) carries NULLs and
    # form-feeds, and a form-feed is a page break — str.splitlines() would split a claim mid-line
    # (silently truncating it) before the illegal char ever reached the docx write. Strip both here.
    biblio_md = _xml_safe(biblio_md)
    doc = Document(str(path))
    bib_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower().startswith(_BIB_HEADING):
            bib_idx = i
            break
    heading_style = doc.paragraphs[bib_idx].style if bib_idx is not None else None
    if bib_idx is not None:
        for p in list(doc.paragraphs[bib_idx:]):
            p._element.getparent().remove(p._element)

    heading, items = _parse_bibliography_md(biblio_md)
    h = doc.add_paragraph()
    if heading_style is not None:
        h.style = heading_style
    h.add_run(_xml_safe(heading))
    n_entries = 0
    for kind, payload in items:
        if kind == "sub":
            sp = doc.add_paragraph()
            sp.add_run(_strip_md(payload)).bold = True  # tier heading (cited / additional)
            continue
        citation, claims = payload
        n_entries += 1
        cp = doc.add_paragraph()
        cp.add_run(_strip_md(citation)).bold = True
        for cl in claims:
            doc.add_paragraph().add_run("•  " + _strip_md(cl))

    doc.save(str(path))
    return {"bib_entries": n_entries, "had_existing_section": bib_idx is not None}


_TOP_HEADING = "most load-bearing sources"
_NARRATIVE_HEADING = "narrative review"


def replace_top_sources(path: Path, block_md: str) -> dict:
    """Rewrite the load-bearing-sources block that opens the review, inserting it if absent.

    The block ranks the sources the review rests on most, and it is only ever true of the draft
    it was computed from — a redline that adds citations, and a graft that adds whole sections,
    both change which sources carry the argument. Every re-draft verb must therefore recompute
    it, which is why this exists beside the bibliography regeneration rather than only in the
    full-render path: the redline path used to skip it, so a reviewed document silently kept the
    ranking of a draft several cycles old, or never grew one at all.

    Like the bibliography, this is a generated artifact rebuilt clean rather than redlined —
    tracked changes across a ranked list are noise the reviewer cannot act on.
    """
    block_md = _xml_safe(block_md or "")
    doc = Document(str(path))
    paras = doc.paragraphs

    def _heading_at(i):
        st = paras[i].style.name if paras[i].style is not None else ""
        return is_heading_style(st)

    start = end = None
    for i, par in enumerate(paras):
        if par.text.strip().lower().startswith(_TOP_HEADING) and _heading_at(i):
            start = i
            break
    if start is not None:
        end = len(paras)
        for j in range(start + 1, len(paras)):
            if _heading_at(j):
                end = j
                break

    # Where a fresh block goes: immediately before the narrative, which is what it is a guide to.
    anchor_idx = None
    for i, par in enumerate(paras):
        if _heading_at(i) and par.text.strip().lower().startswith(_NARRATIVE_HEADING):
            anchor_idx = i
            break
    if anchor_idx is None:
        for i, par in enumerate(paras):
            if _heading_at(i) and i > 0:
                anchor_idx = i
                break

    heading_style = paras[start].style if start is not None else (
        paras[anchor_idx].style if anchor_idx is not None else None)
    # Prefer the document's own body style; fall back to whatever the prose beside the block
    # uses, which on a pandoc render is the post-heading style rather than a running-text one.
    body_style = None
    try:
        body_style = doc.styles["Body Text"]
    except KeyError:
        for i in range(((end if start is not None else anchor_idx) or 0), len(paras)):
            if not _heading_at(i) and paras[i].text.strip():
                body_style = paras[i].style
                break

    if start is not None:
        for par in list(paras[start:end]):
            par._element.getparent().remove(par._element)
    if not block_md.strip():
        doc.save(str(path))
        return {"top_sources": 0, "had_existing_block": start is not None}

    # Re-resolve the anchor: the deletion above invalidates the earlier snapshot's indices.
    before_el = None
    for par in doc.paragraphs:
        st = par.style.name if par.style is not None else ""
        if is_heading_style(st) and par.text.strip().lower().startswith(_NARRATIVE_HEADING):
            before_el = par._element
            break
    if before_el is None:
        for i, par in enumerate(doc.paragraphs):
            st = par.style.name if par.style is not None else ""
            if is_heading_style(st) and i > 0:
                before_el = par._element
                break
    if before_el is None:
        return {"top_sources": 0, "had_existing_block": start is not None,
                "error": "no narrative heading to insert before"}

    n = 0
    for line in block_md.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            par = doc.add_paragraph()
            if heading_style is not None:
                par.style = heading_style
            par.add_run(_strip_md(line[3:]))
        elif line.startswith("- "):
            par = doc.add_paragraph()
            if body_style is not None:
                par.style = body_style
            par.add_run("\u2022  " + _strip_md(line[2:]))
            n += 1
        else:
            par = doc.add_paragraph()
            if body_style is not None:
                par.style = body_style
            par.add_run(_strip_md(line)).italic = True
        before_el.addprevious(par._element)
    doc.save(str(path))
    return {"top_sources": n, "had_existing_block": start is not None}


# ── reply comments ───────────────────────────────────────────────────────────────
# rabbitHole's threaded replies are written by the shared writer, haarpi.redline.add_replies
# (called from revise._reply_to_comments). It shadows each parent comment's exact anchor and
# sets w15:paraIdParent, so a reply nests as a real thread even on a paragraph carrying several
# comments — and it creates commentsExtended.xml when the document lacks it. The two functions
# that used to live here (add_reply_comments / _patch_comments_extended) were the inferior
# second implementation: they anchored to the whole parent paragraph rather than the parent's
# span, and gave up when there was no threading part. They are gone.


# ── orchestration ──────────────────────────────────────────────────────────────

def _next_heading_element(paras, idx: int):
    """The heading that opens the next section after ``idx`` — the element a section grafted
    into ``idx``'s section is inserted before.

    Insertion is anchored to the FOLLOWING heading rather than to a computed offset, so several
    sections grafted in one pass stack in call order ahead of the same boundary instead of
    interleaving with each other's paragraphs. Returns None when ``idx`` is in the last section.
    """
    from .docxio import is_section_heading
    for j in range(idx + 1, len(paras)):
        style = paras[j].style.name if paras[j].style is not None else ""
        # H2 only — the level a litreview SECTION uses. `is_heading_style` is also true of the
        # document title and of the bibliography's H3 sub-headings, neither of which bounds a
        # section, and stopping at one of those would graft into the middle of the bibliography.
        if is_section_heading(style):
            return paras[j]._p
    return None


def apply_edits(src: Path, out: Path, edits: list[dict], author: str = "rabbitHole") -> dict:
    """Apply paragraph edits to a copy of ``src`` and write ``out``.

    Each edit: ``{"para": int, "op": "replace"|"insert_after", "text": str}``. Pure XML
    surgery — no LLM, no network. Returns a small summary dict.
    """
    doc = Document(str(src))
    ids = _Ids(_max_existing_id(doc))
    paras = doc.paragraphs  # snapshot: holds the original <w:p> elements by index
    applied = {"replace": 0, "insert_after": 0, "insert_section": 0, "skipped": 0}
    # Replaces first (they don't change paragraph count), then insertions — and because we
    # index into the snapshot's element objects (not a re-read), later inserts never
    # invalidate earlier indices. That is what lets a single pass mix in-place edits with
    # whole grafted sections: nothing moves until every edit has been decided.
    rank = {"replace": 0, "insert_after": 1, "insert_section": 2}
    for e in sorted(edits, key=lambda e: (rank.get(e["op"], 0), e["para"])):
        p_el = paras[e["para"]]._p
        if e["op"] == "insert_after":
            tracked_insert_after(p_el, e["text"], author, ids)
            applied["insert_after"] += 1
        elif e["op"] == "insert_section":
            before = _next_heading_element(paras, e["para"])
            if before is not None:
                _engine.tracked_insert_section(e["heading"], e["paras"], author, ids,
                                               before_el=before)
            elif paras:
                _engine.tracked_insert_section(e["heading"], e["paras"], author, ids,
                                               after_el=paras[-1]._p)
            else:
                applied["skipped"] += 1
                continue
            applied["insert_section"] += 1
        else:
            ok = tracked_replace_sentencewise(p_el, e["text"], author, ids)
            applied["replace" if ok else "skipped"] += 1
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    applied["comments_preserved"] = len(comments_by_id(out))
    return applied
