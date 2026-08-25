"""rabbitHole `graft` — add a section to a reviewed litreview without disturbing the rest.

The verb between `revise` and `report`. A reviewer who asks for a new theme used to get a whole
new document: any single `section` need routed to `report`, which re-plans the review's sections
and regenerates every one of them. That cost the reviewer a second full read of a 27-page
document to see two new sections, and it silently discarded every comment thread on it, because
`report` writes a fresh .docx with no anchors. The rework was scaled to the heaviest ask in the
set rather than to the ask itself.

`graft` drafts only the requested strand and inserts it into a COPY of the reviewer's own .docx
as a tracked insertion. Existing paragraphs are never passed to a model and never rewritten, so
they come through byte-identical and every comment thread survives — the reviewer reads a diff,
not a document. That is also what makes the diff mean anything to the redline engine: a document
that changes 100% every cycle gives it nothing to anchor against.

Where the section goes, in priority order:

  1. the anchor of the comment that asked for it — where the reviewer wrote is a statement about
     where the ask belongs, and deference is owed to it;
  2. the section whose claim is nearest by embedding, when no usable anchor survives;
  3. the end of the review — always available, never silent. Under a through-line requirement a
     tacked-on section is a real defect, so this one is reported rather than quietly taken.
"""

from __future__ import annotations

import shutil
import sys
import time
from pathlib import Path

from . import config, docxio, runlog
from .brain import Brain
from .models import Candidate
from .summarize import (
    Section, SYNTH_SYS, _assemble, _compact_lines, _cosine, _draft_section, _full_lines,
    _make_citekeys, _polish_section, _shortlist, sections_from_markdown,
)

# The insertion is authored as the tool, which is what puts it inside `redline.TOOL_AUTHORS` —
# so the redline engine treats the grafted prose as tool-authored (editable next cycle) rather
# than as an author span it must preserve verbatim.
_AUTHOR = "rabbitHole"

_SECTION_ASK_SYS = ("You plan ONE new section for an existing literature review. "
                    "Respond with ONLY a JSON array, nothing else.")

_SECTION_ASK_PROMPT = """\
Review topic: {topic}
Focus: {focus}

The review already has these sections, in order:
{existing}

A reviewer asked for this to be added:
{asks}

Plan the NEW section(s) only — never restate or re-plan a section above.

- Each names ONE idea in at most 6 words. Never join concepts with a comma, "and", or "/".
- Each carries a claim: one sentence saying what it argues about THIS topic.
- If the ask is already covered by an existing section, return an empty array.

Return ONLY a JSON array:
[{{"heading": "...", "claim": "..."}}]"""


def _plan_new_sections(brain: Brain, cfg, existing: list[Section], asks: list[str],
                       max_new: int = 3) -> list[Section]:
    """Plan only what the reviewer asked for, told what already exists so it cannot duplicate it."""
    from .summarize import _parse_json_list
    prompt = _SECTION_ASK_PROMPT.format(
        topic=cfg.topic, focus=cfg.focus or "",
        existing="\n".join(f"- {s.heading}" for s in existing) or "(none)",
        asks="\n".join(f"- {a}" for a in asks))
    try:
        raw = brain.coordinator(prompt, _SECTION_ASK_SYS, num_ctx=16384, think=False)
        items = _parse_json_list(raw)
    except Exception as e:  # noqa: BLE001
        print(f"[graft] could not plan the new section ({e})", file=sys.stderr)
        return []
    out: list[Section] = []
    have = {s.heading.strip().lower() for s in existing}
    for it in items:
        if not isinstance(it, dict):
            continue
        h = str(it.get("heading", "")).strip()
        c = str(it.get("claim", "")).strip()
        if h and c and h.lower() not in have:
            out.append(Section(heading=h, claim=c))
            have.add(h.lower())
    return out[:max_new]


def choose_position(brain: Brain, existing: list[Section], new: Section,
                    anchor: int | None) -> tuple[int, str]:
    """Where the new section goes: ``(index_to_insert_after, why)``.

    ``anchor`` is the section index a requesting comment sat in, or None when there was none.

    A comment in the FRONT MATTER (anchor -1: the title block, the focus line, the metrics line)
    is document-scoped, not positional — reviewers put general asks at the top of a document.
    Reading it as "open the review with this" would dress a guess up as deference and, for a set
    of three such asks, stack all three ahead of the section that grounds the review. It falls
    through to the embedding instead, which at least places each one near what it relates to.
    """
    if anchor is not None and 0 <= anchor < len(existing):
        return anchor, "the comment's own anchor"
    try:
        vecs = brain.embed_batch([f"{s.heading}. {s.claim or s.text[:400]}" for s in existing]
                                 + [f"{new.heading}. {new.claim}"])
        target, rest = vecs[-1], vecs[:-1]
        sims = [_cosine(v, target) for v in rest]
        if sims:
            return max(range(len(sims)), key=lambda i: sims[i]), "nearest existing section"
    except Exception:  # noqa: BLE001
        pass
    return len(existing) - 1, "appended at the end — NO position signal survived"


def _insert_section(docx_path: Path, after_index: int, sec: Section, author: str) -> bool:
    """Insert one drafted section into the .docx as a tracked insertion.

    Every paragraph already in the document is left exactly as it is — this only adds. That is
    the whole contract: the reviewer's copy, plus a strand, with their comments still on it.
    """
    try:
        from docx import Document
        from haarpi import redline as hredline
    except ImportError:
        return False
    doc = Document(str(docx_path))
    ids = hredline.ids_for(doc)
    # Counted with docxio's SECTION test, not redline's heading test: the two must index the same
    # list `sections_from_markdown` produces, and redline's is also true of the document title.
    heads = [p for p in doc.paragraphs
             if docxio.is_section_heading(p.style.name if p.style is not None else "")]
    if 0 <= after_index < len(heads):
        # everything from the NEXT heading onward belongs after the graft
        nxt = heads[after_index + 1] if after_index + 1 < len(heads) else None
        anchor_el = nxt._p if nxt is not None else None
    elif after_index < 0:
        anchor_el = heads[0]._p if heads else None
    else:
        anchor_el = None

    paras = [p for p in sec.text.split("\n\n") if p.strip()]
    if not paras:
        return False
    if anchor_el is not None:
        head_el = hredline.tracked_heading_before(anchor_el, sec.heading, author, ids)
        prev = head_el
        for para in paras:
            prev = hredline.tracked_insert_after(prev, para.strip(), author, ids)
    else:
        # No following heading: append after the document's last paragraph.
        last = doc.paragraphs[-1]._p if doc.paragraphs else None
        if last is None:
            return False
        prev = hredline.tracked_insert_after(last, sec.heading, author, ids)
        _style_as_heading(prev)
        for para in paras:
            prev = hredline.tracked_insert_after(prev, para.strip(), author, ids)
    doc.save(str(docx_path))
    return True


def _style_as_heading(p_el, style: str = "Heading2") -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        p_el.insert(0, ppr)
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), style)
    ppr.insert(0, pstyle)


def _base_markdown(docx: Path, paths) -> Path | None:
    """The markdown of the draft an annotated .docx was made from.

    The reviewer's copy carries their initials (``..._ra_DCR.docx``) and has no markdown of its
    own — the markdown belongs to the tool's draft it was made from (``..._ra.md``). Walk the
    initials chain back until a markdown turns up, and look in ``output/old/`` too, since a
    superseded cycle's draft is archived there while its markup stays in play.
    """
    stem = docx.stem
    seen: list[Path] = []
    while stem:
        for d in (docx.parent, paths.output, paths.output / "old"):
            cand = d / f"{stem}.md"
            seen.append(cand)
            if cand.is_file():
                return cand
        if "_" not in stem:
            break
        stem = stem.rsplit("_", 1)[0]      # drop one initials segment and retry
    print(f"[graft] looked for: {', '.join(str(p) for p in dict.fromkeys(seen))}",
          file=sys.stderr)
    return None


def _section_asks(comments: list[dict]) -> list[tuple[str, int]]:
    """Reviewer comments that ask for a new section, with the section index each sits in."""
    from haarpi.planner import _SECTION_ASK
    return [(c["text"], c.get("section", -1)) for c in comments
            if _SECTION_ASK.search(c.get("text", ""))]


def run(directory: str = ".", brain_override: str | None = None,
        file: str | None = None) -> int:
    """CLI entry: add the reviewer's requested section(s) to their annotated .docx."""
    t0 = runlog.start()          # the run clock every shared helper's stamp() reads
    cfg = config.load_project(directory)
    gc = config.load_global()
    paths = config.project_paths(directory)
    docx = Path(file) if file else docxio.find_annotated_docx(paths)
    if docx is None or not Path(docx).is_file():
        print("[graft] no annotated .docx found — nothing to graft onto.", file=sys.stderr)
        return 1
    md = _base_markdown(docx, paths)
    if md is None:
        print(f"[graft] no markdown found for {docx.name} — graft needs the draft the markup "
              "was made from to know the existing sections.", file=sys.stderr)
        return 1

    comments = docxio.read_comments(docx)
    asks = _section_asks(comments)
    if not asks:
        print("[graft] no comment asks for a new section — nothing to do.", file=sys.stderr)
        return 1

    narrative = md.read_text(encoding="utf-8")
    if "## Narrative Review" in narrative:
        narrative = narrative.split("## Narrative Review", 1)[1]
    existing = sections_from_markdown(narrative)
    if not existing:
        print("[graft] could not recover the review's sections from its markdown.",
              file=sys.stderr)
        return 1

    from .revise import _load_corpus, _load_notes
    corpus: list[Candidate] = _load_corpus(paths)
    citekeys = _make_citekeys(corpus)
    notes = _load_notes(paths, corpus, citekeys)
    compact = _compact_lines(corpus, notes, citekeys)
    full = _full_lines(corpus, notes, citekeys)
    corpus_keys = set(citekeys.values())

    brain = Brain(cfg.brain, gc, backend_override=brain_override)
    print(f"  {runlog.stamp()}[graft] {docx.name}", flush=True)
    print(f"  {runlog.stamp()}[graft] base markdown: {md.name} "
          f"({'output/old' if md.parent.name == 'old' else 'output'})", flush=True)
    print(f"  {runlog.stamp()}[graft] {len(existing)} existing section(s); "
          f"{len(comments)} comment(s), {len(asks)} section ask(s)", flush=True)
    print(f"  {runlog.stamp()}[graft] planning the new section(s)...", flush=True)
    new = _plan_new_sections(brain, cfg, existing, [a for a, _ in asks])
    if not new:
        print("[graft] the ask is already covered by an existing section — nothing added.",
              file=sys.stderr)
        return 1

    print(f"  {runlog.stamp()}[graft] planned {len(new)}: "
          f"{', '.join(repr(s.heading) for s in new)}", flush=True)
    # Shortlist over the new sections ONLY; the existing ones keep the evidence they have.
    print(f"  {runlog.stamp()}[graft] shortlisting evidence over {len(corpus)} sources...",
          flush=True)
    _shortlist(brain, new, compact, full)
    out_docx = paths.output / f"{docx.stem}_ra.docx"
    shutil.copyfile(docx, out_docx)

    placed: list[tuple[Section, int, str]] = []
    for i, sec in enumerate(new):
        anchor = asks[i][1] if i < len(asks) else None
        # Timestamp the START of each expensive call, not just its end: a start line with no
        # successor is what makes a stall visible. Drafting and peer review run with the
        # coordinator's chain-of-thought ON (they are judgement work), so each section is three
        # slow calls — roughly 3.5h — and silence between them is normal, not a hang.
        t_sec = time.time()
        print(f"  {runlog.stamp()}[graft] §{i + 1}/{len(new)} drafting {sec.heading!r} "
              f"({len(sec.candidates)} candidate sources)...", flush=True)
        sec.text = _draft_section(brain, cfg, new, i, full, SYNTH_SYS)
        print(f"  {runlog.stamp()}[graft] §{i + 1} drafted {len(sec.text.split())}w in "
              f"{runlog.fmt_dt(time.time() - t_sec)}; polishing...", flush=True)
        t_pol = time.time()
        sec.text = _polish_section(brain, cfg, new, i, full, SYNTH_SYS, corpus_keys)
        print(f"  {runlog.stamp()}[graft] §{i + 1} polished to {len(sec.text.split())}w in "
              f"{runlog.fmt_dt(time.time() - t_pol)}", flush=True)
        at, why = choose_position(brain, existing, sec, anchor)
        after = repr(existing[at].heading) if 0 <= at < len(existing) else "the top"
        print(f"  {runlog.stamp()}[graft] §{i + 1} placing after {after} — {why}", flush=True)
        if not _insert_section(out_docx, at, sec, _AUTHOR):
            print(f"  [warn] could not insert {sec.heading!r} into the .docx", file=sys.stderr)
            continue
        existing.insert(at + 1, sec)
        placed.append((sec, at, why))
        print(f"  {runlog.stamp()}[graft] §{i + 1} inserted as a tracked change "
              f"({runlog.fmt_dt(time.time() - t_sec)} for the section)", flush=True)
        if "NO position signal" in why:
            print(f"  [warn] {sec.heading!r} was appended at the end — no anchor and no "
                  "embedding signal. Check its placement.", file=sys.stderr)

    if not placed:
        out_docx.unlink(missing_ok=True)
        print("[graft] nothing was inserted.", file=sys.stderr)
        return 1

    out_md = out_docx.with_suffix(".md")
    out_md.write_text(_assemble(existing), encoding="utf-8")
    print(f"  {runlog.stamp()}[graft] wrote {out_md.name} "
          f"({len(existing)} sections)", flush=True)

    print()
    print("=" * 60)
    print(f" graft complete  [{runlog.fmt_dt(time.time() - t0)}]")
    print("=" * 60)
    for sec, at, why in placed:
        where = existing[at].heading if 0 <= at < len(existing) else "the top"
        print(f"  + {sec.heading}   (after {where!r} — {why})")
    print(f"  Review (docx): {out_docx}")
    print(f"  Review (md)  : {out_md}")
    print("  Existing sections were not passed to a model and are unchanged; every comment "
          "thread is still on the document.")
    return 0
