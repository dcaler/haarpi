"""Shared docx annotation parsing for `revise` and `parseNplan`.

Reads reviewer annotations out of an annotated .docx: comments, tracked
insertions/deletions, and plain body text. Also locates the newest
user-annotated file in output/ (the one whose trailing suffix is not `ra`).

Naming convention (see ~/.claude/CLAUDE.md):
  {YYMMDD}_{project}_litreview_ra.docx           tool output
  {YYMMDD}_{project}_litreview_ra_DCR.docx       user annotated
  {YYMMDD}_{project}_litreview_ra_DCR_ra.docx    tool revised
A file is "user-annotated" when its last '_'-separated segment is NOT `ra`.
"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


def require_docx() -> None:
    if not _DOCX_OK:
        print("[error] python-docx is required: pip install python-docx", file=sys.stderr)
        raise SystemExit(1)


def read_comments(path: Path) -> list[dict]:
    doc = Document(str(path))
    comments = []
    try:
        for rel in doc.part.rels.values():
            if "comments" not in rel.reltype.lower():
                continue
            for c in rel.target_part._element.findall(".//" + qn("w:comment")):
                author = c.get(qn("w:author"), "reviewer")
                texts = [t.text for t in c.findall(".//" + qn("w:t")) if t.text]
                if texts:
                    comments.append({"author": author, "text": " ".join(texts)})
            break
    except Exception:  # noqa: BLE001
        pass
    return comments


def read_track_changes(path: Path) -> dict:
    doc = Document(str(path))
    body = doc.element.body
    insertions, deletions = [], []
    for ins in body.iter(qn("w:ins")):
        text = "".join(t.text or "" for t in ins.iter(qn("w:t")))
        if text.strip():
            insertions.append(text)
    for dele in body.iter(qn("w:del")):
        text = "".join(t.text or "" for t in dele.iter(qn("w:delText")))
        if text.strip():
            deletions.append(text)
    return {"insertions": insertions, "deletions": deletions}


def read_body_text(path: Path) -> str:
    """Body text with tracked insertions visible, deletions gone (python-docx default)."""
    doc = Document(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def build_revision_context(path: Path) -> str:
    comments = read_comments(path)
    changes = read_track_changes(path)
    parts = []
    if changes["deletions"]:
        lines = "\n".join(f"  - {d}" for d in changes["deletions"])
        parts.append(f"DELETED TEXT (the reviewer removed these):\n{lines}")
    if changes["insertions"]:
        lines = "\n".join(f"  + {i}" for i in changes["insertions"])
        parts.append(f"INSERTED TEXT (the reviewer added these):\n{lines}")
    if comments:
        lines = "\n".join(f"  [{c['author']}]: {c['text']}" for c in comments)
        parts.append(f"REVIEWER COMMENTS:\n{lines}")
    return "\n\n".join(parts)


def find_annotated_docx(paths) -> Path | None:
    """Newest docx in output/ that was annotated by the user.

    rabbitHole's own outputs end in `_ra`; user-annotated files end in their
    initials (e.g. `_DCR`). We pick the most recently modified file whose last
    '_'-separated segment is not `ra`, so we never re-process our own output.
    """
    candidates = [p for p in paths.output.glob("*.docx")
                  if p.stem.rsplit("_", 1)[-1] != "ra"]
    if not candidates:
        return None
    return max(candidates, key=lambda p: p.stat().st_mtime)
