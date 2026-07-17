"""The slate: how the author says where the paper is going.

raconteur's brainstorm will not always find the venue you have in mind. When it doesn't, you
add a row — and that row is a decision, not a suggestion. The tool may propose candidates
beside it; it may never touch, downgrade, or delete it.
"""

from __future__ import annotations

import types

from raconteur import slate
from raconteur.config import VenueConfig

ANALYSIS = """\
## Recommendation
**Primary target: ISMIR.**

## Venue slate

Set `status` to selected for each venue you intend to write for.

| slug | venue | kind | status | source |
|------|-------|------|--------|--------|
| ismir | ISMIR | conference | candidate | (raconteur) |
| jasss | JASSS | journal | selected | (raconteur) |
| nime | NIME 2027 | conference | selected | https://nime.org/2027/cfp |
"""


def test_the_slate_is_read_back_out_of_the_analysis():
    got = slate.parse(ANALYSIS)
    assert set(got) == {"ismir", "jasss", "nime"}
    assert got["jasss"]["status"] == "selected"
    assert got["ismir"]["status"] == "candidate"
    assert got["nime"]["url"] == "https://nime.org/2027/cfp"
    assert got["nime"]["kind"] == "conference"


def test_a_venue_the_tool_never_proposed_is_the_authors_own():
    """The case that has to work: raconteur's brainstorm missed NIME, so the author typed
    it in."""
    venues = {"ismir": VenueConfig(name="ISMIR", origin="raconteur")}
    merged = slate.merge(venues, slate.parse(ANALYSIS))
    assert merged["nime"].by_author
    assert merged["nime"].status == "selected"
    assert merged["nime"].url == "https://nime.org/2027/cfp"


def test_the_slate_moves_a_status_and_nothing_else():
    """It is a decision surface, not a config editor: it may not overwrite a spec."""
    venues = {"jasss": VenueConfig(name="JASSS", status="candidate", page_limit=20,
                                   sources={"page_limit": "cfp"})}
    merged = slate.merge(venues, slate.parse(ANALYSIS))
    assert merged["jasss"].status == "selected"
    assert merged["jasss"].page_limit == 20, "the spec is untouched"
    assert merged["jasss"].sources["page_limit"] == "cfp"


def test_a_row_the_author_struck_is_reported_not_obeyed_silently():
    venues = {"ismir": VenueConfig(name="ISMIR"), "cmj": VenueConfig(name="CMJ")}
    assert slate.dropped(venues, slate.parse(ANALYSIS)) == ["cmj"]


def test_no_slate_means_no_change():
    venues = {"ismir": VenueConfig(name="ISMIR", status="selected")}
    assert slate.parse("## Recommendation\nNo table here.") == {}
    assert slate.merge(venues, {})["ismir"].status == "selected"


def test_a_bolded_cell_from_word_still_parses():
    """The author edits this table in Word; pandoc brings bold and stray spaces back."""
    got = slate.parse("## Venue slate\n\n| slug | venue | status |\n|--|--|--|\n"
                      "| **nime** | *NIME 2027* |  **selected**  |\n")
    assert got["nime"]["status"] == "selected"


def test_the_slate_renders_what_the_author_will_edit():
    text = slate.render({
        "ismir": VenueConfig(name="ISMIR", kind="conference"),
        "nime": VenueConfig(name="NIME 2027", status="selected", origin="author",
                            url="https://nime.org/cfp"),
    })
    assert slate.SLATE_HEADING in text
    assert "| ismir | ISMIR | conference | candidate |" in text
    assert "https://nime.org/cfp" in text
    assert slate.parse(text)["nime"]["status"] == "selected", "round-trips"


# ── the specs come from the venue, or they are unknown ───────────────────────

def test_a_venue_with_no_url_keeps_an_unknown_format():
    v = VenueConfig(name="NIME 2027")
    out = slate.fetch_specs(v, brain=None)
    assert out.page_limit is None and not out.sources


def test_the_cfp_is_the_source_and_says_so(monkeypatch):
    import raconteur.web as web
    monkeypatch.setattr(web, "fetch_page_text",
                        lambda url, email, max_chars=6000: "Papers may be up to 8 pages.")

    brain = types.SimpleNamespace(
        coordinator=lambda *a, **k: '{"page_limit": 8, "columns": 2}')
    v = VenueConfig(name="ISMIR", url="https://ismir.net/cfp")
    out = slate.fetch_specs(v, brain)

    assert out.page_limit == 8 and out.columns == 2
    assert out.sources["page_limit"] == "cfp"
    assert "from the call for papers" in out.spec_line("page_limit")


def test_an_unreadable_cfp_leaves_the_format_unknown(monkeypatch):
    import raconteur.web as web
    monkeypatch.setattr(web, "fetch_page_text", lambda *a, **k: "")
    v = VenueConfig(name="X", url="https://example.invalid/cfp")
    out = slate.fetch_specs(v, brain=None)
    assert out.page_limit is None and not out.sources


def test_the_writer_is_warned_when_nothing_was_verified():
    block = slate.specs_block(VenueConfig(name="NIME 2027", status="selected"))
    assert "unknown" in block
    assert "none of this venue's format has been read" in block


def test_a_verified_spec_carries_no_warning():
    v = VenueConfig(name="ISMIR", page_limit=8, sources={"page_limit": "cfp"})
    block = slate.specs_block(v)
    assert "from the call for papers" in block
    assert "not been read" not in block


# ── resolving which venue a verb is for ──────────────────────────────────────

def _cfg(venues):
    from raconteur.config import ProjectConfig
    return ProjectConfig(short_title="Chords", venues=venues)


def test_one_selected_venue_needs_no_flag():
    cfg = _cfg({"ismir": VenueConfig(name="ISMIR", status="selected")})
    assert slate.resolve(cfg) == "ismir"


def test_several_selected_venues_must_be_disambiguated():
    cfg = _cfg({"ismir": VenueConfig(name="ISMIR", status="selected"),
                "jasss": VenueConfig(name="JASSS", status="selected")})
    with __import__("pytest").raises(SystemExit):
        slate.resolve(cfg)
    assert slate.resolve(cfg, "jasss") == "jasss"


def test_no_venue_chosen_is_not_an_error():
    """A project that has not reached the venue analysis still writes."""
    assert slate.resolve(_cfg({})) == ""


def test_an_unknown_venue_is_refused():
    cfg = _cfg({"ismir": VenueConfig(name="ISMIR", status="selected")})
    with __import__("pytest").raises(SystemExit):
        slate.resolve(cfg, "nosuchvenue")


# ── reading the candidates out of an existing analysis ───────────────────────

def test_the_candidates_are_read_out_of_the_analysis():
    """The whole extraction ran behind a broad `except Exception`, so a NameError in it
    looked exactly like "the model returned nothing" — and the slate came back empty."""
    from raconteur import venue as venue_mod

    brain = types.SimpleNamespace(worker=lambda *a, **k: """[
      {"venue": "Journal of Artificial Societies and Social Simulation", "type": "journal"},
      {"venue": "ISMIR", "type": "conference", "url": "https://ismir.net/cfp"}
    ]""")
    got = venue_mod._candidates_from(brain, "…any analysis text…")

    assert set(got) == {"jasss", "ismir"}
    assert got["ismir"].url == "https://ismir.net/cfp"
    assert got["jasss"].kind == "journal"
    assert all(v.status == "candidate" for v in got.values()), "the tool never selects"
    assert all(v.origin == "raconteur" for v in got.values())


def test_a_model_that_returns_junk_yields_no_candidates():
    from raconteur import venue as venue_mod
    brain = types.SimpleNamespace(worker=lambda *a, **k: "I could not find any venues.")
    assert venue_mod._candidates_from(brain, "…") == {}


# ── the slate is a WORD TABLE, and that is where the author edits it ─────────

def _slate_docx(tmp_path, rows, struck=(), tracked=()):
    """A .docx whose slate is a real Word table — what pandoc actually renders."""
    from docx import Document
    from docx.oxml.ns import qn
    from haarpi.redline import _ins

    d = Document()
    d.add_paragraph("Venue slate", style="Heading 2")
    t = d.add_table(rows=1, cols=5)
    for i, h in enumerate(("slug", "venue", "kind", "status", "source")):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
    for i in struck:                       # the author deleted this row in Word
        row = t.rows[i]
        trpr = row._tr.makeelement(qn("w:trPr"), {})
        trpr.append(row._tr.makeelement(qn("w:del"),
                                        {qn("w:id"): "9", qn("w:author"): "DCR",
                                         qn("w:date"): "2026-07-14T00:00:00Z"}))
        row._tr.insert(0, trpr)
    for i, col, text in tracked:           # the author typed into this cell
        cell = t.rows[i].cells[col]
        cell.text = ""
        cell.paragraphs[0]._p.append(_ins(text, "DCR", 7))
    path = tmp_path / "slate.docx"
    d.save(str(path))
    return path


def test_the_slate_is_read_out_of_the_word_table(tmp_path):
    """Rendered to .docx the slate is a TABLE, and a table is not paragraphs — every reader
    in this codebase walks doc.paragraphs, so a slate parsed from the document's text comes
    back EMPTY and the author's decision goes silently unread."""
    doc = _slate_docx(tmp_path, [
        ("ismir", "ISMIR", "conference", "candidate", "(raconteur)"),
        ("jasss", "JASSS", "journal", "candidate", "(raconteur)"),
    ])
    got = slate.parse_docx(doc)
    assert set(got) == {"ismir", "jasss"}

    from raconteur.revise import read_text
    assert slate.parse(read_text(doc)) == {}, "the paragraph text never sees the table"


def test_a_status_the_author_typed_is_read_as_tracked_text(tmp_path):
    """They type "selected" over "candidate"; Word records it as a tracked insertion."""
    doc = _slate_docx(tmp_path, [
        ("ismir", "ISMIR", "conference", "candidate", "(raconteur)"),
    ], tracked=[(1, 3, "selected")])
    assert slate.parse_docx(doc)["ismir"]["status"] == "selected"


def test_a_row_the_author_struck_in_word_is_gone(tmp_path):
    doc = _slate_docx(tmp_path, [
        ("ismir", "ISMIR", "conference", "selected", "(raconteur)"),
        ("cmj", "Computer Music Journal", "journal", "candidate", "(raconteur)"),
    ], struck=[2])
    got = slate.parse_docx(doc)
    assert set(got) == {"ismir"}, "a struck row is a row they removed"


def test_a_row_the_author_added_names_a_venue_we_never_proposed(tmp_path):
    doc = _slate_docx(tmp_path, [
        ("ismir", "ISMIR", "conference", "candidate", "(raconteur)"),
        ("nime", "NIME 2027", "conference", "selected", "https://nime.org/2027/cfp"),
    ])
    merged = slate.merge({"ismir": VenueConfig(name="ISMIR")}, slate.parse_docx(doc))
    assert merged["nime"].by_author and merged["nime"].status == "selected"
    assert merged["nime"].url == "https://nime.org/2027/cfp"


# ── a call for papers is very often a PDF ────────────────────────────────────

def test_a_pdf_cfp_is_read_as_a_pdf(monkeypatch):
    """`…/CSS2026_CFP_Draft.docx-2.pdf` is a real CFP URL. Handed a PDF, an HTML stripper
    returns the mangled remains of a binary stream — from which a model will confidently
    extract a page limit that is not in the document."""
    from raconteur import web

    class _R:
        status_code = 200
        headers = {"content-type": "application/pdf"}
        content = b"%PDF-1.4 ..."
        text = "\x00\x01garbage"

    class _C:
        def __enter__(self): return self
        def __exit__(self, *a): return False
        def get(self, url, timeout=None): return _R()

    monkeypatch.setattr(web, "_client", lambda email: _C())
    monkeypatch.setattr(web, "_pdf_text",
                        lambda data, mx: "Papers should be 3,000 to 5,000 words.")

    assert "5,000 words" in web.fetch_page_text("https://x/cfp.pdf", "")


def test_a_pdf_is_detected_by_its_bytes_not_its_extension(monkeypatch):
    from raconteur import web

    class _R:
        status_code = 200
        headers = {"content-type": "application/octet-stream"}   # server lies
        content = b"%PDF-1.7 ..."
        text = "junk"

    class _C:
        def __enter__(self): return self
        def __exit__(self, *a): return False
        def get(self, url, timeout=None): return _R()

    monkeypatch.setattr(web, "_client", lambda email: _C())
    monkeypatch.setattr(web, "_pdf_text", lambda data, mx: "read as pdf")
    assert web.fetch_page_text("https://x/whatever", "") == "read as pdf"


# ── the widened fetch: content structure, anonymity, and the template link ────
# A conference CFP mandates more than a format — required sections, double-blind,
# a template to submit in. The first two shape the WRITING (so they ride into the
# specs block the outline/draft read); the template link only pre-fills the human
# fetch task, so it stays OUT of the writer's block.

def test_the_fetch_reads_content_rules_and_the_template_link(monkeypatch):
    import raconteur.web as web
    monkeypatch.setattr(web, "fetch_page_text", lambda *a, **k: "double-blind; needs CCS")
    brain = types.SimpleNamespace(coordinator=lambda *a, **k: (
        '{"required_sections": "CCS concepts and keywords", "anonymized": true, '
        '"template_url": "https://acm.org/kit.zip", "template_kind": "latex-acm"}'))
    out = slate.fetch_specs(VenueConfig(name="CHI", url="https://chi.acm.org/cfp"), brain)
    assert out.required_sections == "CCS concepts and keywords"
    assert out.anonymized is True
    assert out.template_url == "https://acm.org/kit.zip"
    assert out.template_kind == "latex-acm"
    for f in ("required_sections", "anonymized", "template_url", "template_kind"):
        assert out.sources[f] == "cfp"


def test_a_stated_single_blind_is_false_not_unknown(monkeypatch):
    import raconteur.web as web
    monkeypatch.setattr(web, "fetch_page_text", lambda *a, **k: "author names allowed")
    brain = types.SimpleNamespace(coordinator=lambda *a, **k: '{"anonymized": false}')
    out = slate.fetch_specs(VenueConfig(name="V", url="https://v/cfp"), brain)
    assert out.anonymized is False and out.sources["anonymized"] == "cfp"


def test_the_writer_block_carries_content_rules_and_the_blind_directive():
    v = VenueConfig(name="CHI", required_sections="CCS concepts and keywords",
                    anonymized=True, sources={"required_sections": "cfp"})
    block = slate.specs_block(v)
    assert "CCS concepts and keywords" in block
    assert "DOUBLE-BLIND" in block


def test_a_named_but_non_blind_venue_gets_no_blind_directive():
    block = slate.specs_block(VenueConfig(name="V", anonymized=False,
                                          page_limit=8, sources={"page_limit": "cfp"}))
    assert "DOUBLE-BLIND" not in block


def test_the_template_link_is_not_shown_to_the_writer():
    v = VenueConfig(name="CHI", template_url="https://acm.org/kit.zip",
                    page_limit=8, sources={"page_limit": "cfp"})
    assert "kit.zip" not in slate.specs_block(v)


# ── the template brief: the human's scaffolded fetch instructions ─────────────

def test_the_brief_prefills_a_known_template_link_and_kind():
    v = VenueConfig(name="CHI", template_url="https://acm.org/kit.zip",
                    template_kind="latex-acm")
    brief = slate.template_brief(v, "paper/templates/chi")
    assert "https://acm.org/kit.zip" in brief and "latex-acm" in brief
    assert "paper/templates/chi/" in brief


def test_the_brief_falls_back_to_the_cfp_when_no_template_link():
    v = VenueConfig(name="CSS2026", url="https://css/cfp")
    brief = slate.template_brief(v, "paper/templates/css2026")
    assert "https://css/cfp" in brief
    assert "empty" in brief          # 'no template needed' is a valid outcome


def test_the_brief_asks_for_the_blinded_variant_when_double_blind():
    v = VenueConfig(name="CHI", url="https://chi/cfp", anonymized=True)
    assert "DOUBLE-BLIND" in slate.template_brief(v, "paper/templates/chi")
