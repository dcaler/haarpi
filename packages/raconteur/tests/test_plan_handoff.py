"""The word plan crossing the skeleton→outline gate, on the document itself.

The plan the author approves has to live where they approve it. It rides as one comment per
section heading, survives the mint because the mint stopped erasing comments, and is read
back by the next rung — which trusts exactly one number in it and checks the rest.
"""

from __future__ import annotations

import inspect
import tempfile
from pathlib import Path

import pytest
from docx import Document

from haarpi import redline
from raconteur import guards, outline, skeleton
from raconteur.config import ProjectConfig, VenueConfig

SKELETON = ("# T\n\n## Abstract\n\n## Introduction\n\n## Background\n\n### A\n\n### B\n\n"
            "## Methods\n\n### C\n\n### D\n\n### E\n\n## Results\n\n### F\n\n### G\n\n"
            "### H\n\n## Discussion\n\n### I\n\n### J\n\n### K\n\n## Conclusion\n\n"
            "## Acknowledgements\n\n## References\n")


@pytest.fixture
def release(tmp_path):
    proj = tmp_path / "proj"
    (proj / "paper").mkdir(parents=True)
    work = proj / "paper" / "skeleton"
    work.mkdir()
    cfg = ProjectConfig(short_title="myproj", title="T")
    cfg.venues = {"css2026": VenueConfig(name="CSS2026", word_min=3000, word_limit=5000,
                                         status="selected")}
    skeleton._write(proj, cfg, work, SKELETON, venue="css2026")
    markup = next(work.glob("*skeleton_ra.docx"))
    out = proj / "out" / "260721_myproj_css2026_skeleton.docx"
    redline.mint_release(markup, out, md_sibling=False)
    return out


def test_the_plan_survives_the_mint_and_reads_back(release):
    wpb, problems = skeleton.plan_from_release(release)
    assert problems == []
    assert wpb == {"Introduction": 150, "Background": 150, "Methods": 150,
                   "Results": 166, "Discussion": 150, "Conclusion": 150}


def test_a_structure_edited_without_its_plan_is_refused_not_guessed(release):
    """The one thing standing between an authoritative comment and silent drift. The author
    adds a Methods subsection and does not touch the comment; the plan now describes a
    structure that no longer exists, which is worse than no plan because it looks like one."""
    doc = Document(str(release))
    for p in doc.paragraphs:
        if p.text.strip() == "E":
            new = doc.add_paragraph("E2", style="Heading 3")
            p._p.addnext(new._p)
            break
    drifted = release.parent / "drifted.docx"
    doc.save(str(drifted))
    redline._carry_comment_parts(release, drifted)

    wpb, problems = skeleton.plan_from_release(drifted)
    assert "Methods" not in wpb, "a stale plan must never be trusted"
    assert any("3 subsection(s)" in p and "now has 4" in p for p in problems)


def test_a_section_with_no_plan_is_reported(release):
    doc = Document(str(release))
    doc.add_paragraph("Appendix", style="Heading 2")
    extra = release.parent / "extra.docx"
    doc.save(str(extra))
    redline._carry_comment_parts(release, extra)
    _, problems = skeleton.plan_from_release(extra)
    assert any("Appendix" in p and "no word plan" in p for p in problems)


def test_furniture_needs_no_plan(release):
    _, problems = skeleton.plan_from_release(release)
    assert not any(h in " ".join(problems)
                   for h in ("Abstract", "Acknowledgements", "References"))


def test_the_plan_is_consumed_not_carried_forward(release):
    """Comments describe a skeleton. Carried into the outline they would describe a document
    that has moved on."""
    assert redline.comment_threads(release)
    skeleton.strip_plan_comments(release)
    assert redline.comment_threads(release) == {}
    body = Document(str(release)).element.body
    from docx.oxml.ns import qn
    assert not list(body.iter(qn("w:commentRangeStart")))


def test_an_unresolved_plan_comment_blocks_the_gate(release):
    """Resolving is the acknowledgement. A plan nobody read is a plan nobody approved."""
    proj = release.parent.parent
    markup = next((proj / "paper" / "skeleton").glob("*skeleton_ra.docx"))
    check = redline.gate_check(markup)
    assert not check["clean"]
    assert all(c["author"] == "raconteur" for c in check["unresolved"])
    assert len(check["unresolved"]) == 6


# ── the mint reconciles ──────────────────────────────────────────────────────
# The comments are written against the structure as GENERATED; the author then moves
# subsections. Rather than make the next rung refuse — turning a routine edit into an error
# cleared by hand-editing six comments — the mint updates the plan to match what was
# approved. WORDS PER BULLET is carried, never recomputed.

def _add_subsection(path, after, name):
    doc = Document(str(path))
    for p in doc.paragraphs:
        if p.text.strip() == after:
            new = doc.add_paragraph(name, style="Heading 3")
            p._p.addnext(new._p)
            break
    doc.save(str(path))
    redline._carry_comment_parts(path, path)


def test_the_mint_updates_the_plan_to_the_approved_structure(tmp_path):
    proj = tmp_path / "p"
    (proj / "paper").mkdir(parents=True)
    work = proj / "paper" / "skeleton"
    work.mkdir()
    cfg = ProjectConfig(short_title="myproj", title="T")
    cfg.venues = {"css2026": VenueConfig(name="CSS2026", word_min=3000, word_limit=5000,
                                         status="selected")}
    skeleton._write(proj, cfg, work, SKELETON, venue="css2026")
    markup = next(work.glob("*skeleton_ra.docx"))

    _add_subsection(markup, "E", "E2")            # Methods 3 -> 4 subsections
    rel = proj / "out" / "rel.docx"
    redline.mint_release(markup, rel, md_sibling=False, post=skeleton.reconcile_plan)

    wpb, problems = skeleton.plan_from_release(rel)
    assert problems == [], "the release must agree with itself"
    assert wpb["Methods"] == 150, "the rate is carried, not re-derived from the share"
    texts = [c["text"] for c in redline.comments_by_id(rel).values()]
    assert any("1200 words · 4 sub · 8 bullets · 150 each" in t for t in texts)
    assert any("Discussion — 900 words · 3 sub" in t for t in texts), "others untouched"


def test_the_rate_can_never_pin_below_one_paragraph():
    """Two ways to starve a section, both closed. Re-deriving from the share after a
    subsection is added gives 900//8 = 112; and a section GENERATED wider than its share
    affords would pin that number at birth — Background came back with four subsections on
    a 600-word share and pinned 75. The floor holds the rate at one paragraph, so the
    section states its true cost and merging restores the share instead of halving it."""
    secs = [(2, "Methods")] + [(3, f"S{i}") for i in range(4)]
    unfloored = (guards.section_words("Methods", 4000)
                 // (guards.MIN_BULLETS_PER_SUBSECTION * 4))
    assert unfloored == 112, "what an unfloored derivation would give"
    assert skeleton.words_per_bullet(secs, 4000)["Methods"] == guards.WORDS_PER_PARAGRAPH


def test_merging_a_too_wide_section_restores_its_share():
    """The behaviour that makes the finding actionable rather than punitive."""
    wide = [(2, "Background")] + [(3, f"S{i}") for i in range(4)]
    merged = [(2, "Background")] + [(3, f"S{i}") for i in range(2)]
    rate = skeleton.words_per_bullet(wide, 4000)["Background"]
    assert rate == 150, "not 75 — a bad generation must not pin a starved rate"
    assert "1200 words · 4 sub" in skeleton.plan_row("Background", 4, rate)
    assert skeleton.words_per_bullet(merged, 4000)["Background"] == 150
    assert "600 words · 2 sub" in skeleton.plan_row("Background", 2, 150)


def test_the_model_is_told_the_allowance_not_left_to_derive_it():
    """Background came back with four subsections because the prompt stated the BULLET floor
    (100) in a sentence about SUBSECTIONS. 600/4 = 150 cleared that bar; the arithmetic it
    skipped was that four subsections is eight bullets, so 75 words a paragraph."""
    from raconteur import outline as _o
    src = inspect.getsource(_o._budget_block)
    assert "leaf_allowance" in src and "affords" in src
    prompt = skeleton._SKELETON_PROMPT
    assert "{min_words}" in prompt and "at least two paragraphs" in prompt
    assert "min_words=guards.subsection_words()" in inspect.getsource(skeleton)
    assert "min_words=guards.PARAGRAPH_BAND[0]" not in inspect.getsource(skeleton)


def test_the_planner_reconciles_only_the_skeleton_rung():
    import inspect
    from haarpi import planner
    src = inspect.getsource(planner)
    assert 'if deliverable in ("skeleton", "outline"):' in src
    assert "from raconteur.skeleton import reconcile_plan as post" in src
    assert "from raconteur.outline import reconcile_plan as post" in src
    assert "mint_release(" in src
    # The whole paper ladder is off the markdown sibling now — package reads the release
    # .docx through read_release, and takes the abstract off the same file. A second copy
    # of an approved contract can only drift from the one that matters.
    assert 'md_sibling=stage != "paper"' in src


# ── the rate reaches the draft ───────────────────────────────────────────────
# The band is where the pinned rate finally does its work. Assuming a flat
# WORDS_PER_PARAGRAPH makes Methods and Results indistinguishable to the drafter — six
# paragraphs each — when the rate is the only thing telling them apart.

def _outline_md(sections):
    out = ["# T", ""]
    for sec, n in sections:
        out += [f"## {sec}", ""] + [f"- beat {i}" for i in range(n)] + [""]
    return "\n".join(out)


def test_the_band_uses_the_sections_own_rate():
    md = _outline_md([("Methods", 6), ("Results", 6)])
    rates = {"Methods": 150, "Results": 166}
    assert guards.section_band("Methods", md, 4000, rates=rates) == (720, 1080)
    assert guards.section_band("Results", md, 4000, rates=rates) == (797, 1195)


def test_without_the_rate_two_sections_become_one_shape():
    """What E fixes. Same bullet count, different contracts, identical bands."""
    md = _outline_md([("Methods", 6), ("Results", 6)])
    assert (guards.section_band("Methods", md, 4000)
            == guards.section_band("Results", md, 4000) == (720, 1080))


def test_the_contract_band_is_tighter_than_the_paragraph_band_implies():
    """It is not redundant with wide_paragraphs. Six paragraphs inside PARAGRAPH_BAND allow
    600–1200 words; the contract holds Results to 797–1195 and Methods to 720–1080."""
    md = _outline_md([("Results", 6)])
    lo, hi = guards.section_band("Results", md, 4000, rates={"Results": 166})
    implied_lo = 6 * guards.PARAGRAPH_BAND[0]
    implied_hi = 6 * guards.PARAGRAPH_BAND[1]
    assert implied_lo < lo and hi < implied_hi


def test_the_draft_reads_the_rate_off_the_release_and_says_when_it_cannot():
    import inspect
    from raconteur import paper
    src = inspect.getsource(paper.run)
    assert "rates_from(sibling)" in src
    assert "no approved word plan on the outline release" in src, \
        "a silent fall back to a flat rate is how the plan stops mattering"
    assert "rates=outline_rates" in src


def test_section_lengths_judges_against_the_rate_too():
    import inspect
    assert "rates=rates" in inspect.getsource(guards.section_lengths)


# ── a figure is not a beat ───────────────────────────────────────────────────
# Giving a figure its own bullet meant every count that matters had to remember to exclude
# it, and three forgot: the css2026 outline's plan comment priced five figures as
# paragraphs and reported 5,410 words where the paper owed 4,596. Appended to the bullet
# that introduces it, there is nothing to remember — and the figure cannot be separated
# from the prose that points at it.

FIG = ("- Discuss settling dynamics [@zhangTipping2011]. "
       "[[Figure 2: Distance over tolerance x radius (blue = closer to the phrase). "
       "(results/figures/E1_0.png)]] "
       "[[Figure 3: Time to settle (rounds). (results/figures/E1_1.png)]]\n")


def test_a_bullet_carrying_figures_is_one_beat():
    md = "# T\n\n## Results\n\n### Recovery\n\n- Present the landscape.\n" + FIG
    head = [h for h in guards.parse_outline(md) if h.level == 3][0]
    assert (head.beats, head.figures) == (2, 2)


def test_the_delimiter_makes_the_path_exact_not_guessed():
    """Real captions contain parentheses — "(blue = closer to the phrase)", "(rounds)" —
    so without a boundary the path has to be guessed at. With "]]" closing the figure it is
    simply the last parenthesised group before it."""
    paths = [m.group("path") for m in guards.appended_figures(FIG)]
    assert paths == ["results/figures/E1_0.png", "results/figures/E1_1.png"]


def test_the_form_does_not_collide_with_a_citekey():
    """Single brackets were the obvious choice and the wrong one: "[@citekey]" is one, and
    the draft carries them."""
    assert guards.all_citekeys(FIG) == ["zhangTipping2011"]
    assert len(guards.appended_figures(FIG)) == 2


def test_an_appended_figure_is_still_placed_and_counted():
    md = "# T\n\n## Results\n\n### Recovery\n\n" + FIG
    heads = guards.parse_outline(md)
    assert guards.outline_figures(heads, md) == [
        ("results/figures/E1_0.png", "Recovery"),
        ("results/figures/E1_1.png", "Recovery")]
    from raconteur import paper
    assert paper._outline_figure_count(md) == 2


def test_the_standalone_form_still_reads():
    """An outline written the old way still places its figures; only the new form is asked
    for. A guard that stopped seeing them would report every existing figure unplaced."""
    md = ("# T\n\n## Results\n\n### Recovery\n\n- A beat.\n"
          "- Figure 1: A caption. (results/figures/x.png)\n")
    head = [h for h in guards.parse_outline(md) if h.level == 3][0]
    assert (head.beats, head.figures) == (1, 1)
    assert guards.outline_figures(guards.parse_outline(md), md) == [
        ("results/figures/x.png", "Recovery")]


def test_the_prompt_asks_for_the_appended_form():
    from raconteur import outline as _o
    src = _o._OUTLINE_PROMPT if hasattr(_o, "_OUTLINE_PROMPT") else ""
    joined = src or "".join(v for k, v in vars(_o).items()
                            if k.endswith("PROMPT") and isinstance(v, str))
    assert "APPENDING it IN DOUBLE SQUARE BRACKETS to the bullet" in joined
    assert "A figure is not a beat" in joined


def test_bullet_budget_checks_against_the_plan_not_the_share():
    """It was telling four approved Background subsections they were 75 words a paragraph
    and to halve their bullets — arguing with the plan rather than checking against it."""
    md = "# T\n\n## Background\n\n" + "".join(
        f"### S{i}\n\n- one\n- two\n\n" for i in range(4))
    assert len(guards.bullet_budget(md, 4000, None)) == 4      # share-derived, wrong
    assert guards.bullet_budget(md, 4000, None, {"Background": 150}) == []


def test_the_outline_battery_receives_the_rates():
    import inspect
    from raconteur import outline as _o
    assert "rates" in inspect.signature(guards.outline_findings).parameters
    assert "bullet_budget(markdown, budget, shares, rates)" in inspect.getsource(
        guards.outline_findings)
    assert '"rates": rates or {}' in inspect.getsource(_o._outline_guard_inputs)


def test_both_outline_prompts_ask_for_the_same_figure_form():
    """One asked for the appended form and the other for a standalone line. The model can
    only satisfy one, and which it picks is not predictable."""
    from raconteur import outline as _o
    prompts = [v for k, v in vars(_o).items()
               if k.endswith("PROMPT") and isinstance(v, str) and "figure" in v.lower()]
    assert prompts
    for p in prompts:
        assert "own line" not in p, "the standalone form must be gone from every prompt"
    assert any("DOUBLE SQUARE BRACKETS" in p for p in prompts)


def test_the_draft_prompt_reads_the_bracketed_form():
    """Given an outline carrying [[Figure …]] on a bullet, a drafter told to look for a
    standalone line renders no figures at all — and figure_findings then reports every one
    of them missing."""
    from raconteur import paper
    p = paper._DRAFT_SECTION_PROMPT
    assert "[[Figure: <caption> (<path>)]]" in p
    assert "a line of the form" not in p
    # The marker is a REFERENCE, not a placement instruction: the image follows the
    # paragraph that first points at the figure, which is not always the bullet's own.
    # One bullet in the css2026 outline carries two figures.
    assert "is a REFERENCE to that figure" in p
    assert "immediately \\\nAFTER the paragraph containing its FIRST reference" in p \
        or "AFTER the paragraph containing its FIRST reference" in p


# ── a figure is evidence, and evidence is not optional ───────────────────────
# css2026 lost two: the Methods illustration was never placed, because nothing told the
# model Methods owed one, and a recovery-landscape plot rode a beat that was removed.
# Both were reported and the outline shipped anyway, because the finding was advisory.

class _Fig:
    def __init__(self, path): self.path = path


PLANNED = {"Results": [_Fig("results/figures/E1_0.png"), _Fig("results/figures/E1_1.png")],
           "Methods": [_Fig("illustrations/1Dspace.png")]}


def test_a_figure_never_placed_is_a_fault():
    md = ("# T\n## Methods\n### M1\n- a beat [[Figure 1: c. (illustrations/1Dspace.png)]]\n"
          "## Results\n### R1\n- a beat [[Figure 2: c. (results/figures/E1_0.png)]]\n")
    faults = outline.figure_variance(md, PLANNED)
    assert faults == {"Results": ["E1_1.png: not placed anywhere"]}


def test_a_figure_in_the_wrong_section_is_a_fault():
    """The section is not the model's guess — an author's figure names its own, and
    rayleigh's carry the results."""
    md = ("# T\n## Methods\n### M1\n- a beat\n"
          "## Results\n### R1\n- a beat [[Figure 1: c. (illustrations/1Dspace.png)]]\n")
    faults = outline.figure_variance(md, PLANNED)
    assert any("placed under 'Results'" in x and "puts it in 'Methods'" in x
               for x in faults["Methods"])


def test_a_figure_this_paper_does_not_have_is_a_fault():
    md = "# T\n## Results\n### R1\n- a beat [[Figure 9: c. (results/figures/made-up.png)]]\n"
    assert any("not a figure this paper has" in x
               for x in outline.figure_variance(md, {"Results": []})["Results"])


def test_a_matching_outline_reports_nothing():
    md = ("# T\n## Methods\n### M1\n- a beat [[Figure 1: c. (illustrations/1Dspace.png)]]\n"
          "## Results\n### R1\n- one [[Figure 2: c. (results/figures/E1_0.png)]]\n"
          "- two [[Figure 3: c. (results/figures/E1_1.png)]]\n")
    assert outline.figure_variance(md, PLANNED) == {}


def test_figure_faults_join_the_same_converge_or_block_loop():
    import inspect
    src = inspect.getsource(outline._outline_fresh)
    assert "fig_faults = figure_variance(outline, figures)" in src
    assert "if not variance and not fig_faults:" in src
    assert "the mint is blocked" in src
    assert "_write(project_dir, cfg, paper_dir, outline, venue, rates, divergence, fig_faults)" in src


def test_the_fault_rides_the_comment_that_blocks_the_mint():
    import inspect
    src = inspect.getsource(outline.plan_notes)
    assert "FIGURE NOT AS PLANNED" in src and "mint until you resolve it" in src


# ── the prompt has to fit ────────────────────────────────────────────────────
# Ollama discards the BEGINNING of an over-long prompt, so an overrun does not degrade
# gracefully — it removes whatever sits at the top, silently. The css2026 outline ran at
# 11,146 tokens against a 10,649 budget and produced an outline written half-blind.

def test_a_chosen_venue_drops_the_argument_for_choosing_it(tmp_path):
    """Shortlist, tiers, conference options, recommendation and slate are 81% of the file
    and describe a decision already made. What survives is what the analysis says about the
    PAPER."""
    from raconteur.context import load_venue_analysis
    d = tmp_path / "paper" / "venue"
    d.mkdir(parents=True)
    (d / "venue_analysis.md").write_text(
        "## Research question\nq\n\n## Core novelty claim\nc\n\n## Paper profile\np\n\n"
        "## Venue shortlist\n### Tier 1\nlots\n\n## Recommendation\nmore\n\n"
        "## Venue slate\nyet more\n")
    whole = load_venue_analysis(tmp_path)
    chosen = load_venue_analysis(tmp_path, selected=True)
    assert "Venue shortlist" in whole, "the venue stage still sees all of it"
    assert "Research question" in chosen and "Paper profile" in chosen
    for gone in ("Venue shortlist", "Tier 1", "Recommendation", "Venue slate"):
        assert gone not in chosen


def test_the_heading_match_is_case_insensitive():
    """It shipped case-sensitive against "## Venue shortlist" and was a silent no-op —
    the file came back the same length and nothing said so."""
    from raconteur.context import _VENUE_CHOICE_FROM
    assert _VENUE_CHOICE_FROM.search("## Venue Shortlist\n")
    assert _VENUE_CHOICE_FROM.search("## venue slate\n")


def test_the_plan_names_figures_by_path_not_by_caption():
    """The captions are already in key_figures. Repeating them put 1,044 characters of the
    same text in one prompt twice."""
    import inspect
    src = inspect.getsource(outline._per_subsection_plan)
    assert 'f"      {f.path}"' in src
    assert "{f.caption}" not in src


# ── a beat the plan never granted may not be counted into the plan ───────────

def _outline_doc(tmp_path, body: str):
    """A rendered outline .docx from markdown, with real Word list bullets."""
    import subprocess
    md = tmp_path / "o.md"
    md.write_text(body, encoding="utf-8")
    out = tmp_path / "o.docx"
    subprocess.run(["pandoc", str(md), "-o", str(out)], check=True, capture_output=True)
    return Document(str(out))


def test_a_stray_beat_above_the_first_subsection_is_not_counted(tmp_path):
    """reconcile_plan counts what the document has, and _sections_of used to count EVERY
    list paragraph under an H2 — including beats hanging above its first subsection, which
    the plan never granted. So five figure beats dumped on `## Methods` and `## Results`
    were written INTO the plan rather than measured against it: Results read "1660 words ·
    10 bullets" against a plan of 996 and 6, and the document was planned 410 words past
    the venue ceiling with no comment saying anything was wrong."""
    doc = _outline_doc(tmp_path,
                       "# T\n\n## Results\n\n- a stray beat\n\n### C\n\n- one\n- two\n")
    got = {sec: (subs, bullets) for sec, subs, bullets in outline._sections_of(doc)}
    assert got["Results"] == (["C"], 2)


def test_a_section_with_no_subsections_still_counts_its_own_beats(tmp_path):
    """planned_bullets names the section itself where it has no subsections, so there its
    bullets ARE the plan — the rule is "count where the plan puts one", not "ignore H2s"."""
    doc = _outline_doc(tmp_path, "# T\n\n## Conclusion\n\n- one\n- two\n")
    got = {sec: (subs, bullets) for sec, subs, bullets in outline._sections_of(doc)}
    assert got["Conclusion"] == ([], 2)
