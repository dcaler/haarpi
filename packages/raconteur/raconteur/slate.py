"""The venue slate — how the author says WHERE this paper is going.

An outline and a manuscript are written for a particular venue, so the venue has to be
chosen, and choosing is the author's job. The tool proposes candidates; the author decides.
Same rule as accepting a tracked change and resolving a comment: the pipeline never makes
this call on its own.

The slate is a table at the end of the venue analysis, so the choice is made in the document
the author is already reading, through the redline loop that already exists:

    ## Venue slate

    | slug  | venue                    | status    | source                    |
    |-------|--------------------------|-----------|---------------------------|
    | ismir | ISMIR                    | candidate | (raconteur)               |
    | jasss | JASSS                    | candidate | (raconteur)               |
    | nime  | NIME 2027                | selected  | https://nime.org/2027/cfp |

The last row is one the AUTHOR typed. That is the whole point: raconteur's brainstorm will
not always find the venue you have in mind, and when it doesn't, you add a row. A row you
added carries ``origin: author``, and the tool may never touch, downgrade or delete it —
its candidates are suggestions; your row is a decision.

A ``source`` URL is the call for papers. It matters because a venue's format specs must come
from the VENUE, not from prose the model wrote about the venue: today `_update_yaml` asks an
LLM to read its own analysis and extract a page limit, which for a venue it does not really
know produces either a null or a fiction. With a URL we fetch the real page and read the
real numbers — and where there is neither, the spec stays "unknown" and says so, rather than
being invented.
"""

from __future__ import annotations

import re

from .config import VenueConfig, venue_slug
from .log import log

SLATE_HEADING = "## Venue slate"

SLATE_INSTRUCTIONS = (
    "Set `status` to **selected** for each venue you intend to write for; each selected "
    "venue gets its own outline and manuscript. Add a row for any venue not listed — a row "
    "you add is authoritative and raconteur will never change it. Put the call-for-papers "
    "URL in `source` and the format specs (page limit, columns, citation style) are read "
    "from it; without one they stay unknown, and the writing will not assume them."
)

_STATUSES = ("candidate", "selected", "submitted", "published", "rejected", "declined")

_ROW = re.compile(r"^\s*\|(?P<cells>.+)\|\s*$")
_SEP = re.compile(r"^\s*\|[\s:|-]+\|\s*$")
_URL = re.compile(r"https?://\S+")


def render(venues: dict[str, VenueConfig]) -> str:
    """The slate section, for the end of a venue analysis."""
    lines = [SLATE_HEADING, "", SLATE_INSTRUCTIONS, "",
             "| slug | venue | kind | status | source |",
             "|------|-------|------|--------|--------|"]
    for slug, v in venues.items():
        src = v.url or ("(author)" if v.by_author else "(raconteur)")
        lines.append(f"| {slug} | {v.name} | {v.kind or ''} | {v.status} | {src} |")
    return "\n".join(lines) + "\n"


def _row(cells: list[str]) -> tuple[str, dict] | None:
    """One slate row -> (slug, record). None when it is a header, a rule, or unreadable."""
    cells = [c.strip().strip("*_ ") for c in cells]
    if not cells or not cells[0] or cells[0].lower() == "slug":
        return None
    if set(cells[0]) <= set("-: "):                     # a markdown rule
        return None
    slug = (venue_slug(cells[0]) if " " in cells[0]
            else re.sub(r"[^a-z0-9]", "", cells[0].lower()))
    name = cells[1] if len(cells) > 1 and cells[1] else cells[0]
    if not slug or not name:
        return None
    kind = cells[2].lower() if len(cells) > 2 else ""
    status, url = "", ""
    for c in cells[2:]:
        if c.lower() in _STATUSES:
            status = c.lower()
        if (u := _URL.search(c)):
            url = u.group(0)
    return slug, {
        "name": name,
        "kind": kind if kind in ("journal", "conference", "workshop") else "",
        "status": status or "candidate",
        "url": url,
    }


def parse(text: str) -> dict[str, dict]:
    """Read the slate out of an analysis in MARKDOWN.

    Tolerant on purpose: a missing column, or a stray asterisk from a bold run, must not
    lose the whole slate. A row that cannot be read is skipped, never guessed at.
    """
    body = text.split(SLATE_HEADING, 1)[1] if SLATE_HEADING in text else ""
    if not body:
        return {}
    out: dict[str, dict] = {}
    for line in body.splitlines():
        if line.strip().startswith("#"):
            if out:
                break                      # the next section ends the slate
            continue
        m = _ROW.match(line)
        if not m or _SEP.match(line):
            continue
        got = _row(m.group("cells").split("|"))
        if got:
            out[got[0]] = got[1]
    return out


def parse_docx(path) -> dict[str, dict]:
    """Read the slate out of the author's WORD DOCUMENT — which is where they edit it.

    The slate is rendered to a real Word table, and a Word table is not paragraphs: every
    reader in this codebase walks `doc.paragraphs`, and a table is invisible to all of them.
    A slate parsed from the paragraph text of a .docx comes back EMPTY — the author's
    decision, silently unread.

    Cells are read as ACCEPTED text (their tracked insertions kept, their deletions gone),
    and a row the author struck in Word is a row they removed: Word marks it `w:trPr/w:del`
    and it does not survive here either.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from haarpi.redline import _accepted_para_text

    def _cell(c) -> str:
        return " ".join(_accepted_para_text(p._p).strip() for p in c.paragraphs).strip()

    def _struck(row) -> bool:
        trpr = row._tr.find(qn("w:trPr"))
        return trpr is not None and trpr.find(qn("w:del")) is not None

    for table in Document(str(path)).tables:
        header = [_cell(c).lower() for c in table.rows[0].cells]
        if not header or header[0] != "slug":
            continue
        out: dict[str, dict] = {}
        for row in table.rows[1:]:
            if _struck(row):
                continue
            got = _row([_cell(c) for c in row.cells])
            if got:
                out[got[0]] = got[1]
        return out
    return {}


def merge(venues: dict[str, VenueConfig], slate: dict[str, dict]) -> dict[str, VenueConfig]:
    """Fold the author's slate back into the configured venues.

    A venue in the slate but NOT in the config came from the author's pen — raconteur writes
    its own candidates into the config and the slate together, so anything appearing in the
    slate alone was typed there. It is recorded ``origin: author``, and from then on it is
    theirs: the tool may not touch, downgrade, or delete it.

    An existing venue keeps every spec it has. The slate may move its STATUS — that is what
    the slate is FOR — and supply a URL, and nothing else. raconteur's candidates are
    suggestions; the author's rows are decisions.
    """
    out = dict(venues)
    for slug, row in slate.items():
        existing = out.get(slug)
        if existing is None:
            out[slug] = VenueConfig(
                name=row["name"], kind=row["kind"], status=row["status"], url=row["url"],
                origin="author",
            )
            log(f"[raconteur] venue '{slug}' ({row['name']}) — the author's own row, "
                f"status {row['status']}")
            continue
        if row["status"] and row["status"] != existing.status:
            log(f"[raconteur] venue '{slug}': {existing.status} → {row['status']}")
            existing.status = row["status"]
        if row["url"] and row["url"] != existing.url:
            existing.url = row["url"]
        if row["kind"] and not existing.kind:
            existing.kind = row["kind"]
    return out


_CFP_SYS = (
    "You read a venue's call for papers or author guidelines and extract its SUBMISSION "
    "FORMAT. You report only what the page actually states. A format detail the page does "
    "not give is null — never a typical value, never a value from a venue you know. A "
    "fabricated page limit is worse than no page limit: the author cannot tell it is wrong "
    "until the desk reject."
)

_CFP_PROMPT = """\
Venue: {name}
Source page:
{page}

Return ONLY a JSON object with exactly these keys, using null for anything the page does \
not state:
{{"page_limit": null, "word_limit": null, "abstract_limit": null, "columns": null, \
"citation_style": null, "format_notes": null}}

- page_limit / word_limit / abstract_limit: integers, for the MAIN submission (not the \
camera-ready, not the abstract-only track), excluding references if the page says so.
- columns: 1 or 2.
- citation_style: the named style (e.g. "APA", "IEEE", "Chicago author-date"), if stated.
- format_notes: anything else an author must obey — template, anonymisation, section \
limits, supplementary rules. One short paragraph, or null."""


def fetch_specs(venue: VenueConfig, brain, email: str = "") -> VenueConfig:
    """Read a venue's format off its OWN page, and record that that is where it came from.

    The specs used to be extracted by an LLM from the venue analysis — a document the same
    LLM had just written. For a venue it does not truly know, that produces a plausible
    number with nothing behind it, and a plausible number is indistinguishable from a real
    one at the moment it matters. A URL turns the guess into a reading.
    """
    import json

    from . import web

    if not venue.url:
        return venue
    log(f"[raconteur] reading the call for papers: {venue.url}")
    page = web.fetch_page_text(venue.url, email, max_chars=6000)
    if not page.strip():
        log(f"[warn] could not read {venue.url} — {venue.name}'s format stays unknown")
        return venue
    try:
        raw = brain.coordinator(
            _CFP_PROMPT.format(name=venue.name, page=page), system=_CFP_SYS, num_ctx=16384)
        m = re.search(r"\{.*\}", raw, re.S)
        spec = json.loads(m.group(0)) if m else {}
    except Exception as e:  # noqa: BLE001
        log(f"[warn] could not extract {venue.name}'s format from its CFP ({e})")
        return venue
    if not isinstance(spec, dict):
        return venue

    got = []
    for field_name in VenueConfig.SPEC_FIELDS:
        value = spec.get(field_name)
        if value in (None, "", 0):
            continue
        if field_name in ("page_limit", "word_limit", "abstract_limit", "columns"):
            try:
                value = int(value)
            except (TypeError, ValueError):
                continue
        setattr(venue, field_name, value)
        venue.sources[field_name] = "cfp"      # read, not guessed
        got.append(field_name)
    log(f"[raconteur] {venue.name}: {', '.join(got) if got else 'the CFP states no format'}"
        + (" (from the call for papers)" if got else ""))
    return venue


def resolve(cfg, requested: str = "") -> str:
    """Which venue a deliverable is for. '' when the project has not chosen one.

    Exactly one selected venue needs no flag. SEVERAL is a question only the author can
    answer — writing the ISMIR paper when they meant the JASSS one costs a cycle and is not
    obvious from the output — so we refuse rather than pick.
    """
    if requested:
        slug = requested.strip().lower()
        if slug not in cfg.venues:
            known = ", ".join(sorted(cfg.venues)) or "(none configured)"
            log(f"[error] no venue '{slug}' in paper/raconteur.yaml. Known: {known}")
            raise SystemExit(1)
        return slug

    chosen = cfg.selected_venues()
    if len(chosen) == 1:
        return chosen[0]
    if not chosen:
        return ""                       # no venue chosen yet: a venue-free deliverable
    log("[error] this project targets several venues — say which one:\n"
        + "\n".join(f"    --venue {s}   ({cfg.venues[s].name})" for s in chosen))
    raise SystemExit(1)


def specs_block(venue: VenueConfig | None) -> str:
    """The venue's format, as the writer must see it: with its provenance, or as unknown.

    The one thing this must never do is present an inferred number as a fact. A writer told
    "Page limit: 8" cannot tell whether that was read off the call for papers or invented by
    a model summarising its own prose — and the difference only surfaces at the desk reject.
    """
    if venue is None or not venue.name:
        return ""
    lines = [f"Target venue: {venue.name}"
             + (f" ({venue.kind})" if venue.kind else "")]
    if venue.url:
        lines.append(f"Call for papers: {venue.url}")
    lines += [f"{venue.spec_line(f)}" for f in VenueConfig.SPEC_FIELDS
              if f != "format_notes"]
    if venue.format_notes:
        lines.append(f"Format notes: {venue.format_notes}")
    if not any(venue.sources.get(f) == "cfp" for f in VenueConfig.SPEC_FIELDS):
        lines.append("NOTE: none of this venue's format has been read from its call for "
                     "papers. Write to the argument, not to an assumed length, and do not "
                     "state a limit you cannot verify.")
    return "Venue specifications:\n" + "\n".join(lines) + "\n"


def dropped(venues: dict[str, VenueConfig], slate: dict[str, dict]) -> list[str]:
    """Venues the author struck from the slate — deleted rows.

    Reported, never acted on silently. A struck row means "not this one", and the honest
    reading is `declined`; but a slate the author never touched (or a parse that lost a row)
    would otherwise silently un-select a venue they are mid-way through writing for.
    """
    return [s for s in venues if s not in slate]
