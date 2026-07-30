"""The redline revise path: answer each comment with a minimal, in-place tracked change.

This is what `raconteur paper` does by default when it finds an annotated .docx. The clean
rewrite survives as `paper --resynth`.

Three ideas carry the whole module:

  1. THE ANNOTATION BLOB IS DEAD. `revise.build_revision_context` concatenates every comment
     and tracked change into one string handed to every section with "apply only those
     relevant; ignore the rest". That leaves routing to the model, which is why a comment on
     the Discussion used to rewrite the Methods. Here each comment is anchored to the exact
     sentences it spans, in the paragraph it belongs to, and no other sentence in the
     document is ever shown to the reviser.

  2. THE REVISER RETURNS ONLY THE SENTENCES IT CHANGED, keyed by index. A sentence it does
     not return is copied byte-for-byte — the untouched sentences are literally the original
     strings, so minimality is true by construction rather than hoped for from a diff. That
     also makes the touched set exact, so `minimal_edit_violation` can prove a comment on
     sentence 2 did not rewrite sentence 4.

  3. FAIL CLOSED. Malformed JSON, a dropped citekey, a dropped or invented equation, an
     out-of-scope sentence, an exhausted retry budget — any of these and we write NO tracked
     change and say so in the reply. A broken edit under a reply claiming "done" is the worst
     outcome, worse than a visibly skipped comment.

Guards in Python, judgement in the LLM: everything mechanical is decided here, precisely, and
stated as an imperative. The audit call is left with the one question code cannot answer —
does this edit mean what the comment asked for?
"""

from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from haarpi import redline as hredline
from haarpi import redline_engine as _engine
from haarpi import text as htext
from haarpi.redline_engine import Evidence, ParaContext, route_class_of

from . import guards, redline
from .brain import Brain
from .config import ProjectConfig
from .log import log

# ── the per-paragraph reviser ────────────────────────────────────────────────

_PARA_REVISE_SYS = """\
You are revising ONE paragraph of a scholarly paper to satisfy a reviewer's comment(s) on it.
The paragraph is given to you as NUMBERED SENTENCES. You return only the sentences you
changed, keyed by their number — never the whole paragraph. A sentence you do not return
survives word for word, which is the point: it keeps its citations, its grounding, and its
evidence intact.

Make the SMALLEST change that fully and genuinely addresses every comment. Revise the
sentence(s) the comment bears on. Leave the rest alone.

CITATIONS — cite ONLY sources in the bibliography below, and ALWAYS as a [@citekey] tag using
the exact key shown: write "[@smith2021]", NEVER "Smith (2021)" or "(Smith, 2021)". An
author-year citation is invisible to the bibliography and silently unverifies the claim.
Every [@citekey] in a sentence you rewrite must survive in your version unless a comment asks
you to remove that source.

PLACEHOLDERS — a token like ⟦m:1⟧ stands for an equation in the original. Reproduce it
exactly, in the sentence whose claim it supports. Never retype an equation as prose, never
move one to another sentence, and never invent a placeholder of your own.

THE AUTHOR'S OWN WORDS — a token like ⟦a:1⟧ stands for text the AUTHOR wrote BY HAND. It is
the most expensive text on the page, and it is not yours. Reproduce every ⟦a:N⟧ exactly where
it stands: you may not rewrite it, shorten it, absorb it into a sentence of your own, or drop
it. Write AROUND it. It is a fixed point, and your prose must lead into it and out of it
without repeating what it already says. The legend below gives you their words precisely so
you can do that — read them; never retype them.
If, and only if, their text carries an outright ERROR — a typo, a broken construction — you
may SUGGEST a correction. A suggestion is delivered to the author as a comment and is never
applied. Never suggest a change of style, wording, emphasis, or opinion: disagreeing with the
author is not an error in the author.

OUTPUT — a single JSON object mapping sentence number to its replacement text. Use null to
delete a sentence. Return nothing else: no prose, no commentary, no code fence.
  {"2": "The revised second sentence [@smith2021].", "5": null}
You may additionally include the key "copyedits", mapping an author placeholder to your
suggested correction of THEIR sentence:
  {"2": "…", "copyedits": {"a:1": "…their sentence, with the typo fixed…"}}
If no sentence needs to change, return {}."""

_PARA_REVISE_PROMPT = """\
Paper: {title}
Section: {heading}

PARAGRAPH, as numbered sentences. ▶ marks the sentence(s) the reviewer's comment is anchored
to — those are the ones to revise:
{sentences}
{authored_section}
REVIEWER COMMENT(S) on this paragraph (address every one):
{comments}
{context_section}{bib_section}
Return the JSON object of changed sentences only."""

_AUTHORED_BLOCK = """
THE AUTHOR'S OWN SENTENCES — fixed points. Reproduce the placeholder; never the text:
{legend}
"""

_PARA_AUDIT_SYS = """\
You audit ONE revised paragraph of a scholarly paper against the reviewer comment(s) it was
meant to satisfy. Mechanical checks — citation format, dropped citations, equations, which
sentences were touched — have already been made in code and passed; do not repeat them. Judge
only what code cannot: MEANING.

Respond with EXACTLY one of three things, nothing else:
- "OK" — the revision fully and genuinely addresses every comment.
- A line "ROUTE: <class>: <brief reason>" — a comment that CANNOT be satisfied by editing
  this paragraph's prose. <class> is exactly one of:
      section   — asks for a new section or subsection, or material not belonging here
      sources   — asks for literature or citations not present in the bibliography
      evidence  — asks for a result, statistic, or method that does not exist yet
      figure    — asks for a table, chart, or figure
  Do not accept a prose gesture as satisfying such a request.
- Otherwise a numbered list of specific problems: a comment not really addressed, or
  addressed in name only. Quote the text you mean."""

_PARA_AUDIT_PROMPT = """\
Paper: {title}
Section: {heading}

REVIEWER COMMENT(S) the revision must satisfy:
{comments}

ORIGINAL PARAGRAPH:
{paragraph}

REVISED PARAGRAPH (under audit):
{revised}

Judge only against the comment(s): is each fully and genuinely addressed, and is the comment
even satisfiable by editing this paragraph's prose at all? Respond "OK", or
"ROUTE: <class>: <reason>", or a numbered list."""

# What the reviewer is told when a comment cannot be a redline. Answering a request for a
# figure with "this needs sources not in the bibliography" is a false diagnosis: gathering
# papers will never satisfy it. The class is what makes the reply honest.
_ROUTE_CLASSES = ("section", "sources", "evidence", "figure")

_ROUTE_ADVICE = {
    "section": "cannot be a tracked change — it asks for new structure. "
               "Run 'raconteur outline' to revise the structure, then 'raconteur paper'.",
    "sources": "cannot be a tracked change — it asks for literature not in refs.bib. "
               "Run rabbitHole to gather the sources first.",
    "evidence": "cannot be a tracked change — it asks for a result or method that does not "
                "exist. raconteur cannot manufacture evidence: run rayleigh (results) or "
                "raster (methods) first.",
    "figure": "cannot be a tracked change — it asks for a table or figure. "
              "Produce it in rayleigh, then re-render.",
}


def _route_class(verdict: str) -> str:
    """Extract the class from a "ROUTE: <class>: <reason>" verdict."""
    rest = verdict.split(":", 1)[1] if ":" in verdict else ""
    head = rest.strip().split(":", 1)[0].strip().lower()
    return head if head in _ROUTE_CLASSES else "sources"


def _is_ok(verdict: str) -> bool:
    v = verdict.strip().upper()
    return v == "OK" or v.startswith("OK ") or v.startswith("OK.")


# ── sentence-indexed edits ───────────────────────────────────────────────────

def _number_sentences(units: list[str], anchored: set[int]) -> str:
    """Render the paragraph as numbered sentences, marking those a comment bears on."""
    return "\n".join(
        f"{'▶' if i in anchored else ' '} {i + 1}. {u.strip()}"
        for i, u in enumerate(units))


def _apply_sentence_edits(units: list[str], edits: dict) -> str:
    """Rebuild the paragraph from the original units plus the reviser's replacements.

    Every sentence the reviser did not return is copied byte-for-byte. This is what makes the
    sentence-level redline true by construction rather than hoped for from a diff: the
    untouched sentences are literally the original objects.
    """
    out: list[str] = []
    for i, unit in enumerate(units):
        key = str(i + 1)
        if key not in edits:
            out.append(unit)
            continue
        repl = edits[key]
        if repl is None:
            continue  # deleted
        trailing = unit[len(unit.rstrip()):]  # keep the original inter-sentence spacing
        out.append(repl.strip() + trailing)
    return "".join(out)


def _parse_sentence_edits(raw: str, n_units: int,
                          authored: dict | None = None) -> tuple[dict, dict, list[str]]:
    """Parse the reviser's JSON. Returns (edits, copyedits, errors).

    Strict on purpose: a lenient parser that falls back to an empty dict on malformed output
    would look exactly like a well-formed edit of nothing, and we would write no tracked
    change while replying that the comment was addressed.

    ``copyedits`` are suggested corrections to the AUTHOR'S OWN spans. They are never
    applied — they are delivered as comments — so a copyedit naming a span that does not
    exist is dropped rather than argued with.
    """
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    obj = None
    if m:
        try:
            parsed = json.loads(m.group(0))
            obj = parsed if isinstance(parsed, dict) else None
        except json.JSONDecodeError:
            obj = None
    if obj is None:
        return {}, {}, ['Output was not a JSON object. Return only a JSON object mapping '
                        'sentence number to replacement text, e.g. {"2": "…"}.']

    copyedits: dict[str, str] = {}
    raw_copy = obj.pop("copyedits", None)
    if isinstance(raw_copy, dict):
        for k, v in raw_copy.items():
            key = f"⟦{str(k).strip().strip('⟦⟧')}⟧"
            if isinstance(v, str) and v.strip() and (authored is None or key in authored):
                copyedits[key] = v.strip()

    edits, errors = {}, []
    for k, v in obj.items():
        if not str(k).strip().isdigit() or not (1 <= int(k) <= n_units):
            errors.append(f'"{k}" is not a sentence number between 1 and {n_units}.')
            continue
        if v is not None and not isinstance(v, str):
            errors.append(f'The value for sentence {k} must be text or null.')
            continue
        edits[str(int(k))] = v
    return edits, copyedits, errors


# ── the per-paragraph adversary ──────────────────────────────────────────────

def _para_guard_findings(
    old_text: str, new_text: str, touched: set[int], anchored: set[int],
    n_units: int, kind: str, known: set[str], signature: dict | None = None,
) -> list[guards.Finding]:
    """Everything about a paragraph rewrite that Python can decide precisely.

    Note what is NOT here. Density guards (`uncited_paragraphs`, `sparse_paragraphs`) belong
    to the DRAFT phase: on a redline, collateral change is the defect, and a comment asking
    to tighten a sentence does not license injecting citations into it. Note also the
    section-kind gate: a Methods or Results paragraph is grounded in the writeup, not the
    bibliography, so it is not required to carry a citation at all.

    And note what is left for the audit: whether the edit means what the comment asked for.
    Everything else moved into code, where it is exact.
    """
    findings = (guards.author_year_prose(new_text)
                + guards.dropped_citekeys(old_text, new_text)
                + guards.dropped_sentinels(old_text, new_text)
                + guards.invented_sentinels(old_text, new_text)
                + guards.minimal_edit_violation(touched, anchored, n_units))
    if known:
        findings += guards.unresolved_keys(new_text, known)
    if guards.expects_citations(kind) and not guards.CITE_TAG_RE.search(new_text):
        findings.append(guards.Finding(
            "uncited", "paragraph",
            "The paragraph now cites no source — restore a [@citekey] from the bibliography."))
    # "Match this author's voice" is a wish; a transition he has never written is a finding.
    findings += guards.style_findings(new_text, signature or {})
    return findings


class RaconteurPolicy:
    """raconteur's dialect for the shared redline engine: evidence is a refs.bib slice plus
    the heading's section context, the prompts are the paper house style (with the ⟦a:N⟧
    authored-span contract), and the guards are raconteur's full set. Constructed per
    paragraph, because ``context_section`` is heading-specific.

    ``resolve_named_source`` is None for now: raconteur owns no gatherer, so a comment naming a
    source not in refs.bib routes as a missing source (via the audit) rather than being pulled
    in. It will return a line once raconteur may delegate the pull to rabbitHole.
    """

    author = redline.AUTHOR
    # sources-first: the engine falls back to route_classes[0] for an unrecognised audit
    # class, and raconteur's safe default is "missing sources" (matching the original
    # _route_class). Membership is order-independent; the audit prompt lists them itself.
    route_classes = ("sources", "section", "evidence", "figure")

    def __init__(self, title: str, context_section: str, bib_section: str,
                 known: set[str], signature: dict | None = None):
        self._title = title
        self._context_section = context_section
        self._bib_section = bib_section
        self._known = set(known or ())
        self._signature = signature

    def evidence_for(self, ctx: ParaContext) -> Evidence:
        return Evidence(known=set(self._known), context=self._context_section)

    def resolve_named_source(self, citekey: str):
        return None

    def revise_system(self) -> str:
        return _PARA_REVISE_SYS

    def audit_system(self) -> str:
        return _PARA_AUDIT_SYS

    def revise_user(self, ctx: ParaContext, evidence: Evidence,
                    numbered_sentences: str, comment_block: str) -> str:
        legend = "\n".join(f'  {k} = "{v.strip()}"' for k, v in (ctx.authored or {}).items())
        return _PARA_REVISE_PROMPT.format(
            title=self._title, heading=ctx.heading,
            sentences=numbered_sentences,
            authored_section=_AUTHORED_BLOCK.format(legend=legend) if legend else "",
            comments=comment_block,
            context_section=evidence.context,
            bib_section=self._bib_section,
        )

    def audit_user(self, ctx: ParaContext, revised: str, comment_block: str) -> str:
        return _PARA_AUDIT_PROMPT.format(
            title=self._title, heading=ctx.heading, comments=comment_block,
            paragraph=ctx.text, revised=revised)

    def guard_findings(self, old_text: str, new_text: str, touched: set[int],
                       ctx: ParaContext, evidence: Evidence) -> list:
        n_units = len(guards.sentence_units(ctx.text))
        return _para_guard_findings(
            old_text, new_text, touched, ctx.anchored, n_units, ctx.kind,
            evidence.known, self._signature)


def redline_paragraph(
    brain: Brain, title: str, heading: str, paragraph: str, comments: list[str],
    context_section: str, bib_section: str, anchored: set[int], kind: str,
    known: set[str], rounds: int = 2, authored: dict[str, str] | None = None,
    copyedits: dict[str, str] | None = None, signature: dict | None = None,
) -> tuple[str | None, str]:
    """Rewrite one commented paragraph, through the shared engine, in raconteur's dialect.

    The loop, the primals, and the fail-closed contract now live in
    ``haarpi.redline_engine``; this is the thin adapter that keeps raconteur's call signature
    and outcome vocabulary (``edited`` / ``route:<class>`` / ``skipped``) so the orchestration
    and the adversary tests are unchanged.

    ``authored`` maps ⟦a:N⟧ → the author's exact words. ``copyedits``, if given, is filled with
    any correction the reviser proposes to those spans — DELIVERED AS COMMENTS, never applied.
    """
    policy = RaconteurPolicy(title, context_section, bib_section, known, signature)
    ctx = ParaContext(
        heading=heading, text=paragraph, comments=list(comments),
        anchored=set(anchored or ()), named_keys=set(), kind=kind,
        authored=dict(authored or {}))
    new_text, disposition, produced = _engine.redline_paragraph(
        brain, ctx, policy, rounds=rounds)
    if copyedits is not None:
        copyedits.update(produced)
    cls = route_class_of(disposition)
    outcome = f"route:{cls}" if cls else disposition   # engine "routed:x" → raconteur "route:x"
    return new_text, outcome


# ── orchestration ────────────────────────────────────────────────────────────

def _ask_text(ask: dict) -> str:
    """One open ask, with its thread — the whole of what the reviewer said.

    The thread is part of the ask. Under the comment protocol, new information about an
    unmet ask arrives as a REPLY rather than a second comment, so a reviser that reads
    only the top comment reads only half its instructions. And an ask the tool has
    already answered, still open, is not a fresh request — it is one the tool got wrong,
    and it must not be answered the same way twice.
    """
    parts = [ask["text"].strip()]
    for f in ask.get("followups", []):
        parts.append(f"(reviewer, same thread) {f.strip()}")
    if ask.get("repeat"):
        parts.append("** You already answered this and the reviewer left it OPEN: your "
                     "previous answer did not satisfy them. Do not repeat it. **")
        for r in ask.get("prior_tool_replies", []):
            parts.append(f"(your previous, rejected answer) {r.strip()[:300]}")
    return "\n".join(parts)


def _out_path(paper_dir: Path, short_title: str, user_rev: Path) -> Path:
    """A redline is a MINOR version: it keeps the reviewer's datestamp and extends the chain.

    260709_trust_ra_DCR.docx -> 260709_trust_ra_DCR_ra.docx
    """
    from .naming import parse, minor_name
    parsed = parse(user_rev, short_title)
    chain = parsed[1] if parsed else ["ra"]
    datestamp = parsed[0] if parsed else None
    return paper_dir / minor_name(short_title, chain, "docx", datestamp)


def _write_md_sibling(project_dir: Path, out: Path) -> None:
    """The accepted-text .md next to the markup .docx.

    The docx is the reviewer's redline; the md is what downstream LLM consumers
    bind (load_onepager, the outline the draft reads). Without it they would
    silently pick up the previous cycle's pre-annotation md.
    """
    from docx import Document
    md_path = out.with_suffix(".md")
    md_path.write_text(redline.accepted_markdown(Document(str(out))), encoding="utf-8")
    log(f"[raconteur] wrote {md_path.relative_to(project_dir)} (accepted text)")


def redline_revise(
    project_dir: Path,
    cfg: ProjectConfig,
    brain: Brain,
    paper_dir: Path,
    user_rev: Path,
    litrev: str,
    code: str,
    results: str,
    bib_section: str,
    known: set[str],
    context_fn=None,
    md_sibling: bool = False,
) -> tuple[Path, dict[str, str]]:
    """Edit a COPY of the reviewer's .docx in place, one anchored comment at a time.

    Returns (output_path, dispositions) where dispositions maps comment id -> outcome.
    Silence is not a decision: every comment gets a disposition, and every one that could not
    be answered by a tracked change is reported with the reason it could not.

    `context_fn(heading, paragraph_text) -> str` overrides the paper's heading-keyed
    evidence routing — a deliverable whose structure lives in the paragraph rather than
    the heading (the one-pager's beats) routes on the paragraph text instead.
    """
    from docx import Document
    from .paper import _context_for_section

    out = _out_path(paper_dir, cfg.short_title, user_rev)
    shutil.copy2(user_rev, out)

    # Only the OPEN asks. A resolved comment is history, not an instruction — answering
    # it again rewrites prose the reviewer has already accepted. And a thread is part of
    # its ask: the reviewer's follow-up replies come with it, and an ask the tool has
    # already answered (still open) is a failure to repair, not a fresh request.
    asks = {a["id"]: a for a in hredline.open_asks(out)}
    live = set(asks)
    anchors = redline.comment_anchors(out, only=live)
    cmap = redline.comments_by_id(out)
    headings = redline.heading_comments(out, only=live)

    if not anchors and not headings:
        log("[warn] no comment anchors found in the revision — nothing to redline")
        # The reviewer may still have left tracked changes; the accepted md is
        # how those reach the next stage.
        if md_sibling:
            _write_md_sibling(project_dir, out)
        return out, {}

    doc = Document(str(out))
    ids = redline.ids_for(doc)
    body = {rec["index"]: rec for rec in redline.body_paragraphs(doc)}
    dispositions: dict[str, str] = {}
    replies: dict[str, str] = {}
    copyedit_notes: list[tuple[str, str]] = []
    edited = 0

    for anchor in anchors:
        rec = body.get(anchor["index"])
        if rec is None:
            continue

        # A comment whose anchor lies wholly inside deleted text has lost its subject: the
        # reviewer asked "what IS tonal negotiation?" and then deleted the sentence that
        # said it. Handing that to a model as a live instruction buys a confident rewrite
        # of prose that no longer exists. Say so instead.
        obsolete = hredline.anchors_in_deleted_text(rec["para"]._p) & set(anchor["ids"])
        for cid in obsolete:
            dispositions[cid] = "obsolete"
            replies[cid] = ("The text this comment was anchored to has since been deleted, "
                            "so there is nothing here to revise. Resolve the thread if the "
                            "deletion settles it; re-comment on the new text if it does not.")
        ids_live = [c for c in anchor["ids"] if c not in obsolete]
        comments = [_ask_text(asks[c]) for c in ids_live if c in asks]
        if not comments:
            if obsolete:
                log(f"[raconteur] para {anchor['index']}: {len(obsolete)} comment(s) on "
                    f"text you deleted — answered, not revised")
            continue
        heading = anchor["heading"] or "Abstract"
        if context_fn is not None:
            ctx = context_fn(anchor["heading"] or "", anchor["text"])
        else:
            ctx = _context_for_section(anchor["heading"], litrev, code, results)
        context_section = f"\n{ctx}" if ctx else ""

        # The author's own sentences are atoms: readable, reproducible, untouchable. The
        # tool's prose around them stays fully editable — deference is owed to the text a
        # person wrote, not to the paragraph it happens to sit in.
        authored = anchor.get("authored") or {}
        proposed: dict[str, str] = {}

        log(f"[raconteur] redlining '{heading}' para {anchor['index']} "
            f"({len(comments)} comment(s), anchored to sentence(s) "
            f"{[i + 1 for i in anchor['anchored']] or 'all'}"
            f"{f'; {len(authored)} authored span(s) held fixed' if authored else ''})…")

        new_text, outcome = redline_paragraph(
            brain, cfg.title, heading, anchor["text"], comments,
            context_section, bib_section, set(anchor["anchored"]),
            anchor["kind"], known, authored=authored, copyedits=proposed,
        )

        for span, fix in proposed.items():
            original = authored.get(span, "").strip()
            # Word by word, anchored on the offending words, saying only the correction —
            # never the author's whole sentence handed back with the change buried in it.
            copyedit_notes.extend(htext.copyedit_notes(original, fix.strip()))

        for cid in ids_live:
            dispositions[cid] = outcome

        if outcome == "edited" and new_text:
            if redline.tracked_replace_sentencewise(
                    rec["para"]._p, new_text, redline.AUTHOR, ids, protect_authored=True):
                edited += 1
                log("[raconteur]   → tracked change written"
                    + (f" (around {len(authored)} authored span(s))" if authored else ""))
            else:
                for cid in ids_live:
                    dispositions[cid] = "skipped"
                log("[warn]   → no textual change; nothing written")
        elif outcome.startswith("route:"):
            cls = outcome.split(":", 1)[1]
            log(f"[warn]   → {_ROUTE_ADVICE.get(cls, 'cannot be a tracked change.')}")
        else:
            log("[warn]   → skipped: no verifiable edit could be produced; "
                "the paragraph is unchanged")

    for h in headings:
        for cid in h["ids"]:
            dispositions[cid] = "route:section"
        log(f"[warn] comment on heading '{h['heading']}' — {_ROUTE_ADVICE['section']}")

    doc.save(str(out))
    log(f"[raconteur] wrote {out.relative_to(project_dir)} "
        f"({edited} paragraph(s) redlined)")
    if replies:
        n = hredline.add_replies(out, replies, author=redline.AUTHOR)
        log(f"[raconteur] {n} threaded reply/replies delivered")
    if copyedit_notes:
        # The author's text is theirs. A correction to it is a remark, never an edit.
        n = hredline.add_anchored_comments(out, copyedit_notes, author=redline.AUTHOR)
        log(f"[raconteur] {n} suggested copyedit(s) on the author's own text, "
            f"delivered as comments — not applied")
    _report(dispositions, cmap, known, out)
    if md_sibling:
        _write_md_sibling(project_dir, out)
    return out, dispositions


def _report(dispositions: dict[str, str], cmap: dict[str, dict],
            known: set[str], out: Path) -> None:
    """Every comment gets a reply. Silence is not a decision — a comment neither applied nor
    explicitly declined is a defect in the revise pass itself."""
    if not dispositions:
        return
    log("[raconteur] ── comment dispositions ──")
    counts = {"edited": 0, "obsolete": 0, "routed": 0, "declined": 0}
    for cid, outcome in dispositions.items():
        text = cmap.get(cid, {}).get("text", "?")[:60]
        if outcome == "edited":
            verdict = "applied as a tracked change"
            counts["edited"] += 1
        elif outcome == "obsolete":
            verdict = ("the text it was anchored to has been deleted — answered in a reply, "
                       "nothing to revise")
            counts["obsolete"] += 1
        elif outcome.startswith("route:"):
            verdict = _ROUTE_ADVICE.get(outcome.split(":", 1)[1], "routed")
            counts["routed"] += 1
        else:
            verdict = "DECLINED — no verifiable edit could be produced; paragraph unchanged"
            counts["declined"] += 1
        log(f"[raconteur]   [{cid}] {text!r}: {verdict}")
    log(f"[raconteur] {counts['edited']} applied · {counts['obsolete']} obsolete · "
        f"{counts['routed']} routed · {counts['declined']} declined")

    from docx import Document
    doc = Document(str(out))
    md = redline.accepted_markdown(doc)
    log(f"[raconteur] {guards.metrics(md, known)}")

    cited = set(guards.all_citekeys(md))
    if known:
        unresolved = sorted(cited - known)
        if unresolved:
            log(f"[warn] the redlined text cites {len(unresolved)} key(s) with no refs.bib "
                f"entry: {', '.join('[@' + k + ']' for k in unresolved)}")
    # The References list in the .docx was rendered by pandoc/citeproc at draft time. A
    # redline edits prose only, so a newly cited source has no entry there. Rebuilding a
    # styled bibliography inside OOXML is a separate job; say so rather than pretend.
    ref_keys = _references_keys(doc)
    if ref_keys is not None:
        new_keys = sorted(cited - ref_keys)
        if new_keys:
            log(f"[warn] {len(new_keys)} newly cited source(s) have no entry in the rendered "
                f"References list: {', '.join('[@' + k + ']' for k in new_keys)}")
            log("[warn] the References section is not rebuilt by the redline — re-render "
                "with 'raconteur paper --resynth' if the bibliography must be complete")


def _references_keys(doc) -> set[str] | None:
    """Citekeys already present in the rendered References section.

    Returns None when the document has no References section at all — then there is no
    bibliography to have drifted, and warning about it would be noise.

    Best-effort otherwise: pandoc renders entries as prose, so we look for [@key] text, which
    survives when the draft still carries the tags.
    """
    text: list[str] = []
    found = False
    in_refs = False
    for p in doc.paragraphs:
        if redline.is_heading_style(redline._style_name(p)):
            in_refs = guards.is_references((p.text or "").strip())
            found = found or in_refs
            continue
        if in_refs:
            text.append(p.text or "")
    if not found:
        return None
    return set(guards.all_citekeys("\n".join(text)))
