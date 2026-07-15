from __future__ import annotations
import json
import re
import shutil
from typing import NamedTuple
from haarpi import text as htext

from . import guards
from .log import log
from pathlib import Path
from .brain import Brain
from .config import ProjectConfig, GlobalConfig
from .context import (
    load_litreview, load_methods, load_results, load_bib_summary,
    load_bib_keys, load_style_profile, load_style_signature,
    load_figure_manifest, check_prerequisites,
)
from .naming import major_onepager_name, find_latest, find_user_revision
from .render import to_docx

# The one-pager is the first deliverable: the most concise path through the
# paper's narrative — high notes only, at most two figures. A human edits it,
# and 'outline' uses the approved narrative to design the full paper.
#
# Beats are drafted one at a time, each seeing only the evidence that bears on
# it (cf. paper.py:_context_for_section). The beats are fixed rather than parsed
# from an outline, so the routing is a static table. A whole-document critique
# then closes the through-line: a 1-3 sentence beat is too short to stand alone
# the way a paper section does, so coherence is repaired at the end rather than
# assumed.

_SYSTEM = (
    "You are an expert academic writing assistant. You distil a research project "
    "into the single most concise path through its narrative — the story a reader "
    "must follow, and nothing more."
)

# Priority claims are the reflex to suppress in the beats that position the work.
# A bare "avoid overclaiming" gets ignored, so the banned phrasings and their
# replacements are both named.
_NO_PRIORITY_CLAIMS = (
    '- Make no priority claims. Never write "the first to", "no one has", "novel", '
    '"unprecedented", or "opens an entirely new field". Frame the space as '
    'underexamined instead: "this area is underexamined", "has received little '
    'attention", "remains largely untested".'
)


class _Beat(NamedTuple):
    name: str
    intent: str            # what this beat must convey
    sources: tuple[str, ...]   # bulky evidence routed to it
    drop: tuple[str, ...]      # analysis keys withheld from it
    rules: str = ""            # extra constraints


# Approach is given key_design and denied key_findings: it describes how the
# experiments were built, never what they returned.
_BEATS: list[_Beat] = [
    _Beat(
        "Motivation",
        "why this problem matters — the stakes, and what it costs to leave the "
        "problem unaddressed",
        ("description", "litrev"),
        ("key_findings", "key_design", "key_equations"),
    ),
    _Beat(
        "Gap",
        "where this work sits in the literature: what the existing literature has "
        "and has not examined. This is the beat that situates the research product "
        "among what has come before",
        ("description", "litrev", "results"),
        ("key_design", "key_equations"),
        _NO_PRIORITY_CLAIMS,
    ),
    _Beat(
        "Approach",
        "the core idea and method, named specifically, including how the experiments "
        "are designed",
        ("description", "code"),
        ("key_findings",),
        "- Describe the design of the work only. Report no outcomes, values, or "
        "findings — those belong to the Key result(s) beat.",
    ),
    _Beat(
        "Key result(s)",
        "the one or two findings that carry the paper, with the concrete values that "
        "support them",
        ("results",),
        ("key_design", "key_equations"),
        "- Report every value the way the results report it. If you quote a BEST case — a "
        "maximum, a peak, the band where the effect appears — say what it was best OUT OF: "
        "the mean, the range, the sweep it came from. A maximum stated alone is a maximum "
        "hiding its distribution, and the first reviewer to open the results file will see "
        "that it was doing so.",
    ),
    _Beat(
        "Implication",
        "what this work changes or enables",
        ("description", "litrev"),
        ("key_design", "key_equations"),
        "- Return explicitly to the gaps named in the Gap beat above and say how "
        "this work narrows them. Do not raise a new gap here.\n"
        + _NO_PRIORITY_CLAIMS,
    ),
]

# Which kind of the author's prose each beat is. A Results paragraph and a Discussion
# paragraph are not written in the same voice, even by the same person.
_BEAT_KIND = {
    "Motivation": "litrev",
    "Gap": "litrev",
    "Approach": "methods",
    "Key result(s)": "results",
    "Implication": "discussion",
}

# Figures come from rayleigh's results dir, so results are the only beat that can
# carry one — and it inherits the whole two-figure budget.
_FIGURE_BEAT = "Key result(s)"

_FIGURE_RULE = (
    "- You may embed AT MOST TWO figures, and only those that carry the argument. "
    "Choose paths only from the figure list above; do not invent paths. Omit figures "
    "entirely if none is essential.\n"
    "- EVERY figure must be INTRODUCED IN THE TEXT before it appears: a sentence that says "
    "what the reader should see in it (\"Figure 1 shows …\", \"…, as Figure 2 makes "
    "plain\"). A figure the prose never mentions is a figure the reader is never told to "
    "look at.\n"
    "- Number figures in the order they appear, from 1, and write the caption as "
    "\"Figure N: …\" on its own line, in this exact markdown form:\n"
    "  ![Figure 1: what it plots, on what axes, and what the colours mean](figure/path)\n"
    "- The caption must be INFORMATIVE — enough for a reader to interpret the figure "
    "without the text: what is plotted, on which axes, what the colour encoding means, and "
    "what to look for. Adapt the description given with each figure above; do not "
    "invent axes, units, or colours it does not mention."
)
_NO_FIGURE_RULE = "- Do not embed a figure in this beat."

# On a RE-CUT the images are already in the .docx and the count is not the writer's to
# choose. The permissive rule above ("omit figures entirely if none is essential") is right
# for a blank page and wrong here: taken up, it leaves two embedded figures unnumbered,
# uncaptioned and unmentioned, and the beat still passes every check.
_RECUT_FIGURE_RULE = """\
- This document ALREADY CONTAINS EXACTLY {n} figure(s), embedded in this order. You are \
RE-CAPTIONING what is already there. You may not add one and you may not drop one. Their \
captions currently read:
{current}
- Write EXACTLY {n} caption line(s), in that order, each on its own line, in this exact \
markdown form:
  ![Figure 1: what it plots, on what axes, and what the colours mean](figure/path)
  Take the path from the figure list above — match each figure by what its current caption \
says it shows. Do not invent a path.
- INTRODUCE each figure in the prose BEFORE its caption line: a sentence saying what the \
reader should see in it ("Figure 1 shows …", "…, as Figure 2 makes plain"). A figure the \
prose never mentions is a figure the reader is never told to look at.
- Each caption must be INFORMATIVE enough to read the figure by WITHOUT the surrounding \
text: what is plotted, on which axes, what the colour encoding means, and what to look for. \
Build it from the description given with that figure above — do not invent axes, units, or \
colours it does not mention."""

_BEAT_PROMPT = """\
Write the **{beat}** beat of a one-pager: the most concise path through this paper's \
narrative.

This beat must convey: {intent}

Title: {title}
Topic: {topic}
Focus: {focus}
{venue_section}{style_section}Structural analysis:
{analysis}

{evidence_section}{recut_section}{authored_section}{preceding_section}{figure_section}
Rules:
- Aim for 1-3 sentences. This is a beat, not a section — high notes only. But an \
instruction from the author OUTRANKS this: if answering their annotation needs a fourth \
sentence — a definition they asked for, a term they want explained — write it. Never drop \
an annotation to stay short. If an annotation genuinely cannot be honoured in this beat, \
say so plainly in one sentence beginning "CONFLICT:" instead of silently ignoring it.
- Derive every claim from the content above. No generic academic filler: if the \
content does not support a claim, do not make it.
- Do not repeat what the preceding beats already said. Carry the through-line \
forward from them.
- Write prose only. Do not output the beat label, a heading, or a bullet.
{figure_rule}{extra_rules}
- Output only the beat text — no preamble, no closing remarks.
"""

_AUTHORED_BLOCK = """\
THE AUTHOR'S OWN SENTENCES in this beat — they wrote these BY HAND, and they are FIXED:
{legend}

Reproduce every placeholder above EXACTLY as it appears (⟦a:1⟧, and so on), once each, in \
its place in the beat. You may not rewrite, shorten, merge, or drop what the author wrote — \
it is the most expensive text on the page and it is not yours. Write AROUND it: your prose \
must lead into it and out of it, and must not repeat what it already says. The legend gives \
you their exact words so that you can do this coherently; read them, never retype them.

"""

_RECUT_BLOCK = """\
This is a RE-CUT: the author rejected the previous narrative. The previous \
{beat} beat read:
{prior}

The author's annotations on it:
{notes}

Honour every annotation. Re-derive the beat from the evidence above — do not \
patch the old wording. Keep only what the author explicitly marked as good.

Every [@citekey] in the previous beat must SURVIVE in your re-cut, unless an annotation \
asks you to drop that source. A re-cut that loses a citation is rejected outright, and \
the beat is left as it stands.

"""

_RECUT_GENERAL = """\
The author also left these overall annotations on the previous one-pager \
(they apply to every beat):
{notes}

"""

_CRITIQUE_PROMPT = """\
Critique this one-pager. It must be the most concise path through the paper's \
narrative — high notes only, a single unbroken through-line.

Check for:
- Beats that repeat each other, or restate the same claim in different words.
- Breaks in the through-line: a beat that does not follow from the one before it.
- The Implication beat failing to return to the gaps the Gap beat named.
- Priority claims — "the first to", "no one has", "novel", "unprecedented". The \
work must be positioned as entering an underexamined area, not as unprecedented.
- Generic academic filler: any sentence that could appear in any paper in this field.
- Claims the structural analysis does not support.
- Total length over ~500 words, or more than two embedded figures.

One-pager:
{onepager}

List the specific problems as a numbered list. If there are none, say so.
Output only the list."""

_REVISE_PROMPT = """\
Revise this one-pager to address every point in the critique.

One-pager:
{onepager}

Critique:
{critique}

Rules:
- Keep the beat structure: each beat is a "## <beat name>" heading followed by 1-3 \
sentences of plain prose. No bold text.
- Cut every sentence that is not essential to the through-line. Do not exceed ~500 words.
- Keep concrete values. Do not add new claims, and do not add figures — keep at most two.
- Output only the revised one-pager — no preamble."""

# ── helpers ───────────────────────────────────────────────────────────────────

def _style_block(style_profile: str) -> str:
    """The voice block is self-describing now — a palette and real prose, not 'guidance'."""
    return f"{style_profile}\n\n" if style_profile else ""


def _ensure_style(project_dir: Path, cfg: ProjectConfig) -> None:
    """Style is required: ensure an author voice profile exists before drafting.

    Trains it now if missing (non-interactive when the author's papers were
    confirmed at init). Hard-errors if it cannot be produced.
    """
    from .style import STYLE_PROFILE_PATH
    if not STYLE_PROFILE_PATH.exists():
        log("[raconteur] style profile required and missing — training now…")
        from .style import run as style_run
        style_run(project_dir)
        if not STYLE_PROFILE_PATH.exists():
            log("[error] style profile could not be created — configure Zotero and "
                "run 'raconteur style', then retry")
            raise SystemExit(1)
    # Style is required, so keep it applied through every downstream stage.
    if not cfg.use_style:
        cfg.use_style = True
        cfg.save(project_dir)


def _figure_section(figures) -> str:
    """The figures, WITH what rayleigh says each one shows.

    A path alone tells the writer nothing — it cannot name an axis it has never seen, so it
    invents one. rayleigh's caption names the axes, the colour encoding, and what to look
    for; that is the raw material for a caption a reader can actually use.
    """
    if not figures:
        return ""
    lines = []
    for f in figures:
        path = f.path if hasattr(f, "path") else str(f)
        caption = getattr(f, "caption", "")
        lines.append(f"- {path}\n    what it shows: {caption}" if caption
                     else f"- {path}\n    what it shows: (rayleigh recorded no description "
                          f"— do not invent one; describe only what the results support)")
    return ("Available figures (embed at most two, only if essential, using these exact "
            "paths):\n" + "\n".join(lines) + "\n")


def _evidence_for_beat(
    sources: tuple[str, ...], description: str, litrev: str, code: str, results: str
) -> str:
    """The evidence bundle for one beat — only what bears on it."""
    parts = []
    if "description" in sources and description:
        parts.append(
            "Research description (the author's own framing of the project — the "
            f"source of record for why this work matters):\n{description}"
        )
    if "litrev" in sources and litrev:
        parts.append(f"Literature review:\n{litrev}")
    if "code" in sources and code:
        parts.append(f"Methods (raster writeup):\n{code}")
    if "results" in sources and results:
        parts.append(f"Results content (rayleigh):\n{results}")
    return ("\n\n".join(parts) + "\n\n") if parts else ""


def _preceding_block(beats: list[tuple[str, str]],
                     has_authored: set[str] = frozenset()) -> str:
    if not beats:
        return ""
    written = "\n".join(
        f"**{name}**"
        f"{' (carries sentences the author wrote BY HAND — fixed points)' if name in has_authored else ''}"
        f" — {text}"
        for name, text in beats)
    return (
        "Beats already written (continue from these; do not repeat them):\n"
        f"{written}\n\n"
    )


_LABEL_DASHES = "—-–:"


def _strip_label(name: str, text: str) -> str:
    """Drop a beat label the model emitted despite being told not to."""
    t = text.strip()
    for prefix in (f"- **{name}**", f"**{name}**", f"## {name}", f"- {name}", name):
        if t.lower().startswith(prefix.lower()):
            t = t[len(prefix):].lstrip().lstrip(_LABEL_DASHES).lstrip()
            break
    return t


_FIG_LINE = re.compile(r"[ \t]*(!\[[^\]]*\]\([^)]+\))[ \t]*$", re.M)


def _assemble(beats: list[tuple[str, str]]) -> str:
    """Each beat under its own section heading, figures standing as their own blocks.

    Headings, not inline bold labels: a bold lead-in run would donate its
    formatting to any tracked replacement of the paragraph, and headings are
    what the redline reader routes by."""
    out = []
    for name, text in beats:
        body = _FIG_LINE.sub(r"\n\n\1\n", text.strip()).strip()
        out.append(f"## {name}\n\n{body}")
    return "\n\n".join(out)


def _write(project_dir: Path, cfg: ProjectConfig, paper_dir: Path, text: str) -> None:
    output = f"# {cfg.title} — one-pager\n\n{text.strip()}\n"
    out_path = paper_dir / major_onepager_name(cfg.short_title, "md")
    out_path.write_text(output, encoding="utf-8")
    log(f"[raconteur] wrote {out_path.relative_to(project_dir)}")

    bib_path = (project_dir / cfg.litrev_dir / "output" / "refs.bib") if cfg.litrev_dir else None
    docx = to_docx(out_path, bib_path=bib_path, resource_path=project_dir)
    if docx:
        log(f"[raconteur] wrote {docx.relative_to(project_dir)}")


_VERIFY_SYS = """\
You check whether a reviewer's ask was ACTUALLY satisfied by the delivered text. Be
adversarial. A gesture toward the ask is not the ask.

  * "define X" is satisfied only if the text now DEFINES X. Using X again is not defining
    it. Deleting X does not define it either — it removes the question without answering it.
  * "add Y" is satisfied only if Y is present AND NEW. If Y was already in the previous
    text, adding nothing did not address the ask, and claiming to have added it is false.
  * A vaguer or shorter sentence is not an explanation.

Respond with EXACTLY one line, and nothing else:
  SATISFIED: <the words in the delivered text that satisfy it, quoted>
  NOT SATISFIED: <what is still missing — one sentence>"""

_VERIFY_PROMPT = """\
The reviewer's ask, on the "{beat}" beat:
{ask}

The words their comment was anchored to — this is what "this", "it" and "that" refer to:
{on}
{deleted}
The beat as it read BEFORE:
{before}

The beat as DELIVERED now:
{after}

Was the ask satisfied?"""

_DELETED_NOTE = ("\nNOTE: the re-cut DELETED the text the comment was anchored to. Deleting "
                 "the words a question was asked about does not answer the question, unless "
                 "the reviewer asked for the deletion.\n")

_OUTCOME_NOTE = {
    "rewritten": "",
    "refused": ("The re-cut of this beat was REJECTED by a check — it would have dropped a "
                "citation, an equation, or one of your own sentences — so the beat is "
                "unchanged. "),
    "untouched": "This beat was not changed. ",
}


def _verify_replies(brain: Brain, asks: dict[str, dict], beat_of: dict[str, str],
                    before: dict[str, str], after: dict[str, str],
                    outcomes: dict[str, str], on: dict[str, str] | None = None,
                    deleted: set[str] | None = None) -> dict[str, str]:
    """One reply per open ask, and every reply is CHECKED against what was delivered.

    Nothing here is taken on trust. The tool once replied "We explicitly added the recovery
    of Beethoven's Fifth to the Approach" about a phrase that was already in the paragraph
    and which it had not touched — a report of a change derived from the intention to make
    it, not from the change. So each ask is put to an adversary along with the text before
    and the text after, and the reply is built from the VERDICT.

    A reply that says "not addressed" is worth more than a reply that says "addressed" and
    is wrong: the reviewer can act on the first.

    The verdict is SAID, either way. A satisfied verdict used to arrive with the word
    SATISFIED stripped off it, leaving a bare quotation of the delivered text — the one ask
    the tool believed it had met read, to the reviewer, as the tool saying nothing at all.
    """
    on = on or {}
    deleted = deleted or set()
    out: dict[str, str] = {}
    for cid, ask in asks.items():
        beat = beat_of.get(cid, "")
        outcome = outcomes.get(beat, "untouched")
        prefix = _OUTCOME_NOTE.get(outcome, "")
        b, a = before.get(beat, ""), after.get(beat, "")
        if not a or not beat:
            out[cid] = (prefix or "This comment was not addressed in the re-cut. ") + \
                       "Re-comment on the delivered text if it still stands."
            continue
        try:
            verdict = brain.coordinator(
                _VERIFY_PROMPT.format(
                    beat=beat, ask=ask["text"].strip(),
                    on=on.get(cid) or "(the comment covers the whole paragraph)",
                    deleted=_DELETED_NOTE if cid in deleted else "",
                    before=b or "(this beat did not exist)", after=a),
                system=_VERIFY_SYS, num_ctx=8192).strip()
        except Exception as e:  # noqa: BLE001
            log(f"[warn] reply verification failed for comment {cid} ({e})")
            out[cid] = (prefix + "The re-cut addressed this beat; please check the tracked "
                        "changes against your comment — this reply could not be verified.")
            continue
        line = verdict.splitlines()[0].strip() if verdict else ""
        said = line.split(":", 1)[-1].strip() if ":" in line else line
        if line.upper().startswith("SATISFIED"):
            out[cid] = (prefix + "Addressed — " + said).strip()
        else:
            out[cid] = (prefix + "NOT addressed — " + said).strip()
            log(f"[warn] comment {cid} on '{beat}' is NOT satisfied by the re-cut: {said}")
    return out


def _strip_spurious_bold(doc) -> int:
    """Un-bold paragraphs whose every text run is bold. Returns paragraphs cleaned.

    Uniform run-level bold across a whole paragraph is the label era's scar: a
    bold lead-in label once donated its rPr to the full tracked replacement.
    The one-pager's body is plain prose — headings get their weight from their
    style, not their runs — so a wholly-bold paragraph is an artifact, while
    partial bold (real emphasis) is kept. Paragraphs carrying the reviewer's
    live tracked changes are left entirely alone — frozen means frozen, even
    for formatting; the scar heals there once the human accepts their edits."""
    from docx.oxml.ns import qn
    from haarpi import redline as hrl

    def _bolded(r):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            return None
        b = rpr.find(qn("w:b"))
        return b if b is not None and b.get(qn("w:val")) not in ("0", "false", "none") else None

    cleaned = 0
    for p in doc.paragraphs:
        if hrl.has_foreign_markup(p._p):
            continue
        text_runs = [r for r in p._p.iter(qn("w:r"))
                     if any((t.text or "") for t in r.iter(qn("w:t")))
                     or any((t.text or "") for t in r.iter(qn("w:delText")))]
        if not text_runs or not all(_bolded(r) is not None for r in text_runs):
            continue
        for r in text_runs:
            rpr = r.find(qn("w:rPr"))
            for tag in ("w:b", "w:bCs"):
                el = rpr.find(qn(tag))
                if el is not None:
                    rpr.remove(el)
        cleaned += 1
    return cleaned


_COPYEDIT_SYS = """\
You proofread sentences the AUTHOR of a paper wrote by hand. You are looking ONLY for
outright errors: a typo, a misspelling, a broken grammatical construction, a word plainly
used in place of another ("it's" for "its", "has dates" for "dates").

You are NOT an editor. Style, word choice, emphasis, hedging, length, and opinion are the
author's own — disagreeing with them is not an error in them. If a sentence is merely one
you would have written differently, it has no error.

OUTPUT — a single JSON object mapping the label (S1, S2, …) to the corrected sentence,
containing ONLY the sentences that carry a real error. Return {} if there are none. No other
output."""

_COPYEDIT_PROMPT = """\
The author's sentences:
{legend}

Return the JSON object of corrections."""


def _copyedit_notes(brain: Brain, authored: dict[str, dict[str, str]]
                    ) -> list[tuple[str, str]]:
    """Typos in the author's own text, offered as comments — never as edits.

    Their text is frozen against the tool's pen, which would otherwise leave an obvious
    misspelling standing forever with no way to raise it. So the tool raises it: it says
    so, and the author decides. The correction is never applied.

    The pass is keyed on labels of its OWN (S1, S2, …), not on the sentinels. Sentinels are
    numbered per PARAGRAPH — every beat has its own ⟦a:1⟧ — so a legend flattened across
    beats collides, and the last beat silently wins: the model dutifully corrects the typo in
    Motivation's ⟦a:1⟧, the lookup hands back Key result(s)'s ⟦a:1⟧, and the diff between two
    unrelated sentences arrives as a blob anchored on the wrong paragraph. Which is exactly
    what the author got.
    """
    words: dict[str, str] = {}
    lines: list[str] = []
    for beat, spans in authored.items():
        for v in spans.values():
            tag = f"S{len(words) + 1}"
            words[tag] = v.strip()
            lines.append(f'  {tag} ({beat}) = "{v.strip()}"')
    if not words:
        return []
    try:
        raw = brain.coordinator(_COPYEDIT_PROMPT.format(legend="\n".join(lines)),
                                system=_COPYEDIT_SYS, num_ctx=8192)
    except Exception as e:  # noqa: BLE001
        log(f"[warn] copyedit pass failed ({e}); no suggestions offered.")
        return []
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    if not m:
        return []
    try:
        obj = json.loads(m.group(0))
    except json.JSONDecodeError:
        return []
    if not isinstance(obj, dict):
        return []

    notes: list[tuple[str, str]] = []
    for key, fix in obj.items():
        original = words.get(str(key).strip().strip("⟦⟧"), "")
        if not original or not isinstance(fix, str) or not fix.strip():
            continue
        # Word by word, anchored on the offending words, saying only the correction. The
        # whole sentence restated — with a missing "t" or a stray apostrophe buried
        # somewhere inside it — is the tool handing the author their own prose back and
        # making them find the difference.
        got = htext.copyedit_notes(original, fix.strip())
        if not got:
            log(f'[raconteur] copyedit suggestion for {key} rewrote the sentence rather than '
                f'correcting it — dropped. The author\'s prose is his own.')
        notes.extend(got)
    return notes


_CAPTION_STYLES = ("image caption", "caption")


def _caption_paragraphs(doc) -> list:
    """The document's existing figure captions, in order.

    The figures are ALREADY in the .docx as real images — pandoc embedded them at genesis.
    A re-cut must therefore edit the caption that is there, not emit ![caption](path) into
    the prose, which would land in the reader's Word document as literal markdown text.
    """
    return [p for p in doc.paragraphs
            if (p.style.name or "").strip().lower() in _CAPTION_STYLES]


def _document_figures(user_rev: Path) -> list[str]:
    """The captions of the figures ALREADY embedded in the reviewer's document, in order.

    What the re-cut is re-captioning. The writer is shown these so it can tell which figure
    is which — the manifest lists four and the one-pager carries two, and only the caption
    says which two.
    """
    from docx import Document

    from . import redline as rl

    return [t for p in _caption_paragraphs(Document(str(user_rev)))
            if (t := rl._accepted_para_text(p._p).strip())]


def _recaption(doc, beat_text: str, ids, author: str) -> tuple[str, int]:
    """Lift the figure captions out of a beat's markdown and onto the document's figures.

    Returns (prose_without_the_image_lines, captions_rewritten). Order is the only sound
    correspondence available here: the one-pager carries at most two figures and they are
    already embedded, in order. If the writer produced a different number of figures than
    the document holds, that is a structural change a redline cannot make — say so and
    touch nothing.
    """
    from . import redline as rl

    figs = list(guards.FIGURE_MD_RE.finditer(beat_text))
    prose = guards.FIGURE_MD_RE.sub("", beat_text)
    prose = re.sub(r"\n{3,}", "\n\n", prose).strip()
    if not figs:
        return prose, 0

    caps = _caption_paragraphs(doc)
    if len(caps) != len(figs):
        log(f"[warn] the re-cut wrote {len(figs)} figure caption(s) but the document holds "
            f"{len(caps)} figure(s) — captions left alone. Adding or removing a figure is "
            f"not something a redline can do; produce it in rayleigh and re-render.")
        return prose, 0

    n = 0
    for i, (m, cap_para) in enumerate(zip(figs, caps), start=1):
        caption = (m.group("caption") or "").strip()
        if not guards.FIGURE_NUM_RE.match(caption):
            caption = f"Figure {i}: {caption}"
        if rl.tracked_replace(cap_para._p, caption, author, ids):
            n += 1
    return prose, n


def _recut_guard_findings(old: str, new: str, known: set[str],
                          spans: dict[str, str] | None = None) -> list[guards.Finding]:
    """What a re-cut of one beat may not do, decided in code.

    The re-cut path used to run NO guards at all — it inherited the redline machinery's
    rendering half and not its checking half — which is exactly how a beat was rewritten
    whole and silently dropped its only citation.

    ``spans`` are the author's sentences in this beat. Passed on the WRITE path, where the
    text may be the tightened one-pager rather than the per-beat draft the guards already
    cleared: the tightener reads his sentences as ordinary prose and is quite capable of
    working a copy of one into the paragraph around it.
    """
    findings = (guards.author_year_prose(new)
                + guards.dropped_citekeys(old, new)
                + guards.dropped_sentinels(old, new)
                + guards.invented_sentinels(old, new)
                + guards.echoed_spans(new, spans or {}))
    if known:
        findings += guards.unresolved_keys(new, known)
    return findings


def _write_recut(project_dir: Path, cfg: ProjectConfig, paper_dir: Path,
                 brain: Brain, text: str, beats: list[tuple[str, str]],
                 user_rev: Path, authored: dict[str, dict[str, str]] | None = None,
                 skipped: set[str] | None = None,
                 known: set[str] | None = None) -> None:
    """Deliver the re-cut as a redline on the reviewer's own document.

    A re-cut replaces the narrative. It must not replace the CONVERSATION, and it must not
    write a word of the author's prose. Acceptance is a human act: the reviewer's tracked
    changes are never accepted here, and every sentence they typed by hand is held as a
    fixed point (an atom) that the re-cut writes AROUND. The tool's prose beside their
    sentences is fair game — the beat is not abandoned merely because they touched it.

    Every beat replacement is checked before it is written: a re-cut that would drop a
    citation, drop an equation, or lose one of the author's sentences is refused, and the
    beat is left exactly as they left it.

    Major version — new datestamp, chain reset to onepager_ra — but the threads and the
    reviewer's own markup ride along into the new cycle.
    """
    from docx import Document
    from haarpi import redline as hrl

    from . import redline as rl

    authored = authored or {}
    skipped = set(skipped or ())
    known = known or set()

    out = paper_dir / major_onepager_name(cfg.short_title, "docx")
    shutil.copy2(user_rev, out)

    doc = Document(str(out))
    if n := _strip_spurious_bold(doc):
        log(f"[raconteur] un-bolded {n} label-era paragraph(s)")
    ids = rl.ids_for(doc)

    final = _beats_from_md(text)             # beat -> body prose, no label/heading
    fallback = dict(beats)
    replaced, refused, recaptioned, done = 0, 0, 0, set()
    outcomes: dict[str, str] = {}            # beat -> rewritten | refused | untouched
    before: dict[str, str] = {}              # what the reviewer's beat said
    after: dict[str, str] = {}               # what we actually delivered
    for rec in rl.body_paragraphs(doc):
        # accepted text, not .text — a label the reviewer tracked in is
        # invisible to python-docx's .text
        accepted = rl._accepted_para_text(rec["para"]._p).strip()
        beat = _beat_of_para(rec["heading"], accepted)
        if beat is None or beat in done:
            continue
        legacy = _label_led(accepted) is not None   # inline label → migrate it
        spans = authored.get(beat, {})
        p_el = rec["para"]._p
        before[beat] = _strip_label(beat, accepted) if legacy else accepted

        # The label is the TOOL's scaffolding, not the author's prose — so a beat carrying
        # the author's sentences still gets its heading, because we rewrite the beat AROUND
        # them and the label leaves with the text we replaced. A beat we did NOT rewrite
        # keeps its old shape: adding a heading above a paragraph that still opens with
        # "Approach — " gives the reader the word twice.
        def _migrate_heading(el=p_el, name=beat, is_legacy=legacy):
            if is_legacy:
                hrl.tracked_heading_before(el, name, rl.AUTHOR, ids)

        def _stands(el=p_el, name=beat):
            outcomes[name] = outcomes.get(name, "untouched")
            after[name] = rl._accepted_para_text(el).strip()
            done.add(name)

        if beat in skipped:
            log(f"[raconteur] beat '{beat}' left untouched (the re-cut could not keep the "
                f"author's sentence(s) intact)")
            _stands()
            continue

        # A beat the author has a hand in is written from the validated per-beat draft,
        # which still carries their placeholders; the tightened markdown has expanded them
        # back into prose and can no longer be laid into the document safely.
        new = (fallback.get(beat) if spans else (final.get(beat) or fallback.get(beat)))
        if not new:
            _stands()
            continue
        new = new.replace("**", "")

        # The figures are already embedded in this document. Their captions are edited in
        # place; the image markdown never reaches the prose, where Word would render it as
        # the literal text "![Figure 1: …](results/figures/….png)".
        new, n_caps = _recaption(doc, new, ids, rl.AUTHOR)
        if n_caps:
            recaptioned += n_caps

        old = hrl.paragraph_text(p_el, protect_authored=True)
        findings = _recut_guard_findings(old, new, known, spans)
        if findings:
            refused += 1
            outcomes[beat] = "refused"
            log(f"[warn] beat '{beat}' REFUSED — the re-cut would have: "
                + "; ".join(f.imperative for f in findings))
            log("[warn]   → left as it stands. Nothing written for this beat.")
            _stands()
            continue

        if rl.tracked_replace_sentencewise(p_el, new, rl.AUTHOR, ids,
                                           protect_authored=True):
            replaced += 1
            outcomes[beat] = "rewritten"
            # The old inline label left with the text we replaced; the beat name comes
            # back as a real section heading.
            _migrate_heading()
            log(f"[raconteur] beat '{beat}' re-cut as tracked changes"
                + (f", around {len(spans)} authored span(s)" if spans else ""))
        _stands()

    doc.save(str(out))
    log(f"[raconteur] wrote {out.relative_to(project_dir)} "
        f"(re-cut: {replaced} beat(s) rewritten, {refused} refused"
        + (f", {recaptioned} figure caption(s) rewritten" if recaptioned else "")
        + f"; the author's sentences held fixed throughout)")

    # Every reply is checked against what was actually delivered, beat by beat.
    log("[raconteur] verifying the re-cut against each open ask…")
    asks = {a["id"]: a for a in hrl.open_asks(user_rev)}
    beat_of: dict[str, str] = {}
    for anchor in rl.comment_anchors(user_rev, only=set(asks)):
        b = _beat_of_para(anchor.get("heading") or "", anchor["text"])
        for cid in anchor["ids"]:
            if b:
                beat_of[cid] = b
    # What each ask POINTS AT, and whether the re-cut struck it out from under them. An ask
    # whose anchor now lies wholly inside a w:del was asked about text that no longer exists,
    # and the adversary has to be told — deleting the sentence someone asked you to explain
    # is not an explanation.
    on = anchor_words(user_rev, only=set(asks))
    struck: set[str] = set()
    for rec in rl.body_paragraphs(doc):
        struck |= hrl.anchors_in_deleted_text(rec["para"]._p)
    replies = _verify_replies(brain, asks, beat_of, before, after, outcomes,
                              on=on, deleted=struck & set(asks))
    n = hrl.add_replies(out, replies, author=rl.AUTHOR)
    log(f"[raconteur] {n} threaded repl{'y' if n == 1 else 'ies'} written")

    if authored:
        notes = _copyedit_notes(brain, authored)
        if notes:
            n = hrl.add_anchored_comments(out, notes, author=rl.AUTHOR)
            log(f"[raconteur] {n} suggested copyedit(s) on the author's own text, "
                f"delivered as comments — not applied")

    md_path = paper_dir / major_onepager_name(cfg.short_title, "md")
    md_path.write_text(rl.accepted_markdown(Document(str(out))).strip() + "\n",
                       encoding="utf-8")
    log(f"[raconteur] wrote {md_path.relative_to(project_dir)} "
        f"(accepted view of the delivered redline)")


def _beat_of(text: str) -> str | None:
    """The beat a heading or a legacy label-led paragraph belongs to."""
    t = text.lstrip("#*-–—• ").lower()
    for beat in _BEATS:
        if t.startswith(beat.name.lower()):
            return beat.name
    return None


def _label_led(text: str) -> str | None:
    """The beat of a paragraph carrying an EXPLICIT inline label ('Gap — …').

    Stricter than _beat_of: the beat name must be followed by a separator, so
    body prose that merely opens with a beat word is never mistaken for a
    label. An explicit label outranks heading scoping — a frozen legacy
    paragraph sitting under some other beat's migrated heading still belongs
    to ITS beat."""
    t = text.lstrip("#*-–—• ").strip()
    for beat in _BEATS:
        if t.lower().startswith(beat.name.lower()):
            rest = t[len(beat.name):].lstrip("*").lstrip()
            if rest[:1] in _LABEL_DASHES:
                return beat.name
    return None


def _beat_of_para(heading: str, text: str) -> str | None:
    """The beat a document paragraph belongs to: inline label first, then heading."""
    return _label_led(text) or _beat_of(heading or "")


def _beats_from_md(md: str) -> dict[str, str]:
    """Beat-name -> body prose (heading/label stripped), from beat-structured markdown.

    Understands both shapes: '## Gap' section headings with the body in the
    blocks that follow (current), and legacy '**Gap** — …' label-led paragraphs.
    Mixed documents are real: a re-cut migrates replaced beats to headings while
    the author's frozen beats keep their legacy label paragraphs. So heading-mode
    still adopts a label-led paragraph that sits OUTSIDE any beat section, for a
    beat no heading has claimed — but body prose inside a headed section is never
    mistaken for a label."""
    blocks = [b.strip() for b in md.split("\n\n") if b.strip()]
    out: dict[str, str] = {}
    if any(b.startswith("#") and _beat_of(b) for b in blocks):
        current: str | None = None
        for b in blocks:
            if b.startswith("#"):
                beat = _beat_of(b)
                current = beat if beat and beat not in out else None
                continue
            lab = _label_led(b)          # a legacy label paragraph among headed
            if lab and lab not in out:   # beats — a frozen beat, kept verbatim;
                out[lab] = _strip_label(lab, b)   # its label outranks the
                current = None                     # enclosing heading's scope
                continue
            if current:
                out[current] = f"{out[current]}\n\n{b}" if current in out else b
        return out
    for b in blocks:
        beat = _beat_of(b)
        if beat and beat not in out:
            out[beat] = _strip_label(beat, b)
    return out


def _prior_serialized(user_rev: Path) -> dict[str, str]:
    """beat -> the previous paragraph AS THE GUARDS MUST SEE IT: authored spans as ⟦a:N⟧.

    There are two views of the same beat, and handing the wrong one to a guard is fatal in a
    way that looks like nothing at all.

    The ACCEPTED PROSE (``_annotation_brief``'s ``prior``) is what a reader sees: the
    author's sentence spelled out in words. Compared against THAT, every ⟦a:1⟧ the draft was
    ORDERED to carry is a placeholder the model invented — and `invented_sentinels` rejects
    the beat for obeying its instructions. It happened: three of the five beats in the
    2026-07-14 re-cut were killed this way, twice each, and reported only as "NOT re-cut",
    while the write path (which compares against the serialized text, correctly) went on to
    report "0 refused". The feature silently disabled itself on precisely the beats it exists
    for — the ones the author had touched.

    So the guards get the SERIALIZED view, where ⟦a:1⟧ genuinely exists.
    """
    from docx import Document
    from haarpi import redline as hrl

    from . import redline as rl

    out: dict[str, str] = {}
    for rec in rl.body_paragraphs(Document(str(user_rev))):
        accepted = rl._accepted_para_text(rec["para"]._p).strip()
        beat = _beat_of_para(rec["heading"], accepted)
        if beat is None or beat in out:
            continue
        out[beat] = hrl.paragraph_text(rec["para"]._p, protect_authored=True)
    return out


def _beat_integrity_problems(draft: str, spans: dict[str, str], was: str,
                             known: set[str], signature: dict | None = None) -> list[str]:
    """The problems that make a beat unsafe to WRITE — a dropped citation, a dropped or
    invented authored span, the author's own sentence retyped, an off-voice draft.

    These are FATAL: a beat that trips one is left exactly as it stands, because a hole
    where the author's words were — or a beat with its only citation quietly gone — is worse
    than a stale beat. ``was`` must be the SERIALIZED previous paragraph (see
    ``_prior_serialized``), never the accepted prose.
    """
    errs = _sentinel_errors(draft, spans)
    errs += [f.imperative for f in guards.echoed_spans(draft, spans)]
    errs += [f.imperative for f in guards.style_findings(draft, signature or {})]
    if was:
        errs += [f.imperative for f in _recut_guard_findings(was, draft, known)]
    return errs


def _beat_figure_problems(draft: str, expect_figures: int | None = None) -> list[str]:
    """A figure the prose never numbers, introduces, or captions. SOFT: a quality shortfall,
    not an integrity breach, so it drives a retry but never abandons the beat — the prose
    that answers the reviewer's comments must not be held hostage to a caption the model
    could not get right. The figures are re-captioned in place at write time regardless.
    ``expect_figures`` is how many figures the document already holds; see
    ``guards.figure_findings``.
    """
    return [f.imperative for f in guards.figure_findings(draft, expect=expect_figures)]


def _beat_problems(draft: str, spans: dict[str, str], was: str,
                   known: set[str], signature: dict | None = None,
                   expect_figures: int | None = None) -> list[str]:
    """Integrity and figure problems together. The drafting loop keeps them apart — only
    integrity can abandon a beat — but callers that want the whole list use this."""
    return (_beat_integrity_problems(draft, spans, was, known, signature)
            + _beat_figure_problems(draft, expect_figures))


# The reviewer's asks are verified against the DELIVERED text at the end of the run, and the
# reply says, honestly, which ones the re-cut missed. But an honest "NOT addressed" is a
# consolation, not a fix. The same judgement, run against the DRAFT while it can still be
# redone, turns a reported failure into a retried one: a "define this" the model ignored is
# put back to it with the reason, up to a bounded number of times.
_MAX_BEAT_TRIES = 3

_INLOOP_VERIFY_SYS = (
    "You judge whether a DRAFT paragraph answers a reviewer's asks, before it is delivered. "
    "Be strict and literal. An ask to DEFINE a term is answered only if the draft says, in "
    "the text, what the term MEANS — not by using it again, and not by deleting it. If the "
    "draft does the thing asked, say OK; otherwise say what it still needs, in one clause.")

_INLOOP_VERIFY_PROMPT = """A reviewer left these asks on the "{beat}" beat. For each, decide whether the DRAFT below answers it.

Asks:
{asks}

The beat as it read BEFORE:
{before}

The DRAFT under review:
{draft}

Output exactly one line per ask, numbered, and nothing else:
  <n>: OK
  <n>: MISSING — <what the draft still needs to do, one clause>"""

_MISSING_RE = re.compile(r"(\d+)\s*[:.]?\s*MISSING\b[\s—:-]*(.*)", re.IGNORECASE)


def _unaddressed_asks(brain: Brain, beat: str, before: str, draft: str,
                      beat_asks: list[dict]) -> list[str]:
    """Which of a beat's open asks the DRAFT still leaves unanswered — the verify verdict fed
    back into the drafting loop, so an ask the model ignored is retried, not merely reported.

    Returns one feedback string per unaddressed ask; ``[]`` when the draft answers them all,
    when the beat has no open asks, or when the judge cannot be reached — a verify hiccup
    must never block a structurally-sound beat from being delivered.
    """
    if not beat_asks:
        return []
    lines = []
    for i, a in enumerate(beat_asks, start=1):
        ref = (a.get("on") or "").strip()
        lines.append(f'{i}. "{a["text"].strip()}"'
                     + (f' (what "this"/"it" refers to: {ref})' if ref else ""))
    try:
        verdict = brain.coordinator(
            _INLOOP_VERIFY_PROMPT.format(
                beat=beat, asks="\n".join(lines),
                before=before or "(this beat did not exist)", draft=draft),
            system=_INLOOP_VERIFY_SYS, num_ctx=8192).strip()
    except Exception as e:  # noqa: BLE001
        log(f"[warn] in-loop verify failed for beat '{beat}' ({e}) — accepting the draft")
        return []
    missing: list[str] = []
    for line in verdict.splitlines():
        m = _MISSING_RE.match(line.strip())
        if not m:
            continue
        idx = int(m.group(1)) - 1
        if 0 <= idx < len(beat_asks):
            why = m.group(2).strip() or "not addressed"
            missing.append(f'"{beat_asks[idx]["text"].strip()}" — {why}')
    return missing


def _authored_by_beat(user_rev: Path) -> dict[str, dict[str, str]]:
    """beat -> {⟦a:N⟧: the author's exact words} for the sentences they typed into it.

    The unit of deference is the SPAN, not the beat. A beat carrying the author's hand is
    not abandoned — it is re-cut AROUND their sentences, which stand as fixed points. The
    tool's own prose beside them stays fair game, which is the whole difference between
    deferring to the author and freezing whatever paragraph they happened to touch.

    Sentinels are numbered per paragraph, and each beat is one paragraph, so a beat's
    legend is self-consistent.
    """
    from docx import Document
    from haarpi import redline as hrl

    from . import redline as rl

    out: dict[str, dict[str, str]] = {}
    for rec in rl.body_paragraphs(Document(str(user_rev))):
        # accepted text, not .text — a label the reviewer tracked in is
        # invisible to python-docx's .text
        accepted = rl._accepted_para_text(rec["para"]._p).strip()
        beat = _beat_of_para(rec["heading"], accepted)
        if beat is None or beat in out:
            continue
        atoms = hrl.authored_atoms(rec["para"]._p)
        if atoms:
            out[beat] = atoms
    return out


def _expand_spans(text: str, authored: dict[str, str]) -> str:
    """A beat as a reader sees it: the author's placeholders back in their own words."""
    for key, words in authored.items():
        text = text.replace(key, words.strip())
    return text


def _sentinel_errors(text: str, authored: dict[str, str]) -> list[str]:
    """Why a generated beat may not be written: it lost, duplicated, or invented a span.

    The author's sentences are the one thing a re-cut may not get wrong, so this fails
    closed rather than delivering a beat with a hole where their words were.
    """
    errors = []
    for key in authored:
        n = text.count(key)
        if n == 0:
            errors.append(f"{key} is missing — the author's sentence must appear, in place.")
        elif n > 1:
            errors.append(f"{key} appears {n} times — it must appear exactly once.")
    for found in set(guards.SENTINEL_RE.findall(text)):
        if found not in authored:
            errors.append(f"{found} is not a real placeholder — do not invent one.")
    return errors


def _adopt_title(project_dir: Path, cfg: ProjectConfig, user_rev: Path) -> None:
    """If the reviewer retitled the document, the config follows.

    The title is config-owned (every _write emits `# {cfg.title}`), so a tracked
    change to the heading can only persist by updating cfg.title itself.
    """
    from docx import Document
    from . import redline
    md = redline.accepted_markdown(Document(str(user_rev)))
    for line in md.splitlines():
        if line.startswith("#"):
            if _beat_of(line):
                continue   # a beat heading is never the title
            title = re.sub(r"\s*[—–-]+\s*one-pager\s*$", "", line.lstrip("# ").strip())
            if title and title != cfg.title:
                log(f"[raconteur] adopting reviewer's title: {title!r}")
                cfg.title = title
                cfg.save(project_dir)
            return


_ANCHOR_CAP = 180
_WHOLE_PARA = 0.9      # an anchor this much of the paragraph is not pointing at anything


def anchor_words(path: Path, only: set[str] | None = None) -> dict[str, str]:
    """comment id -> the exact words the reviewer's comment is anchored to.

    Half of a comment is WHERE it points. "define this" is not an instruction; "define this",
    pinned to *anchored distance*, is. Rendered without its anchor the reviewer's three-word
    ask reached the writer as a riddle — and reached the adversary that checks the answer as
    one too, which is how it came to complain that the text "fails to define the term Key
    result(s)". It was reading the heading, for want of anything else to read.

    A comment covering the WHOLE paragraph is omitted: it is not pointing at anything, and
    quoting the paragraph back into a brief that already carries the paragraph is the same
    noise, from the other direction.
    """
    from docx import Document
    from haarpi import redline as hrl

    from . import redline as rl

    out: dict[str, str] = {}
    for rec in rl.body_paragraphs(Document(str(path))):
        p_el = rec["para"]._p
        # The same serialization comment_spans measures its offsets against, or the words
        # come back shifted.
        text = hrl.paragraph_text(p_el, protect_authored=True)
        for cid, (start, end) in hrl.comment_spans(p_el, protect_authored=True).items():
            if only is not None and cid not in only:
                continue
            frag = text[start:end].strip()
            if not frag or len(frag) >= _WHOLE_PARA * len(text.strip()):
                continue
            out[cid] = frag if len(frag) <= _ANCHOR_CAP else frag[:_ANCHOR_CAP] + "…"
    return out


def _ask_block(ask: dict, on: str = "") -> str:
    """One reviewer ask, rendered for the brief — with its anchor, its thread, its history.

    A thread is part of the ask. New information about an unmet ask arrives as a REPLY
    (that is the protocol), so a brief that carries only the top comment carries only
    half of what the reviewer said. And an ask the tool has already answered, still
    open, is not a fresh task: it is a task the tool got wrong. Saying so is what stops
    the next pass from confidently repeating the last one's mistake.
    """
    lines = [f'- on "{on}" — {ask["text"].strip()}' if on else f"- {ask['text'].strip()}"]
    for f in ask.get("followups", []):
        lines.append(f"  (reviewer, in the same thread) {f.strip()}")
    if ask.get("repeat"):
        lines.append("  ** YOU ALREADY ANSWERED THIS ONCE AND THE REVIEWER LEFT IT OPEN. "
                     "Your previous answer did not satisfy them — do not repeat it. **")
        for r in ask.get("prior_tool_replies", []):
            lines.append(f"  (your previous, rejected answer) {r.strip()[:300]}")
    return "\n".join(lines)


def _annotation_brief(user_rev: Path) -> tuple[dict[str, str], dict[str, str], str]:
    """The re-cut briefing from the reviewer's docx.

    Returns (prior, notes, general): each beat's previous text with the reviewer's
    tracked changes accepted, the OPEN asks anchored to it, and the asks that anchor to
    no recognisable beat (title, figures) — those apply to the whole narrative.

    Only open asks. A resolved comment is history, not an instruction: re-firing it
    makes the tool rewrite prose the reviewer has already accepted, and a settled
    "this sentence is ok" is quite capable of provoking a full re-cut of the paragraph
    it blessed.
    """
    from docx import Document
    from haarpi import redline as hredline

    from . import redline

    prior = _beats_from_md(redline.accepted_markdown(Document(str(user_rev))))

    asks = {a["id"]: a for a in hredline.open_asks(user_rev)}
    live = set(asks)
    on = anchor_words(user_rev, only=live)

    notes: dict[str, str] = {}
    general: list[str] = []

    def _add(beat: str | None, block: str) -> None:
        if not block:
            return
        if beat is None:
            general.append(block)
        else:
            notes[beat] = f"{notes[beat]}\n{block}" if beat in notes else block

    for anchor in redline.comment_anchors(user_rev, only=live):
        block = "\n".join(_ask_block(asks[c], on.get(c, ""))
                          for c in anchor["ids"] if c in asks)
        _add(_beat_of_para(anchor.get("heading") or "", anchor["text"]), block)
    for h in redline.heading_comments(user_rev, only=live):
        block = "\n".join(_ask_block(asks[c]) for c in h["ids"] if c in asks)
        _add(_beat_of(h["heading"]), block)
    return prior, notes, "\n".join(general)


# ── fresh one-pager ───────────────────────────────────────────────────────────

def _onepager_fresh(
    project_dir: Path, cfg: ProjectConfig, brain: Brain, paper_dir: Path,
    brief: tuple[dict[str, str], dict[str, str], str] | None = None,
    user_rev: Path | None = None,
) -> None:
    from .outline import _analyze_structure, _build_venue_section, analysis_view

    litrev = load_litreview(project_dir, cfg.litrev_dir) if cfg.litrev_dir else ""
    code = load_methods(project_dir) if cfg.use_methods else ""
    results = load_results(project_dir, cfg.results_dir) if cfg.results_dir else ""
    figures = load_figure_manifest(project_dir, cfg.results_dir or "results")
    signature = load_style_signature(project_dir)
    known = load_bib_keys(project_dir, cfg.litrev_dir) if cfg.litrev_dir else set()

    log("[raconteur] analysing paper structure…")
    analysis = _analyze_structure(brain, cfg.description, litrev, code, results)

    venue_section = _build_venue_section(cfg, project_dir)
    figure_list = _figure_section(figures)

    # On a re-cut the figures are already embedded in the reviewer's .docx, so their number
    # is settled and the writer's job is to caption and introduce the ones that are there —
    # not to decide whether to have any. Without a path list it cannot write a caption line
    # at all, so an absent manifest falls back to leaving the captions alone.
    doc_figures = _document_figures(user_rev) if user_rev is not None else []
    fig_expect = len(doc_figures) if (doc_figures and figure_list) else None

    prior, beat_notes, general = brief if brief else ({}, {}, "")
    general_block = _RECUT_GENERAL.format(notes=general) if general else ""

    # The author's own sentences are fixed points the beat is rebuilt AROUND — not a
    # reason to abandon the beat. Deference is owed to the text they wrote, not to the
    # paragraph it happens to sit in.
    authored = _authored_by_beat(user_rev) if user_rev is not None else {}
    # The guards compare a draft against the PREVIOUS paragraph, and must see it the way the
    # draft was told to write it: with the author's spans as placeholders.
    prior_ser = _prior_serialized(user_rev) if user_rev is not None else {}
    if authored:
        log("[raconteur] the author's hand is in: "
            + ", ".join(f"{b} ({len(a)} span(s))" for b, a in sorted(authored.items()))
            + " — held fixed, written around")

    # The open asks on each beat, with what each one points at — so a draft can be checked
    # against them AS IT IS WRITTEN and retried when it leaves one unanswered, instead of only
    # being told, after delivery, that it failed. (Re-cut only; genesis has no asks.)
    asks_by_beat: dict[str, list[dict]] = {}
    if user_rev is not None and brief:
        from haarpi import redline as hrl

        from . import redline as rl

        live = {a["id"]: a for a in hrl.open_asks(user_rev)}
        on = anchor_words(user_rev, only=set(live))
        for anchor in rl.comment_anchors(user_rev, only=set(live)):
            b = _beat_of_para(anchor.get("heading") or "", anchor["text"])
            if b is None:
                continue
            for cid in anchor["ids"]:
                if cid in live:
                    asks_by_beat.setdefault(b, []).append(
                        {"text": live[cid]["text"], "on": on.get(cid, "")})

    beats: list[tuple[str, str]] = []
    skipped: set[str] = set()          # beats the re-cut could not safely touch
    for beat in _BEATS:
        can_embed = bool(figure_list) and beat.name == _FIGURE_BEAT
        # The figure beat of a re-cut is re-captioning a document that already holds N.
        expect = fig_expect if beat.name == _FIGURE_BEAT else None
        if expect:
            figure_rule = _RECUT_FIGURE_RULE.format(
                n=expect,
                current="\n".join(f'  {i}. "{c}"'
                                  for i, c in enumerate(doc_figures, start=1)))
        else:
            figure_rule = _FIGURE_RULE if can_embed else _NO_FIGURE_RULE
        spans = authored.get(beat.name, {})
        # a beat is a kind of prose, and his Results voice is not his Discussion voice
        style_section = _style_block(
            load_style_profile(project_dir, kind=_BEAT_KIND.get(beat.name, "")))
        recut_section = ""
        if brief:
            # Show the previous beat as the model must WRITE it: with the author's spans as
            # placeholders, in the positions they occupied. Spelled out, the model reads
            # their sentence as ordinary prose it may reword — and it is looking straight at
            # the words it has just been forbidden to touch.
            was = (prior_ser.get(beat.name) if spans else None) \
                or prior.get(beat.name, "(this beat did not exist)")
            recut_section = _RECUT_BLOCK.format(
                beat=beat.name,
                prior=was,
                notes=beat_notes.get(beat.name, "(none on this beat)"),
            ) + general_block
        legend = "\n".join(f'  {k} = "{v.strip()}"' for k, v in spans.items())
        prompt = _BEAT_PROMPT.format(
            beat=beat.name,
            intent=beat.intent,
            title=cfg.title,
            topic=cfg.topic,
            focus=cfg.focus,
            venue_section=venue_section,
            style_section=style_section,
            analysis=analysis_view(analysis, beat.drop),
            evidence_section=_evidence_for_beat(
                beat.sources, cfg.description, litrev, code, results
            ),
            recut_section=recut_section,
            authored_section=_AUTHORED_BLOCK.format(legend=legend) if spans else "",
            preceding_section=_preceding_block(beats, set(authored)),
            figure_section=figure_list if can_embed else "",
            figure_rule=figure_rule,
            extra_rules=f"\n{beat.rules}" if beat.rules else "",
        )
        log(f"[raconteur] drafting beat '{beat.name}'…"
            + (f" ({len(spans)} authored span(s) held fixed)" if spans else ""))

        beat_asks = asks_by_beat.get(beat.name, [])
        before_prose = prior.get(beat.name, "")

        # The "previous text" the integrity guards compare against is the SERIALIZED paragraph
        # — the author's spans as ⟦a:N⟧, not spelled out. Against the accepted prose, every
        # placeholder the draft was ordered to carry reads as one it invented.
        def _integrity(draft: str, was: str = prior_ser.get(beat.name, "")) -> list[str]:
            return _beat_integrity_problems(draft, spans, was, known, signature)

        # Retry on anything wrong — a broken guard, a figure it did not caption, or an ask it
        # left unanswered — but only an INTEGRITY breach can abandon the beat. A figure the
        # model could not caption, or a definition it would not write, is delivered as prose
        # and reported honestly; it never costs the reviewer the beat's other answers.
        text = _strip_label(beat.name, brain.coordinator(prompt, system=_SYSTEM,
                                                         num_ctx=16384))
        integrity: list[str] = []
        figs: list[str] = []
        unans: list[str] = []
        for attempt in range(1, _MAX_BEAT_TRIES + 1):
            integrity = _integrity(text)
            figs = _beat_figure_problems(text, expect)
            unans = ([] if integrity
                     else _unaddressed_asks(brain, beat.name, before_prose,
                                            _expand_spans(text, spans), beat_asks))
            problems = integrity + figs + unans
            if not problems or attempt == _MAX_BEAT_TRIES:
                break
            log(f"[warn]   → beat '{beat.name}' attempt {attempt}: "
                + "; ".join(problems[:3]) + " — retrying")
            retry = prompt + ("\n\nYour previous attempt still has problems:\n"
                              + "\n".join(f"- {p}" for p in problems)
                              + "\nReturn the beat again, fixing every one. Keep every "
                                "citation and every sentence the author wrote.")
            text = _strip_label(beat.name, brain.coordinator(retry, system=_SYSTEM,
                                                             num_ctx=16384))

        if integrity:
            # Fail closed. A hole where the author's words were, or a dropped citation, is
            # worse than a beat left alone.
            log(f"[warn]   → beat '{beat.name}' NOT re-cut: {'; '.join(integrity[:2])}. "
                f"Left exactly as it stands.")
            skipped.add(beat.name)
            text = prior.get(beat.name, "").strip() or text
        else:
            if figs:
                log(f"[raconteur]   → beat '{beat.name}' delivered; figures fell short "
                    f"({figs[0]}) — prose kept, captions handled at write time")
            if unans:
                log(f"[raconteur]   → beat '{beat.name}' delivered; {len(unans)} ask(s) may "
                    f"still be unmet — the reply will say which")
        beats.append((beat.name, text))

    # The critique reads the one-pager as a reader would — the author's sentences in full,
    # not as placeholders. Only the document write needs the sentinels.
    draft = _assemble([(n, _expand_spans(t, authored.get(n, {}))) for n, t in beats])

    log("[raconteur] critiquing one-pager…")
    critique = brain.coordinator(
        _CRITIQUE_PROMPT.format(onepager=draft),
        system=_SYSTEM,
        num_ctx=8192,
    )
    log(f"[raconteur] critique findings:\n{critique.strip()}")

    log("[raconteur] tightening one-pager…")
    tighten_prompt = _REVISE_PROMPT.format(onepager=draft, critique=critique)
    if authored:
        # The tightener reads the author's sentences as prose, so it must be told which
        # prose is theirs. It may reword the beat around them; it may not touch them.
        tighten_prompt += (
            "\n\nThese sentences were written by the AUTHOR, by hand. Reproduce each one "
            "VERBATIM, whatever the critique says — you may rewrite the prose around them, "
            "never the sentences themselves:\n"
            + "\n".join(f'- "{v.strip()}"'
                        for spans in authored.values() for v in spans.values()))
    tightened = brain.coordinator(tighten_prompt, system=_SYSTEM, num_ctx=8192)

    # The tightener can quietly undo what the beats got right — dropping the sentence that
    # introduced a figure, or a caption's number. The per-beat drafts already passed those
    # checks, so a tightened version that fails them is not an improvement.
    broke = guards.figure_findings(tightened, expect=fig_expect)
    if broke and not guards.figure_findings(draft, expect=fig_expect):
        log("[warn] the tightening pass broke a figure ("
            + "; ".join(f.kind for f in broke) + ") — keeping the drafted beats instead")
        tightened = draft

    if user_rev is not None:
        _write_recut(project_dir, cfg, paper_dir, brain, tightened, beats, user_rev,
                     authored=authored, skipped=skipped, known=known)
    else:
        _write(project_dir, cfg, paper_dir, tightened)


# ── user-annotation revision ──────────────────────────────────────────────────

def _revise(
    project_dir: Path,
    cfg: ProjectConfig,
    brain: Brain,
    paper_dir: Path,
    user_rev: Path,
) -> None:
    """Answer each anchored comment with an in-place tracked change (paper parity).

    Edits a copy of the reviewer's .docx — comments stay anchored and get
    dispositions, the reviewer's own tracked changes survive, and every
    un-flagged paragraph is byte-for-byte untouched. The accepted-text .md
    sibling is what 'outline' binds.
    """
    from .paper import _bib_block
    from .redline_revise import redline_revise

    litrev = load_litreview(project_dir, cfg.litrev_dir) if cfg.litrev_dir else ""
    code = load_methods(project_dir) if cfg.use_methods else ""
    results = load_results(project_dir, cfg.results_dir) if cfg.results_dir else ""
    bib_summary = load_bib_summary(project_dir, cfg.litrev_dir) if cfg.litrev_dir else ""
    bib_keys = load_bib_keys(project_dir, cfg.litrev_dir) if cfg.litrev_dir else set()

    def beat_context(heading: str, para_text: str) -> str:
        # Beats live under their own section headings; legacy one-pagers carry
        # the beat name as an inline lead-in label instead. Either way, route
        # each beat its own evidence bundle.
        name = _beat_of_para(heading, para_text)
        for beat in _BEATS:
            if beat.name == name:
                return _evidence_for_beat(
                    beat.sources, cfg.description, litrev, code, results
                )
        return ""

    out, _ = redline_revise(project_dir, cfg, brain, paper_dir, user_rev,
                            litrev, code, results, _bib_block(bib_summary), bib_keys,
                            context_fn=beat_context, md_sibling=True)

    # Label-era documents carry wholly-bold beat paragraphs, and any tracked
    # edit written into one inherited that bold — normalise the delivered copy.
    from docx import Document
    doc = Document(str(out))
    if n := _strip_spurious_bold(doc):
        doc.save(str(out))
        log(f"[raconteur] un-bolded {n} label-era paragraph(s)")


# ── entry point ───────────────────────────────────────────────────────────────

def run(project_dir: Path, resynth: bool = False) -> None:
    if not ProjectConfig.exists(project_dir):
        log("[error] no paper/raconteur.yaml found — run 'raconteur init' first")
        raise SystemExit(1)

    cfg = ProjectConfig.load(project_dir)
    gcfg = GlobalConfig.load()
    paper_dir = project_dir / "paper"
    paper_dir.mkdir(exist_ok=True)

    if not cfg.description:
        log("[error] no research description — run 'raconteur init' first")
        raise SystemExit(1)

    check_prerequisites(project_dir, cfg)
    _ensure_style(project_dir, cfg)

    brain = Brain(gcfg, coordinator=cfg.brain.coordinator_model)

    if not cfg.topic or not cfg.focus:
        from .outline import _parse_description
        log("[raconteur] extracting topic and focus…")
        parsed = _parse_description(brain, cfg.description)
        if parsed.get("topic"):
            cfg.topic = parsed["topic"]
        if parsed.get("focus"):
            cfg.focus = parsed["focus"]
        if not cfg.title and parsed.get("title"):
            cfg.title = parsed["title"]
        cfg.save(project_dir)
        log(f"  title : {cfg.title}")
        log(f"  topic : {cfg.topic}")
        log(f"  focus : {cfg.focus}")

    user_rev = find_user_revision(paper_dir, cfg.short_title, chain_includes="onepager")
    existing = find_latest(paper_dir, cfg.short_title, "md",
                           last_initials="ra", chain_includes="onepager")

    if not existing:
        _onepager_fresh(project_dir, cfg, brain, paper_dir)
    elif user_rev:
        log(f"[raconteur] found revision: {user_rev.name}")
        if resynth:
            # Narrative-level rejection: the annotations are a brief for a fresh
            # cut, not line edits. Major version — new datestamp, chain reset —
            # but delivered as a redline on the reviewer's file: their edits
            # persist, their comment threads ride along and get replies.
            log("[raconteur] --resynth: re-cutting the narrative from the annotations")
            _adopt_title(project_dir, cfg, user_rev)
            _onepager_fresh(project_dir, cfg, brain, paper_dir,
                            brief=_annotation_brief(user_rev), user_rev=user_rev)
        else:
            _revise(project_dir, cfg, brain, paper_dir, user_rev)
    else:
        log("[raconteur] one-pager already exists — annotate the docx with your "
            "initials and re-run to revise, or run 'raconteur outline'")
        return

