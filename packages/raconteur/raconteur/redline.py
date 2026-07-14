"""In-place, comment-preserving revision with tracked changes.

This is what `paper` does by default when it finds an annotated .docx. The alternative
(`paper --resynth`) re-synthesises the whole manuscript from markdown and renders a fresh
.docx — which discards the reviewer's Word comments and gives them no redline to read the
tool's edits against. This module instead edits a COPY of the annotated .docx in place: it
answers each comment by rewriting only the paragraph(s) that comment is anchored to, records
every rewrite as a Word tracked change attributed to `raconteur`, and leaves the comments
anchored and every un-flagged paragraph byte-for-byte untouched.

Two failures fall out of the clean-rewrite default, and this module exists to kill both:

  COLLATERAL DRIFT — a comment on the Discussion causes the Methods section to be rewritten.
    The reviewer sees changes they never asked for, in sections they approved. A tracked
    change that altered an untouched region, under a reply claiming success, is worse than
    no edit at all.

  NO REDLINE — the reviewer annotated a specific sentence; a clean rewrite throws the
    sentence away and writes a new paragraph. There is no way to see what changed, accept
    some edits and reject others, or trust that the parts they liked survived.

This file is the deterministic machinery only: XML surgery, GPU-free and unit-testable. The
LLM call that turns a comment into revised paragraph text lives in `revise`.

OOXML notes:
  * A comment is anchored by ``<w:commentRangeStart w:id=N/>`` … ``<w:commentRangeEnd
    w:id=N/>`` markers bracketing a run range, plus a ``<w:commentReference w:id=N/>`` run;
    the text lives in comments.xml. python-docx preserves all of these across an open/save,
    so we only manipulate the body XML.
  * A tracked deletion wraps the old run(s) in ``<w:del>`` and turns ``<w:t>`` into
    ``<w:delText>``; a tracked insertion wraps new run(s) in ``<w:ins>``. Both carry an
    author and date, and Word renders them as an accept/rejectable redline.

A paragraph is modelled as an ordered stream of TEXT and OPAQUE atoms (equations, footnote
references, drawings), NOT as the text inside its ``w:r/w:t`` runs. That older model is blind
to everything else in the paragraph: an equation is a SIBLING of the text runs, so a differ
built on ``w:t`` alone sees prose with holes where every number had been, no sentence can
match, and each rewrite collapses to a whole-paragraph replacement — with the equations left
stranded at the paragraph tail, severed from the claims they verified. raconteur hits this
immediately: a Results section is full of inline statistics rendered as OMML.

Atoms serialize to sentinels (``⟦m:1⟧``) for the differ and for the LLM, and expand back to
their original elements on write.

    INVARIANT: raconteur never authors an equation; it only edits the prose around one.

An atom is re-laid as ACCEPTED content between the redlined prose, never inside a
``w:ins``/``w:del``. ``guards.dropped_sentinels`` and ``guards.invented_sentinels`` fail the
edit closed if the model breaks the invariant.

Known limitations (documented, not bugs):
  * Multiple comments on one paragraph are coarsened to bracket the whole revised paragraph —
    every comment stays valid and anchored, but loses sub-paragraph precision. ``comment_spans``
    recovers that precision on the way IN, which is what tells the minimal-edit guard which
    sentences a comment actually bears on.
  * Assumes the annotated draft has no still-open tracked changes from a prior cycle (true for
    a freshly rendered _ra draft the reviewer annotated).
"""

from __future__ import annotations

import copy
import datetime
import difflib
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from . import guards

# ── the shared engine (haarpi.redline), re-exported under the old names ──────
from haarpi import redline as _engine
from haarpi.redline import (  # noqa: F401
    _MATH, _XML_SPACE, _SENTINEL_SPLIT, _OPAQUE_RUN_CHILDREN,
    _now, _Ids, _max_existing_id, ids_for, _rpr_clone, _text_run, _ins, _del,
    _is_text_run, _is_ref_run, _is_opaque, _sentinel_kind,
    serialize_paragraph, paragraph_text, atom_text, flatten_paragraph,
    _render, _segments, _redline_chunk, _relay,
    comments_by_id, comment_spans, anchored_sentences, is_heading_style,
    _accepted_para_text,
)


AUTHOR = "raconteur"


def is_title_style(style_name: str) -> bool:
    """The document title is not a section heading. It must be skipped like a heading, but it
    must not become the enclosing section of the abstract that follows it."""
    return (style_name or "").lower() == "title"


# ── raconteur document structure ──────────────────────────────────────────────
# rabbitHole's narrative is a flat run of body paragraphs. raconteur's .docx has a title, an
# abstract, ## section headings, ### subsection headings, and a References list. The redline
# must touch body prose only, and must know which section a paragraph belongs to so the
# reviser gets the right context bundle (a Methods sentence needs the methods writeup; a
# Background sentence needs the bib).


def _style_name(p) -> str:
    try:
        return p.style.name or ""
    except Exception:
        return ""


def body_paragraphs(doc) -> list[dict]:
    """Body prose paragraphs, each tagged with its enclosing section heading.

    Skips headings and the title (never redline a heading), everything inside a References
    section (a bibliography entry is a generated artifact, not prose the reviewer redlines),
    and empty paragraphs.

    The abstract IS included: it is body prose, and a comment on it must produce a tracked
    change like any other. Returns ``[{index, para, heading, kind}]`` in document order.
    """
    out: list[dict] = []
    heading = ""
    in_references = False
    for i, p in enumerate(doc.paragraphs):
        style = _style_name(p)
        if is_heading_style(style):
            # The title is skipped but does not open a section: the abstract that follows it
            # belongs to no section, not to a section named after the paper.
            # Accepted text, not p.text: a tracked-inserted heading's runs live
            # inside w:ins, which python-docx's .text does not see.
            if not is_title_style(style):
                heading = _accepted_para_text(p._p).strip()
                in_references = guards.is_references(heading)
            continue
        # Accepted text, not p.text: a paragraph whose content sits wholly
        # inside live tracked markup (a full tracked replacement, a reviewer
        # rewrite) is still a body paragraph — .text alone reads it as empty.
        if in_references or not _accepted_para_text(p._p).strip():
            continue
        out.append({
            "index": i,
            "para": p,
            "heading": heading,
            "kind": guards.section_kind(heading),
        })
    return out


def comment_anchors(path: Path) -> list[dict]:
    """Body paragraphs carrying a comment anchor, with comment ids and current text.

    Returns ``[{index, para, heading, kind, ids, text, anchored}]`` in document order.
    ``text`` is the serialized paragraph (atoms as sentinels) — the exact string the reviser
    is asked to revise. ``anchored`` is the sorted set of sentence indices the paragraph's
    comments actually bear on; that set is what the minimal-edit guard enforces against.

    Paragraphs with no comment are omitted, as are headings and References — a comment on a
    heading usually means "add a section" or "find more sources", which is a routing decision,
    not a paragraph edit. The caller handles those; see ``heading_comments``.
    """
    doc = Document(str(path))
    out = []
    for rec in body_paragraphs(doc):
        p_el = rec["para"]._p
        ids = [s.get(qn("w:id")) for s in p_el.findall(qn("w:commentRangeStart"))]
        if not ids:
            continue
        text = paragraph_text(p_el)
        anchored: set[int] = set()
        for span in comment_spans(p_el).values():
            anchored |= anchored_sentences(text, span)
        out.append({**rec, "ids": ids, "text": text, "anchored": sorted(anchored)})
    return out


def heading_comments(path: Path) -> list[dict]:
    """Comments anchored to a heading. These cannot be answered by a paragraph edit — they
    ask for a section to be added, split, or resourced. The caller routes them."""
    doc = Document(str(path))
    out = []
    for i, p in enumerate(doc.paragraphs):
        if not is_heading_style(_style_name(p)):
            continue
        ids = [s.get(qn("w:id")) for s in p._p.findall(qn("w:commentRangeStart"))]
        if ids:
            out.append({"index": i, "heading": _accepted_para_text(p._p).strip(), "ids": ids})
    return out


def accepted_body_text(doc) -> str:
    """The manuscript as it reads with every tracked change accepted."""
    parts = [t for p in doc.paragraphs if (t := _accepted_para_text(p._p)).strip()]
    return "\n\n".join(parts)


def accepted_markdown(doc) -> str:
    """The post-edit manuscript rendered back to markdown, headings and all.

    ``guards`` reason over markdown (``## `` opens a section, which is how the citation floor
    gets gated on section kind and how References are excluded). Recovering that structure
    from the .docx is what lets the redline path emit the same metrics line as the draft path.
    """
    parts: list[str] = []
    for p in doc.paragraphs:
        style = _style_name(p)
        text = _accepted_para_text(p._p).strip()
        if is_title_style(style):
            if text:
                parts.append(f"# {text}")
            continue
        if is_heading_style(style):
            if text:
                parts.append(f"## {text}")
            continue
        accepted = _accepted_para_text(p._p).strip()
        if accepted:
            parts.append(accepted)
    return "\n\n".join(parts)


def tracked_replace(p_el, new_text: str, author: str = AUTHOR, ids: _Ids | None = None) -> bool:
    return _engine.tracked_replace(p_el, new_text, author, ids)


def tracked_replace_sentencewise(p_el, new_text: str, author: str = AUTHOR,
                                 ids: _Ids | None = None) -> bool:
    return _engine.tracked_replace_sentencewise(p_el, new_text, author, ids)


def tracked_insert_after(p_el, text: str, author: str = AUTHOR, ids: _Ids | None = None):
    return _engine.tracked_insert_after(p_el, text, author, ids)
