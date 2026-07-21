"""Phase one of the outline: the paper's sections and subsections, and nothing else.

The outline used to be written in one pass — structure and content beats together — and the
structure was the part nobody could check until a draft had been written from it. A
19-subsection outline against a 5,000-word CFP produced a 6,975-word manuscript, and the
only signal was 4.5 GPU-hours later.

Phase one emits headings alone. That is enough to compute the entire word plan: each
section's share of the body budget, and therefore how many paragraphs each subsection can
afford. The author redlines THAT — cheap to fix, and fixed before a bullet is written.
Phase two (``outline``) adds the bullets, one per manuscript paragraph.

Two conventions the skeleton carries:

  * Headings carry NO numbers. The .docx style supplies them (Heading 1 -> "1",
    Heading 2 -> "1.1"), so a literal "2.1" in the text would render as "2.1 2.1". It also
    retires a whole failure mode: an outline numbered 1.1, 1.3 reads to a drafting model as
    a missing 1.2, and it will helpfully invent one — which is exactly how the css2026 draft
    acquired a section nobody asked for.
  * IBMRDC — Introduction, Background, Methods, Results, Discussion, Conclusion — unless the
    venue says otherwise. An Abstract precedes it; Acknowledgements and References are
    permanent furniture, never planned here and never deleted downstream.
"""

from __future__ import annotations

import re
from pathlib import Path

from . import guards
from .brain import Brain
from .config import ProjectConfig
from .log import log

# The default spine. A venue may override it (`VenueConfig.section_structure`) — Nature's
# format is not IBMRDC, and a venue that mandates its own order is stating a spec, not a
# preference.
IBMRDC = ("Introduction", "Background", "Methods", "Results", "Discussion", "Conclusion")

# Never planned, never deleted, never budgeted. The bibliography is rendered by pandoc at
# write time and the CRediT statement is built from the author list, so neither is the
# outline's business — but both must exist in the manuscript.
FURNITURE = ("Acknowledgements", "References")

_SYSTEM = (
    "You plan the structure of academic papers. You return headings only — never prose, "
    "never bullets, never commentary."
)

_REVISE_PROMPT = """\
Revise a paper's SECTION STRUCTURE to answer the reviewer's annotations. Headings only.

Title: {title}
{budget_section}
Current structure:
{current}

Reviewer annotations — tracked edits they made, and comments they left:
{annotations}

Rules:
- Answer every annotation. A comment asking for a section, a merge, a split or a rename is
  an instruction, not a suggestion — the author owns this structure.
- Keep every heading the annotations do not ask you to change, EXACTLY as written.
- Use EXACTLY these top-level sections, in this order, each as a `## ` heading:
{spine}
- Do NOT number any heading. The document style numbers them.
- Do NOT write bullets, beats, prose, or any body text. Headings only.
- Do NOT add Abstract, Acknowledgements or References: they are added automatically.
- Output only the headings, one per line, starting with the first `## `.
"""


_SKELETON_PROMPT = """\
Plan the SECTION AND SUBSECTION STRUCTURE of an academic paper. Headings only.

Title: {title}
Topic: {topic}
Focus: {focus}
{venue_section}
Structural analysis:
{analysis}
{narrative_section}
{budget_section}
Rules:
- Use EXACTLY these top-level sections, in this order, each as a `## ` heading:
{spine}
- Under each, add `### ` subsections named from this paper's actual content — the \
background_pillars, method_steps and results_structure in the analysis above. A section \
whose material is one continuous argument takes NO subsections; an Introduction of ~2 \
paragraphs never needs them.
- Give each section only as many subsections as its word share affords — the budget above \
states that number outright for each one. Do not exceed it. A subsection costs about \
{min_words} words, because it is at least two paragraphs; one thinner than that is a \
heading with nothing under it, so merge rather than split.
- Do NOT number any heading. The document style numbers them: `## Methods` renders as \
"3 Methods" and `### The Model` as "3.1 The Model". A number you write yourself renders \
twice.
- Do NOT write bullets, beats, prose, or any body text. Headings only — that is the whole \
output. The content plan is a separate pass.
- Do NOT add Abstract, Acknowledgements or References: they are added automatically.
- Output only the headings, one per line, starting with the first `## `.
"""


def spine_for(cfg: ProjectConfig, venue: str = "") -> tuple[str, ...]:
    """The top-level sections this paper uses.

    IBMRDC unless the venue states its own structure. A venue's mandated order is a SPEC
    read off its call for papers, not a preference — see ``VenueConfig.section_structure``.
    """
    v = cfg.venue(venue) if venue else None
    stated = (getattr(v, "section_structure", "") or "").strip() if v else ""
    if not stated:
        return IBMRDC
    parts = [p.strip() for p in stated.replace("\n", ",").split(",") if p.strip()]
    return tuple(parts) or IBMRDC


def assemble(sections: list[tuple[int, str]], title: str) -> str:
    """The skeleton document: title, abstract, the planned sections, then the furniture.

    The Abstract and the furniture are added HERE rather than asked of the model: they are
    fixed, and a model that can add them can also forget them, rename them, or number them.
    """
    lines = [f"# {title}", "", "## Abstract", ""]
    for level, text in sections:
        lines += ["#" * level + " " + text, ""]
    for f in FURNITURE:
        lines += [f"## {f}", ""]
    return "\n".join(lines).rstrip() + "\n"


def parse_headings(raw: str, spine: tuple[str, ...]) -> list[tuple[int, str]]:
    """(level, text) for every heading the model returned, numbers stripped.

    A model told not to number will number anyway often enough to matter, and a stray "2.1"
    in the text renders as "2.1 2.1" once the style adds its own. Stripping is cheaper than
    a retry and cannot be got wrong.
    """
    import re
    out: list[tuple[int, str]] = []
    for raw_line in raw.splitlines():
        line = raw_line.strip()
        if not line.startswith("#"):
            continue
        level = len(line) - len(line.lstrip("#"))
        text = re.sub(r"^\d+(\.\d+)*\.?\s*", "", line[level:].strip()).strip()
        if not text or level < 2:
            continue
        if guards.is_abstract(text) or guards.is_references(text) \
                or guards.is_acknowledgements(text):
            continue                      # added by assemble(), never by the model
        out.append((min(level, 4), text))
    return out


def findings(sections: list[tuple[int, str]], spine: tuple[str, ...],
             budget: int, shares: dict | None = None) -> list[guards.Finding]:
    """What is mechanically wrong with a skeleton. PHASE: skeleton.

    The whole point of phase one: every one of these is computable from headings alone,
    before a bullet or a sentence exists.
    """
    out: list[guards.Finding] = []
    tops = [t for lvl, t in sections if lvl == 2]

    missing = [s for s in spine if s.lower() not in [t.lower() for t in tops]]
    for s in missing:
        out.append(guards.Finding(
            "missing-section", s,
            f'The structure calls for a "{s}" section and the skeleton has none. Add it.'))
    extra = [t for t in tops if t.lower() not in [s.lower() for s in spine]]
    for t in extra:
        out.append(guards.Finding(
            "invented-section", t,
            f'"{t}" is not part of this paper\'s structure. Remove it, or fold its '
            f"material into the section that does cover it."))

    if budget > 0:
        counts: dict[str, int] = {}
        current = None
        for lvl, text in sections:
            if lvl == 2:
                current = text
                counts.setdefault(current, 0)
            elif current is not None:
                counts[current] += 1
        lo, hi = guards.PARAGRAPH_BAND
        for top, n in counts.items():
            words = guards.section_words(top, budget, shares)
            if not words:
                continue
            # One row of the word plan, checked. A subsection costs MIN_BULLETS_PER_SUBSECTION
            # paragraphs, so the section's paragraph width follows from how many subsections
            # it has — and the author is choosing that width whether or not they know it.
            bullets = guards.MIN_BULLETS_PER_SUBSECTION * max(n, 1)
            each = words // bullets
            want = max(1, round(words / guards.subsection_words()))
            if each > hi:
                out.append(guards.Finding(
                    "subsections-too-few", top,
                    f'"{top}" carries {words} words across {n or 1} subsection(s) — at '
                    f"{guards.MIN_BULLETS_PER_SUBSECTION} bullets each that is {bullets} "
                    f"paragraphs of {each} words, over the {hi} a paragraph may run. Split "
                    f"it into {want} subsection(s)."))
            elif each < lo:
                out.append(guards.Finding(
                    "subsections-too-thin", top,
                    f'"{top}" carries {words} words across {n} subsections — at '
                    f"{guards.MIN_BULLETS_PER_SUBSECTION} bullets each that is {bullets} "
                    f"paragraphs of {each} words, under the {lo} a paragraph needs. Merge "
                    f"to {want} subsection(s)."))
    return out


def plan_table(sections: list[tuple[int, str]], budget: int,
               shares: dict | None = None) -> str:
    """The word plan the skeleton implies — what the author is really approving."""
    rows, current, subs = [], None, []
    ordered: list[tuple[str, list[str]]] = []
    for lvl, text in sections:
        if lvl == 2:
            current = text
            subs = []
            ordered.append((current, subs))
        elif current is not None:
            subs.append(text)
    for top, kids in ordered:
        words = guards.section_words(top, budget, shares)
        if not words:
            continue          # furniture: no body-prose share, so nothing to plan
        bullets = guards.MIN_BULLETS_PER_SUBSECTION * max(len(kids), 1)
        rows.append(f"  {top:<28}{words:>6} words  {len(kids):>2} sub  "
                    f"{bullets:>3} bullets  {words // bullets:>4} each")
    return "\n".join(rows)


# ── the verb ─────────────────────────────────────────────────────────────────

def _render_docx(md_path: Path, project_dir: Path,
                 notes: list[tuple[str, str]] | None = None) -> Path | None:
    """Render the .docx and drop the markdown behind it.

    The skeleton's deliverable is the .docx. It is what the author marks up, what the gate
    reads, and what every later stage reads. The .md was pandoc's input and this stage's
    "have I run?" marker, and nothing else ever read it — ``outline`` reads the RELEASE, and
    ``_revise`` reads the annotated .docx. A second file per rung that nothing consumes can
    only drift from the one that matters.

    The .md survives a FAILED render: without pandoc it is the only output there is, and
    deleting it would leave the author nothing to look at.
    """
    from .refdoc import render
    docx = render(md_path, project_dir)
    if docx is None:
        log(f"[raconteur] wrote {md_path.relative_to(project_dir)} (no .docx — pandoc)")
        return None
    md_path.unlink(missing_ok=True)
    if notes:
        from haarpi import redline as _rl
        try:
            n = _rl.add_anchored_comments(docx, notes, author="raconteur", initials="ra")
            log(f"[raconteur] attached {n} word-plan comment(s)")
        except Exception as e:                # noqa: BLE001 — a comment must not fail a render
            log(f"[warn] could not attach word-plan comments: {type(e).__name__}: {e}")
    log(f"[raconteur] wrote {docx.relative_to(project_dir)}")
    return docx


# ── the word plan, carried on the document ───────────────────────────────────
# The plan the author approves has to live where they approve it. A log line is seen once,
# at generation, by whoever was watching the terminal; the .docx is what gets opened, marked
# up and gated. So each section heading carries a comment stating its row of the plan.
#
# WORDS PER BULLET is the invariant, pinned here from the section's share and the structure
# as generated. Everything else derives from it and the headings, which is why adding a
# subsection ADDS words rather than thinning the paragraphs already there: a new subsection
# is two more bullets at this section's established rate. Recomputing the rate from the
# share instead would hold the section still and starve it, which is the opposite of what a
# structural edit means.

_PLAN_RE = re.compile(r"(\d+)\s*(?:words\s*)?each", re.IGNORECASE)
_PLAN_SUB_RE = re.compile(r"(\d+)\s*sub\b", re.IGNORECASE)


def words_per_bullet(sections: list[tuple[int, str]], budget: int,
                     shares: dict | None = None) -> dict[str, int]:
    """Each section's pinned rate: its share divided by the bullets its structure affords."""
    out: dict[str, int] = {}
    for top, kids in _by_section(sections):
        words = guards.section_words(top, budget, shares)
        if not words:
            continue
        # Floored at one paragraph. A section generated wider than its share affords would
        # otherwise pin a rate below what a paragraph needs — Background came back with four
        # subsections on a 600-word share and pinned 75 — and that rate then outlives the
        # mistake, so merging back to two would HALVE the section instead of restoring it.
        # Flooring lets the section state its true cost and makes merging the fix it should
        # be: two subsections at 150 is exactly the 600 the share intended.
        out[top] = max(guards.WORDS_PER_PARAGRAPH,
                       words // (guards.MIN_BULLETS_PER_SUBSECTION * max(len(kids), 1)))
    return out


def _by_section(sections: list[tuple[int, str]]) -> list[tuple[str, list[str]]]:
    ordered: list[tuple[str, list[str]]] = []
    current, subs = None, []
    for lvl, text in sections:
        if lvl == 2:
            current, subs = text, []
            ordered.append((current, subs))
        elif current is not None:
            subs.append(text)
    return ordered


def plan_row(top: str, n_subs: int, wpb: int) -> str:
    """One section's plan, as the author reads it on the heading."""
    bullets = guards.MIN_BULLETS_PER_SUBSECTION * max(n_subs, 1)
    return (f"{top} — {bullets * wpb} words · {n_subs} sub · {bullets} bullets · "
            f"{wpb} each. Add a subsection and this section grows by "
            f"{guards.MIN_BULLETS_PER_SUBSECTION * wpb} words; remove one and it shrinks by "
            f"the same. Words per bullet is fixed for this section.")


def plan_notes(sections: list[tuple[int, str]], budget: int,
               shares: dict | None = None) -> list[tuple[str, str]]:
    """(heading, comment) for every section that spends body prose."""
    wpb = words_per_bullet(sections, budget, shares)
    return [(top, plan_row(top, len(kids), wpb[top]))
            for top, kids in _by_section(sections) if top in wpb]


def read_plan(text: str) -> tuple[int | None, int | None]:
    """(words per bullet, subsections) as a plan comment states them.

    Forgiving on purpose: the author edits these by hand and will annotate their reasoning
    beside the number. The first "N each" and "N sub" win; anything else in the comment is
    theirs. Returns (None, None) when there is no number to find, which the caller reports
    as an instruction rather than guessing.
    """
    w = _PLAN_RE.search(text or "")
    n = _PLAN_SUB_RE.search(text or "")
    return (int(w.group(1)) if w else None, int(n.group(1)) if n else None)


def document_words(sections: list[tuple[int, str]], wpb: dict[str, int]) -> int:
    """Body prose the structure now implies, at each section's pinned rate."""
    return sum(guards.MIN_BULLETS_PER_SUBSECTION * max(len(kids), 1) * wpb[top]
               for top, kids in _by_section(sections) if top in wpb)


def plan_from_release(path) -> tuple[dict[str, int], list[str]]:
    """The word plan the author gated, read back off the minted skeleton.

    Returns (words-per-bullet by section, problems). The comment is authoritative for ONE
    number — the rate — because it is the only one that cannot be recovered from the
    document. Recomputing it from the share after a subsection was added would hold the
    section still and thin its paragraphs, when adding a subsection is meant to add words.

    Everything else in the comment is CHECKED against the document, never trusted. The
    comment states how many subsections the plan was written for; the headings state how
    many there are now. When they disagree the author edited one and not the other, and the
    only safe move is to stop and say so — a plan describing a structure that no longer
    exists is worse than no plan, because it looks like one.
    """
    from docx import Document
    from haarpi import redline as hrl
    # heading_comments, not comment_anchors: the latter omits headings by design, because a
    # comment on one asks for a section to be added or split rather than for prose to be
    # rewritten. The word plan is exactly that kind of comment.
    from .redline import heading_comments

    doc = Document(str(path))
    paras = doc.paragraphs
    subs: dict[str, int] = {}
    current = None
    for para in paras:
        text = para.text.strip()
        style = para.style.name if para.style is not None else ""
        if not text or not style.startswith("Heading"):
            continue
        level = int("".join(c for c in style if c.isdigit()) or 1)
        if level == 2:
            current = text
            subs.setdefault(current, 0)
        elif level >= 3 and current is not None:
            subs[current] = subs.get(current, 0) + 1

    cmap = hrl.comments_by_id(path)
    wpb: dict[str, int] = {}
    problems: list[str] = []
    for a in heading_comments(path):
        head = a["heading"]
        if head not in subs:
            continue
        for cid in a["ids"]:
            rec = cmap.get(str(cid))
            if not rec:
                continue
            rate, stated = read_plan(rec.get("text", ""))
            if rate is None:
                problems.append(
                    f'"{head}": its word-plan comment states no rate. It must contain '
                    f'"N each" — the words one bullet is worth in this section.')
            elif stated is not None and stated != subs[head]:
                problems.append(
                    f'"{head}": its word plan was written for {stated} subsection(s) and '
                    f"the skeleton now has {subs[head]}. Whichever you changed, the other "
                    f"was not updated — fix the comment or the headings so they agree.")
            else:
                wpb[head] = rate
    for head, n in subs.items():
        if head in wpb:
            continue
        if guards.is_abstract(head) or guards.is_references(head) \
                or guards.is_acknowledgements(head):
            continue
        if not any(head in pr for pr in problems):
            problems.append(
                f'"{head}" carries no word plan. Every section that spends body prose is '
                f"planned at the skeleton; one that is not has no gated length.")
    return wpb, problems


def strip_plan_comments(path) -> int:
    """Consume the plan: the outline has read it, so it stops travelling.

    The comments describe a skeleton. Carried into the outline they would describe a
    document that has moved on, and a stale plan beside live content is the drift this
    whole mechanism exists to prevent.
    """
    from docx import Document
    from haarpi import redline as hrl
    doc = Document(str(path))
    n = hrl.strip_comment_anchors(doc)
    doc.save(str(path))
    hrl._clear_comment_parts(path)
    return n


def reconcile_plan(path) -> int:
    """Rewrite each section's word plan to match the structure as APPROVED. The mint's job.

    The comments were written against the structure as GENERATED. The author then moved
    subsections around, and a plan describing a shape the document no longer has is worse
    than no plan, because it looks like one. Rather than make the next rung refuse — which
    would turn a routine structural edit into an error the author has to clear by hand-
    editing six comments — the mint reconciles: the release always agrees with itself.

    WORDS PER BULLET is carried, never recomputed. That is the whole invariant: re-deriving
    it from the share after a subsection was added would hold the section still and thin its
    paragraphs, when adding a subsection is supposed to add words. It is read back off the
    comment the author approved — and if they edited that number, theirs is the one that
    survives. A section they added outright has no pinned rate, so it takes one from its
    share, exactly as a fresh skeleton would.
    """
    from docx import Document
    from haarpi import redline as hrl
    from .redline import heading_comments

    doc = Document(str(path))
    subs: dict[str, int] = {}
    current = None
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style is not None else ""
        if not text or not style.startswith("Heading"):
            continue
        level = int("".join(c for c in style if c.isdigit()) or 1)
        if level == 2:
            current = text
            subs.setdefault(current, 0)
        elif level >= 3 and current is not None:
            subs[current] = subs.get(current, 0) + 1

    cmap = hrl.comments_by_id(path)
    updates: dict[str, str] = {}
    for a in heading_comments(path):
        head = a["heading"]
        if head not in subs:
            continue
        for cid in a["ids"]:
            rec = cmap.get(str(cid))
            if not rec:
                continue
            rate, _ = read_plan(rec.get("text", ""))
            if rate is None:
                continue                 # unreadable: leave the author's words alone
            updates[str(cid)] = plan_row(head, subs[head], rate)
    return hrl.set_comment_text(path, updates)


def _write(project_dir: Path, cfg: ProjectConfig, work_dir: Path, text: str,
           venue: str = "") -> Path:
    from .naming import major_skeleton_name
    from .refdoc import render
    # No byline. A skeleton is a title and a heading structure; authors, affiliations and
    # the corresponding address are derived from the project manifest by
    # ``load_authors_block`` and regenerated at every stage that needs them. Writing a copy
    # here gave authorship a second, older home that rode the whole ladder: the css2026
    # skeleton released two affiliations and a gmail address against the manifest's three
    # and an institutional one, and the author had touched none of those paragraphs,
    # because there is nothing in them to review. An edit made to them would have been
    # silently overwritten downstream — the one span the redline contract could not keep
    # its promise about.
    out = work_dir / major_skeleton_name(cfg.short_title, "md", venue=venue)
    out.write_text(text, encoding="utf-8")
    sections = [(len(h) - len(h.lstrip("#")), h.lstrip("# ").strip())
                for h in text.splitlines() if h.startswith("##")]
    notes = plan_notes(sections, _budget_for(cfg, venue), cfg.section_shares or None)
    return _render_docx(out, project_dir, notes) or out


def _revise(project_dir: Path, cfg: ProjectConfig, brain, work: Path,
            user_rev: Path, venue: str = "") -> None:
    """Answer a redlined skeleton — a minor version, keeping the source's datestamp.

    A structural redline is cheap to answer and expensive to ignore: it is the author
    saying the paper is organised wrongly, at the one moment when fixing that costs a
    heading rather than a draft. Without this the gate dead-ended, and an author whose
    markup carried a comment rather than a clean edit had nowhere to go.
    """
    from .naming import minor_name, parse
    from .revise import build_revision_context, read_text

    annotations = build_revision_context(user_rev)
    if not annotations:
        log("[warn] no annotations in the redlined skeleton — nothing to answer")
        log("[error] nothing to do: this run made no changes (exit 3)")
        raise SystemExit(3)

    spine = spine_for(cfg, venue)
    budget = _budget_for(cfg, venue)
    from .outline import _budget_block
    raw = brain.coordinator(
        _REVISE_PROMPT.format(
            title=cfg.title,
            budget_section=_budget_block(cfg, project_dir, venue),
            current=read_text(user_rev),
            annotations=annotations,
            spine="\n".join(f"  - {s}" for s in spine),
        ),
        system=_SYSTEM, num_ctx=16384)

    sections = parse_headings(raw, spine)
    for f in findings(sections, spine, budget, cfg.section_shares or None):
        log(f"[warn] {f.kind} — {f.where}: {f.imperative}")
    if budget:
        log("[raconteur] word plan:")
        for line in plan_table(sections, budget, cfg.section_shares or None).splitlines():
            log(line)

    parsed = parse(user_rev, cfg.short_title)
    chain, datestamp = (parsed[1], parsed[0]) if parsed else ([], None)
    out = work / minor_name(cfg.short_title, chain, "md", datestamp)
    out.write_text(assemble(sections, cfg.title), encoding="utf-8")
    _render_docx(out, project_dir,
                 plan_notes(sections, budget, cfg.section_shares or None))


def _budget_for(cfg: ProjectConfig, venue: str) -> int:
    """This venue's body-prose budget, or 0 where it states no length."""
    v = cfg.venue(venue) if venue else None
    if not v or not v.word_limit:
        return 0
    return guards.prose_budget(guards.word_target(v.word_min, v.word_limit))


def run(project_dir: Path, venue: str = "") -> None:
    """Phase one: plan the paper's sections and subsections against the word budget."""
    from .brain import Brain
    from .config import GlobalConfig
    from .context import (check_prerequisites, load_litreview, load_methods,
                          load_onepager, load_results)
    from .naming import deliverable_dir, find_latest, find_user_revision
    from .outline import _analyze_structure, _build_venue_section, _budget_block

    if not ProjectConfig.exists(project_dir):
        log("[error] no paper/raconteur.yaml found — run 'raconteur init' first")
        raise SystemExit(1)
    cfg = ProjectConfig.load(project_dir)
    gcfg = GlobalConfig.load()
    check_prerequisites(project_dir, cfg)

    narrative = load_onepager(project_dir, cfg.short_title)
    if not narrative:
        log("[error] no one-pager found — run 'raconteur onepager' first")
        raise SystemExit(1)

    from . import slate
    venue = slate.resolve(cfg, venue)
    if venue:
        log(f"[raconteur] planning structure for {cfg.venues[venue].name} ({venue})")

    work = deliverable_dir(project_dir / "paper", "skeleton", venue)
    work.mkdir(parents=True, exist_ok=True)
    scope = ([venue] if venue else []) + ["skeleton"]
    others = [v for v in cfg.venues if v != venue]
    # The .docx is the deliverable now — the .md is deleted after it renders. Counting
    # markdown here would report every completed skeleton as absent and re-draft it,
    # spending a GPU run to overwrite work the author may already have marked up.
    existing = find_latest(work, cfg.short_title, "docx", last_initials="ra",
                           chain_includes=scope, chain_excludes=others)
    user_rev = find_user_revision(work, cfg.short_title, chain_includes=scope,
                                  chain_excludes=others)
    brain = Brain(gcfg, coordinator=cfg.brain.coordinator_model)

    if existing and user_rev:
        log(f"[raconteur] answering annotations on {user_rev.name}")
        _revise(project_dir, cfg, brain, work, user_rev, venue)
        return
    if existing:
        log("[raconteur] skeleton exists — annotate the docx with your initials, "
            "then run `haarpi next`")
        log("[error] nothing to do: this run made no changes (exit 3)")
        raise SystemExit(3)
    litrev = load_litreview(project_dir, cfg.litrev_dir) if cfg.litrev_dir else ""
    code = load_methods(project_dir) if cfg.use_methods else ""
    results = load_results(project_dir, cfg.results_dir) if cfg.results_dir else ""

    log("[raconteur] analysing paper structure…")
    analysis = _analyze_structure(brain, cfg.description, litrev, code, results,
                                  narrative, None, project_dir)

    spine = spine_for(cfg, venue)
    budget = _budget_for(cfg, venue)

    log("[raconteur] planning sections…")
    raw = brain.coordinator(
        _SKELETON_PROMPT.format(
            title=cfg.title, topic=cfg.topic, focus=cfg.focus,
            venue_section=_build_venue_section(cfg, project_dir, venue),
            analysis=analysis,
            narrative_section=f"Narrative spine (author-approved):\n{narrative}\n",
            budget_section=_budget_block(cfg, project_dir, venue),
            spine="\n".join(f"  - {s}" for s in spine),
            min_words=guards.subsection_words(),
        ),
        system=_SYSTEM, num_ctx=16384)

    sections = parse_headings(raw, spine)
    for f in findings(sections, spine, budget, cfg.section_shares or None):
        log(f"[warn] {f.kind} — {f.where}: {f.imperative}")
    if budget:
        log("[raconteur] word plan:")
        for line in plan_table(sections, budget, cfg.section_shares or None).splitlines():
            log(line)
    _write(project_dir, cfg, work, assemble(sections, cfg.title), venue)
