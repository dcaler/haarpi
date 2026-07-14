"""The author's voice, measured from what they DO write.

The first cut of this module tried to learn what the author never writes — contrast their
corpus against the tool's, and ban what appears in one and not the other. It produced three
thousand "nevers", of which the top entries were `jaccard distance`, `harmonic similarity`
and `chord distance metrics`. The author had never written those because he had never
written about music. The method could not tell STYLE from TOPIC, and banning a paper's own
vocabulary in the name of the author's voice is a spectacular way to fail.

Absence is a bad instrument. It needs a baseline, it needs a frequency floor, it stays
fragile, and it is a weak lever besides: a model told not to write "moreover" writes
"furthermore".

What the author DOES write is directly observable, needs no baseline, and cannot be
contaminated by topic — provided we measure the right things. So everything counted here is
CLOSED-CLASS: discourse markers, hedges, intensifiers, sentence rhythm, punctuation. There
are only so many ways to say "however", and none of them is a domain term. A paper about
chords and a paper about solar rebates draw on the same small stock of connectives, and how
an author moves between ideas is style in its purest, most portable form.

The output is a PALETTE, not a prohibition: these are the transitions this author uses, at
these rates; this is the rhythm of their sentences; here are passages of the real thing. A
model handed a palette paints with it. A model handed a blacklist finds a synonym.
"""

from __future__ import annotations

import re
import statistics
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from haarpi.redline import TOOL_AUTHORS
from haarpi.text import sentence_units

from .log import log

# ── the closed classes ───────────────────────────────────────────────────────
# Style-bearing and topic-proof. A word here cannot be domain vocabulary, so a count is a
# fact about the author's voice and nothing else.

CONNECTIVES = (
    "however", "moreover", "furthermore", "additionally", "in addition",
    "thus", "therefore", "hence", "accordingly", "consequently", "as a result",
    "nonetheless", "nevertheless", "still", "yet",
    "in contrast", "by contrast", "conversely", "on the other hand",
    "similarly", "likewise", "in turn", "meanwhile",
    "overall", "in short", "in sum", "taken together", "on balance",
    "finally", "ultimately", "in the end",
    "importantly", "notably", "crucially", "critically", "of note",
    "indeed", "in fact", "that said", "even so", "to be clear",
    "first", "second", "third", "for example", "for instance", "such as",
    "that is", "in particular", "specifically", "more broadly",
    "as such", "given this", "to this end", "in this way",
)

HEDGES = (
    "may", "might", "could", "can", "appears", "appear", "seems", "seem",
    "suggests", "suggest", "indicates", "indicate", "implies", "imply",
    "likely", "unlikely", "plausibly", "arguably", "presumably", "possibly",
    "tends to", "tend to", "broadly", "largely", "generally", "typically",
    "we think", "we believe", "we argue", "we find", "we show", "we note",
    "to some extent", "at least in part", "in principle",
)

INTENSIFIERS = (
    "very", "highly", "extremely", "remarkably", "strikingly", "profoundly",
    "vitally", "essentially", "fundamentally", "clearly", "obviously",
    "substantially", "considerably", "markedly", "dramatically",
)

_WORD = re.compile(r"[a-z][a-z'’-]*")

# Where markup lives. Walking a whole project means walking results/data — tens of thousands
# of json files — to find a handful of .docx.
_SKIP_DIRS = ("data", "pdfs", "work", "figures", "@eaDir", "__pycache__", ".git")


_REFS_TAIL = re.compile(r"\n\s*(references|bibliography|works cited)\s*\n", re.I)
_CITATION = re.compile(r"\([^()]*\d{4}[a-z]?(;[^()]*)?\)")     # (Smith 2020; Jones 2019)
_NUM_CITE = re.compile(r"\[\d+(?:[,–-]\s*\d+)*\]")             # [3], [3, 5], [3–7]


_FURNITURE = re.compile(
    r"^(figure|table|fig\.|tab\.)\s*\d"      # captions
    r"|^\s*\d+\s*$"                          # page numbers
    r"|^(downloaded|licensed|copyright|©|doi:|https?://)", re.I)


def clean_prose(text: str) -> str:
    """A PDF's raw text is not prose, and measuring it measures the PDF.

    Page numbers, captions, reference lists and citations are all "text", and every one of
    them lies about how a person writes: they make sentences look four words long and
    parentheses three times more common than they are (an author-year citation is a
    citation, not a parenthetical aside).

    A PDF is also HARD-WRAPPED — a line is a typesetting artefact, not a unit of thought —
    so paragraph structure cannot be recovered from it and this does not pretend to. The
    lines are rejoined into flowing prose, hyphens broken across lines are healed, and what
    comes back supports sentence-level measurement and nothing finer.
    """
    text = _REFS_TAIL.split(text, maxsplit=1)[0]
    text = _CITATION.sub("", text)
    text = _NUM_CITE.sub("", text)

    keep: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line or _FURNITURE.match(line):
            continue
        if line.isupper() and len(line.split()) < 12:      # running head / section title
            continue
        digits = sum(c.isdigit() for c in line)
        if digits > len(line) * 0.4:                       # a table row, not a sentence
            continue
        keep.append(line)

    prose = " ".join(keep)
    prose = re.sub(r"(\w)-\s+(\w)", r"\1\2", prose)        # heal hyphens broken by wrapping
    return re.sub(r"\s+", " ", prose).strip()


def pdf_prose(path: Path) -> str:
    """The prose of a PDF, WITH its paragraphs — read from the file, not from an index.

    A PDF's text is laid out, not written: a line is where the typesetter broke it, and a
    flat text dump (Zotero's fulltext index, or `page.get_text()`) hands back those breaks as
    if they meant something. Measure that and you measure the typesetting — four-word
    sentences, paragraphs fifteen sentences deep.

    PyMuPDF's BLOCK extraction gives what is actually wanted: a block is a paragraph, laid
    out as one, with its own bounding box. Rejoin the lines inside each block, keep the
    blocks that are prose, and the author's paragraph structure comes back intact.

    rabbitHole has read PDFs with PyMuPDF all along (see rabbithole/pdfs.py); it never needed
    this because it hands the mess to a model and asks for a quote. Measurement is less
    forgiving.
    """
    try:
        import fitz
    except ImportError:
        log("[warn] PyMuPDF not installed — cannot read the author's PDFs")
        return ""
    try:
        doc = fitz.open(str(path))
    except Exception as e:  # noqa: BLE001
        log(f"[warn] could not open {Path(path).name}: {e}")
        return ""

    paras: list[str] = []
    for page in doc:
        for block in page.get_text("blocks"):
            text = (block[4] or "").strip()
            if not text or _FURNITURE.match(text) or text.isupper():
                continue
            # rejoin the lines the layout broke, healing hyphenation
            text = re.sub(r"-\n(\w)", r"\1", text)
            text = re.sub(r"\s*\n\s*", " ", text).strip()
            if len(text.split()) < 25:            # a caption, a heading, an author line
                continue
            digits = sum(c.isdigit() for c in text)
            if digits > len(text) * 0.25:         # a table
                continue
            paras.append(text)
    doc.close()

    body = "\n\n".join(paras)
    body = _REFS_TAIL.split(body, maxsplit=1)[0]
    return _tidy(body)


def _tidy(text: str) -> str:
    """Heal what the typesetter and the citation-stripper broke.

    A soft hyphen is how a PDF hyphenates across a line break — "policy diffu\xadsion" — and
    it survives extraction invisibly, so the word looks whole and reads as two. And removing
    a citation leaves its punctuation behind: "water reuse (Smith 2020), water reuse" becomes
    "water reuse , water reuse". Both make prose that is fine to count and grotesque to quote,
    which matters the moment an exemplar of it is put in front of a model as the voice to
    match.
    """
    text = text.replace("\xad", "").replace("​", "")
    for lig, plain in (("ﬁ", "fi"), ("ﬂ", "fl"), ("ﬀ", "ff"), ("ﬃ", "ffi"), ("ﬄ", "ffl"),
                       ("’", "'"), ("‘", "'"), ("“", '"'), ("”", '"'), ("''", '"')):
        text = text.replace(lig, plain)          # PDF ligatures survive extraction as one glyph
    text = _CITATION.sub("", text)
    text = _NUM_CITE.sub("", text)
    text = re.sub(r"\(\s*[;,]?\s*\)", "", text)        # parentheses emptied of a citation
    text = re.sub(r"\s+([,.;:)])", r"\1", text)        # space left before punctuation
    text = re.sub(r"([(\[])\s+", r"\1", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def pick_exemplars(corpus: str, n: int = 3, lo: int = 45, hi: int = 110) -> list[str]:
    """Passages of the author's PUBLISHED prose, to show a model what to sound like.

    Published, deliberately. The author's hand-typed insertions in a redline are their voice
    too, and in domain — but they are a DRAFT: they carry the typos we are separately
    offering to fix ("new though patterns", "it's visual output"). Holding those up as the
    voice to imitate would teach the model the author's mistakes and then flag them.

    Prefers paragraphs that carry the author's own connectives — those are the ones where
    the prose is doing the work of moving between ideas, which is what a model most needs to
    see.
    """
    paras = [p.strip() for p in corpus.split("\n\n")
             if lo <= len(p.split()) <= hi]
    # A layout block can begin mid-sentence, where the previous column left off. Quoting
    # "public reaction to the response. The first source of uncertainty…" at a model as the
    # voice to match teaches it to start paragraphs in the middle of a thought.
    paras = [p for p in paras if p[:1].isupper()]
    if not paras:
        return []

    def score(p: str) -> tuple:
        marks = sum(count_phrase(p, c) for c in CONNECTIVES)
        sents = len([u for u in sentence_units(p) if len(u.split()) >= 4])
        return (marks >= 1, 2 <= sents <= 6, marks)

    return [p for p in sorted(paras, key=score, reverse=True)[:n]]


def author_corpus(fulltexts: list[str]) -> str:
    """Clean each paper, THEN join them.

    Cleaning a CONCATENATION truncates at the first paper's reference list and silently
    discards every paper after it — nine papers in, one paper out, and a signature that
    confidently reports the author never uses words they use on every page.
    """
    return "\n\n".join(p for p in (clean_prose(t) for t in fulltexts) if p)


def _words(text: str) -> list[str]:
    return _WORD.findall((text or "").lower())


def _rate(count: int, n_words: int) -> float:
    """Occurrences per 10,000 words — comparable across corpora of any size."""
    return round(count / n_words * 10_000, 2) if n_words else 0.0


def count_phrase(text: str, phrase: str) -> int:
    """Word-normalised, so punctuation and case cannot hide an occurrence."""
    hay = " ".join(_words(text))
    needle = " ".join(_words(phrase))
    if not needle:
        return 0
    return len(re.findall(rf"(?<!\w){re.escape(needle)}(?!\w)", hay))


# ── the measured signature ───────────────────────────────────────────────────

def inventory(text: str, candidates: tuple[str, ...]) -> dict[str, float]:
    """Which of a closed class the author uses, and how often (per 10k words).

    Only what they USE. A zero is reported by omission, and a caller that wants to know
    whether the author ever writes "moreover" asks whether it is in the palette — which is a
    measured fact about their prose, not an inference from a model's impression of it.
    """
    n = len(_words(text))
    out = {}
    for c in candidates:
        k = count_phrase(text, c)
        if k:
            out[c] = _rate(k, n)
    return dict(sorted(out.items(), key=lambda kv: -kv[1]))


def rhythm(text: str) -> dict:
    """The metre of the prose: sentence length, its spread, and paragraph shape.

    A draft of uniform 18-word sentences does not read like a person wrote it, and no amount
    of "match this author's voice" in a prompt will tell a model that its rhythm is flat. A
    number will.

    Paragraph shape is only measured when the source HAS paragraphs — ``pdf_prose`` recovers
    them from the PDF's layout blocks. Zotero's flat fulltext index has none (nine papers
    come back as nine blobs of 288 sentences each), and a figure computed from that would be
    a fact about the indexer.
    """
    units = [u for u in sentence_units(text) if len(u.split()) >= 4]
    if len(units) < 40:
        return {}
    lengths = sorted(len(u.split()) for u in units)
    out = {
        "sentence_words_mean": round(statistics.mean(lengths)),
        "sentence_words_p10": lengths[int(len(lengths) * 0.10)],
        "sentence_words_p90": lengths[int(len(lengths) * 0.90)],
    }
    paras = [p for p in text.split("\n\n") if len(p.split()) >= 25]
    if len(paras) >= 20:          # real paragraphs, not one blob per paper
        per = [len([u for u in sentence_units(p) if len(u.split()) >= 4]) for p in paras]
        per = [n for n in per if n]
        if per and statistics.mean(per) < 20:
            out["sentences_per_paragraph"] = round(statistics.mean(per), 1)
    return out


def punctuation(text: str) -> dict[str, float]:
    """Habits a model will not copy unless told: dashes, semicolons, colons, parentheses."""
    n = len(_words(text))
    if not n:
        return {}
    return {
        "em_dash": _rate(len(re.findall(r"—|--", text)), n),
        "semicolon": _rate(text.count(";"), n),
        "colon": _rate(len(re.findall(r":(?!\d)", text)), n),
        "parenthetical": _rate(text.count("("), n),
    }


def signature(text: str, clean: bool = True) -> dict:
    """Everything countable about how this author writes. Topic-proof by construction.

    ``clean`` strips the PDF's furniture first (see ``clean_prose``) — without it the
    numbers describe the typesetting, not the author.
    """
    prose = clean_prose(text) if clean else text
    if len(_words(prose)) < 2000:
        log(f"[warn] voice: only {len(_words(prose))} words of usable prose — "
            f"the signature will be noisy")
    return {
        "corpus_words": len(_words(prose)),
        **rhythm(prose),
        "connectives": inventory(prose, CONNECTIVES),
        "hedges": inventory(prose, HEDGES),
        "intensifiers": inventory(prose, INTENSIFIERS),
        "punctuation": punctuation(prose),
    }


def outside_palette(text: str, palette: dict[str, float],
                    candidates: tuple[str, ...]) -> list[str]:
    """Closed-class markers this draft uses that the author does not.

    The positive form of the old "never" list, and the reason it is sound: the candidate set
    is closed, so a hit is a style fact and never a domain word. `harmonic similarity` cannot
    appear here; `moreover` can.
    """
    return sorted({c for c in candidates
                   if c not in palette and count_phrase(text, c)})


def style_block(sig: dict, exemplars: list[str], analysis: str = "",
                budget: int = 3800) -> str:
    """The voice, as the drafter must receive it: a palette and some real prose.

    Exemplars FIRST and never truncated mid-passage. The profile used to be capped at 2,000
    characters with the verbatim excerpts at the END of the file, so they were exactly what
    got cut: the model was handed a 350-word DESCRIPTION of the author's prose and not one
    sentence of the prose itself. A description of a style is a feeble lever on a language
    model; three sentences of the real thing are worth more than all of it.
    """
    if not sig and not exemplars:
        return ""
    parts: list[str] = ["Writing style — THIS AUTHOR'S. Match it."]

    if sig.get("sentence_words_mean"):
        line = (f"\nRHYTHM: sentences average {sig['sentence_words_mean']} words and range "
                f"from {sig['sentence_words_p10']} to {sig['sentence_words_p90']}. Vary them "
                f"— a paragraph of even-length sentences does not read like this author.")
        if sig.get("sentences_per_paragraph"):
            line += (f" His paragraphs run about "
                     f"{sig['sentences_per_paragraph']} sentences.")
        parts.append(line)

    def _palette(label: str, key: str, candidates: tuple[str, ...], n: int = 10) -> None:
        pal = sig.get(key) or {}
        if not pal:
            return
        uses = ", ".join(list(pal)[:n])
        avoids = [c for c in candidates if c not in pal][:8]
        line = f"\n{label} HE USES: {uses}."
        if avoids:
            line += (f"\n{label} HE NEVER USES — do not write them: "
                     f"{', '.join(avoids)}.")
        parts.append(line)

    _palette("TRANSITIONS", "connectives", CONNECTIVES)
    _palette("HEDGES", "hedges", HEDGES)
    _palette("INTENSIFIERS", "intensifiers", INTENSIFIERS, n=6)

    p = sig.get("punctuation") or {}
    if p.get("semicolon", 0) > 10 or p.get("em_dash", 0) > 10:
        habits = [n for n, k in (("semicolons", "semicolon"), ("em-dashes", "em_dash"))
                  if p.get(k, 0) > 10]
        parts.append(f"\nPUNCTUATION: he uses {' and '.join(habits)} freely.")

    head = "\n".join(parts)
    out = head

    # Exemplars come first and are never cut in half. They are the point: three sentences of
    # the real thing outweigh any amount of description, and it was precisely the exemplars
    # that the old 2,000-character cap threw away.
    if exemplars:
        # Reserve room for the analysis rather than letting a third exemplar eat it: two
        # passages already show the voice, and what the analysis says — how he opens a
        # paragraph, when he hedges, where the number goes — is exactly what counting cannot.
        reserve = min(len(analysis or ""), 1300)
        room = max(budget - len(head) - reserve, 700)
        body = ["\nHIS PROSE — this is the voice to match:"]
        used = 0
        for ex in exemplars:
            ex = " ".join(ex.split())
            if used + len(ex) > room and used:
                break
            body.append(f"  > {ex}")
            used += len(ex)
        out += "\n" + "\n".join(body) + "\n"

    # The analysis is the first thing to cut, not the last — and it is cut at a paragraph,
    # never mid-word.
    if analysis and (room := budget - len(out)) > 500:
        kept: list[str] = []
        for para in analysis.strip().split("\n\n"):
            para = para.strip()
            if sum(len(p) for p in kept) + len(para) > room:
                break
            kept.append(para)
        if kept:
            out += "\nHOW HE WRITES:\n" + "\n\n".join(kept) + "\n"
    return out


# ── the corpora on disk ──────────────────────────────────────────────────────

def _markup_files(root: Path) -> list[Path]:
    return sorted(f for f in Path(root).rglob("*.docx")
                  if not f.name.startswith(("~", "."))
                  and not any(part in _SKIP_DIRS for part in f.parts))


def _dedupe(texts: list[str]) -> list[str]:
    """The same span survives into every later cycle of the same document."""
    seen, out = set(), []
    for t in texts:
        k = " ".join(_words(t))
        if k and k not in seen:
            seen.add(k)
            out.append(t)
    return out


def harvest_authored(root: Path, tool_authors=TOOL_AUTHORS) -> list[str]:
    """Every sentence the AUTHOR typed by hand into a redline.

    Their voice, in domain, about this paper — a better style sample than a paragraph lifted
    from a paper they wrote in 2019 about something else, and it costs nothing: the pipeline
    has been recording it as tracked changes all along.

    Coalesced, not run by run: Word splits one typed sentence across several ``w:ins``
    elements, and harvesting those raw yields "A fresh presentation of an old idea can" — a
    fragment, which teaches a model nothing about how the author finishes a thought.
    """
    from haarpi.redline import authored_atoms

    out: list[str] = []
    for f in _markup_files(root):
        try:
            doc = Document(str(f))
        except Exception:  # noqa: BLE001 — a corrupt docx must not stop the harvest
            continue
        for p in doc.paragraphs:
            for text in authored_atoms(p._p, tool_authors).values():
                if len(text.split()) >= 8:
                    out.append(text.strip())
    return _dedupe(out)


def harvest_struck(root: Path, tool_authors=TOOL_AUTHORS) -> list[str]:
    """Every passage the AUTHOR deleted from the tool's prose.

    NOT a ban list. A deletion says the author did not want THAT TEXT — which may be because
    it was badly written, or because the claim was wrong, and the tracked change cannot tell
    us which. ("Static distance metrics in music theory fail to capture…" was struck for
    being false, not for being ugly.) It is a signal for a human to read, and the material
    for the one judgement worth spending a model on.
    """
    tool = {a.lower() for a in tool_authors}
    out: list[str] = []
    for f in _markup_files(root):
        try:
            doc = Document(str(f))
        except Exception:  # noqa: BLE001
            continue
        for el in doc.element.body.iter(qn("w:del")):
            if (el.get(qn("w:author")) or "").lower() in tool:
                continue                   # the tool deleting its own prose proves nothing
            text = "".join(t.text or "" for t in el.iter(qn("w:delText"))).strip()
            if len(text.split()) >= 5:
                out.append(text)
    return _dedupe(out)
