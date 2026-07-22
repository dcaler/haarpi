"""The shared redline engine: in-place, comment-preserving revision with tracked changes.

Deterministic OOXML machinery only — GPU-free and unit-testable. The LLM call that
turns a reviewer comment into revised paragraph text lives in each tool's revise
layer, and each tool keeps its own comment-routing (`comment_anchors` and friends)
in its redline binding; what lives here is the paragraph model and the XML surgery
rabbitHole and raconteur previously maintained as a fork pair.

A paragraph is modelled as an ordered stream of TEXT and OPAQUE atoms (equations,
footnote references, drawings), NOT as the text inside its ``w:r/w:t`` runs: an
equation is a SIBLING of the text runs, so a differ built on ``w:t`` alone sees
prose with holes where every number had been. Atoms serialize to sentinels
(``⟦m:1⟧``) for the differ and for the LLM, and expand back to their original
elements on write.

    INVARIANT: the tool never authors an equation; it only edits the prose
    around one. An atom is re-laid as ACCEPTED content between the redlined
    prose, never inside a ``w:ins``/``w:del``.

OOXML notes:
  * A comment is anchored by ``<w:commentRangeStart w:id=N/>`` …
    ``<w:commentRangeEnd w:id=N/>`` markers bracketing a run range, plus a
    ``<w:commentReference w:id=N/>`` run; the text lives in comments.xml.
    python-docx preserves all of these across an open/save.
  * A tracked deletion wraps the old run(s) in ``<w:del>`` and turns ``<w:t>``
    into ``<w:delText>``; a tracked insertion wraps new run(s) in ``<w:ins>``.
"""

from __future__ import annotations

import copy
import datetime
import difflib
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .text import sentence_units

# OOXML math. An equation is a sibling of the text runs, NOT inside one — see module docstring.
_MATH = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# python-docx's nsmap does not register the reserved ``xml`` prefix, so qn("xml:space")
# would KeyError — use the literal namespaced attribute name.
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

# Everyone else is a human. The tools sign their tracked changes and their comments, and
# that signature is the only thing separating what the machine wrote from what the author
# did. It decides which spans are editable and which are not.
TOOL_AUTHORS = ("rabbitHole", "raconteur", "raster", "rayleigh", "haarpi", "the tool")


def is_tool_author(author: str | None, tool_authors=TOOL_AUTHORS) -> bool:
    return (author or "").lower() in {a.lower() for a in tool_authors}


def _is_authored_span(el, tool_authors=TOOL_AUTHORS) -> bool:
    """A live tracked insertion by a human — text the author put there by hand.

    The most expensive text in the document. The tool preserves it and never authors it,
    which is exactly the contract it already honours for an equation.
    """
    return (el.tag == qn("w:ins")
            and not is_tool_author(el.get(qn("w:author")), tool_authors))


def _now() -> str:
    return datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")


# ── id allocation ─────────────────────────────────────────────────────────────

class _Ids:
    """Hand out w:id values that don't collide with existing comment/change ids."""

    def __init__(self, start: int):
        self._n = start

    def next(self) -> int:
        self._n += 1
        return self._n


def _max_existing_id(doc) -> int:
    """Highest w:id already used by comments or tracked changes, body + comments part."""
    ids = [0]

    def _scan(root):
        for tag in ("w:comment", "w:commentRangeStart", "w:commentRangeEnd",
                    "w:commentReference", "w:ins", "w:del"):
            for el in root.iter(qn(tag)):
                v = el.get(qn("w:id"))
                if v and v.lstrip("-").isdigit():
                    ids.append(int(v))

    _scan(doc.element.body)
    for rel in doc.part.rels.values():
        if rel.reltype.lower().endswith("/comments"):
            _scan(rel.target_part._element)
            break
    return max(ids)


def ids_for(doc) -> _Ids:
    return _Ids(_max_existing_id(doc))


# ── element builders ──────────────────────────────────────────────────────────

def _rpr_clone(run_el):
    rpr = run_el.find(qn("w:rPr"))
    return copy.deepcopy(rpr) if rpr is not None else None


def _dominant_rpr(els):
    """Formatting of the text run carrying the most characters.

    A paragraph often opens with a short specially-formatted run (a bold
    lead-in label, an italic term); the first run must not donate its
    formatting to a whole replacement. The run holding the bulk of the prose
    is what the replacement should look like — even when that run has no rPr
    at all (plain beats bold)."""
    best, best_len = None, -1
    for r in els:
        if r.tag == qn("w:r") and _is_text_run(r):
            n = sum(len(t.text or "") for t in r.findall(qn("w:t")))
            if n > best_len:
                best, best_len = _rpr_clone(r), n
    return best


def _text_run(text: str, rpr=None):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(_XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    return r


def _ins(text: str, author: str, wid: int, rpr=None):
    el = OxmlElement("w:ins")
    el.set(qn("w:id"), str(wid))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), _now())
    el.append(_text_run(text, rpr))
    return el


def _del(text: str, author: str, wid: int, rpr=None):
    el = OxmlElement("w:del")
    el.set(qn("w:id"), str(wid))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), _now())
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    dt = OxmlElement("w:delText")
    dt.set(_XML_SPACE, "preserve")
    dt.text = text
    r.append(dt)
    el.append(r)
    return el


# ── the paragraph as an atom stream ───────────────────────────────────────────

_SENTINEL_SPLIT = re.compile(r"(⟦[a-z]+:\d+⟧)")


_OPAQUE_RUN_CHILDREN = ("w:footnoteReference", "w:endnoteReference",
                        "w:drawing", "w:pict", "w:object")


def _is_text_run(r) -> bool:
    """A run carrying visible text (not a comment-reference marker run)."""
    return r.find(qn("w:t")) is not None and r.find(qn("w:commentReference")) is None


def _is_ref_run(r) -> bool:
    return r.tag == qn("w:r") and r.find(qn("w:commentReference")) is not None


def _is_opaque(el) -> bool:
    """Content we preserve but never author."""
    if el.tag in (f"{{{_MATH}}}oMath", f"{{{_MATH}}}oMathPara"):
        return True
    if el.tag == qn("w:hyperlink"):
        return True
    if el.tag == qn("w:r"):
        return any(el.find(qn(t)) is not None for t in _OPAQUE_RUN_CHILDREN)
    return False


def _sentinel_kind(el) -> str:
    if el.tag.startswith(f"{{{_MATH}}}"):
        return "m"
    if el.tag == qn("w:hyperlink"):
        return "h"
    return "x"


def serialize_paragraph(p_el, protect_authored: bool = False,
                        tool_authors=TOOL_AUTHORS) -> tuple[str, dict[str, object], list]:
    """Render a paragraph as (text_with_sentinels, sentinel -> element(s), consumed children).

    Comment plumbing and ``w:pPr`` are left alone — they are re-attached around the rebuilt
    body. Prior tracked deletions are dropped (the paragraph reads as it currently stands);
    prior insertions read as accepted text.

    ``protect_authored`` makes a HUMAN's live tracked insertion an atom instead of anonymous
    prose. This is the whole of span-level deference. Without it, provenance is erased right
    here — by the time the differ and the model see the paragraph, the author's sentences are
    indistinguishable from the tool's, and the only way to defend them is to freeze the entire
    paragraph, which also freezes the tool's own prose sitting beside them. As an atom, the
    author's span is content the tool preserves but never authors — the identical contract an
    equation already has, guards and all — and the re-cut writes AROUND it.

    Word splits one typed sentence across several ``w:ins`` elements, so ADJACENT authored
    insertions coalesce into a single atom: what the author wrote is one span, not five.
    """
    parts: list[str] = []
    smap: dict[str, object] = {}
    consumed: list = []
    n = 0
    pending: list = []          # a run of adjacent authored insertions

    def _flush_authored() -> None:
        nonlocal n
        if not pending:
            return
        n += 1
        key = f"⟦a:{n}⟧"
        smap[key] = list(pending)
        parts.append(key)
        consumed.extend(pending)
        pending.clear()

    for child in list(p_el):
        tag = child.tag
        if tag == qn("w:pPr") or tag in (qn("w:commentRangeStart"), qn("w:commentRangeEnd")):
            continue
        if _is_ref_run(child) or tag == qn("w:del"):
            continue
        if protect_authored and _is_authored_span(child, tool_authors):
            pending.append(child)
            continue
        _flush_authored()
        if _is_opaque(child):
            n += 1
            key = f"⟦{_sentinel_kind(child)}:{n}⟧"
            smap[key] = child
            parts.append(key)
            consumed.append(child)
        elif tag == qn("w:ins"):
            parts.append("".join(t.text or "" for t in child.iter(qn("w:t"))))
            consumed.append(child)
        elif tag == qn("w:r") and child.findall(qn("w:t")):
            parts.append("".join(t.text or "" for t in child.findall(qn("w:t"))))
            consumed.append(child)
    _flush_authored()
    return "".join(parts), smap, consumed


def authored_atoms(p_el, tool_authors=TOOL_AUTHORS) -> dict[str, str]:
    """``{sentinel: the author's exact words}`` for every protected span in a paragraph.

    The model must READ the author's text to write around it coherently — to pick up its
    thread, and not to repeat it — while being unable to change a character of it. So the
    sentinel travels with a legend, and the guards enforce that every one comes back.
    """
    _, smap, _ = serialize_paragraph(p_el, protect_authored=True, tool_authors=tool_authors)
    return {k: atom_text(v) for k, v in smap.items() if k.startswith("⟦a:")}


def paragraph_text(p_el, protect_authored: bool = False) -> str:
    """The paragraph as the reviser and the differ see it: prose with sentinels for atoms."""
    return serialize_paragraph(p_el, protect_authored)[0]


def _atom_els(value) -> list:
    """An atom is one element, or several — an authored span Word split across runs."""
    return list(value) if isinstance(value, list) else [value]


def atom_text(value) -> str:
    """The visible text inside an atom (an equation, a figure, an authored span).

    An equation's characters live in ``m:t``, not ``w:t``, so anything reading only ``w:t``
    sees a paragraph with holes where every number was. Used to flatten a paragraph for
    callers that want plain prose rather than sentinels.

    One ordered pass over both tags: no atom carries m:t and w:t together today, but an
    atom that ever did — a hyperlink inside an equation — would reorder the prose under a
    tag-at-a-time read, and the damage would look like a model error rather than a reader bug.
    """
    return "".join(t.text or ""
                   for el in _atom_els(value)
                   for t in el.iter(f"{{{_MATH}}}t", qn("w:t")))


def flatten_paragraph(p_el) -> str:
    """The paragraph as plain prose, with each atom rendered as its own text.

    Deletions are dropped and insertions read as accepted, exactly as ``serialize_paragraph``
    defines it — but atoms come back as their characters instead of ``⟦m:1⟧``, so the result
    is readable markdown rather than a sentinel stream.
    """
    text, smap, _ = serialize_paragraph(p_el)
    for key, el in smap.items():
        text = text.replace(key, atom_text(el))
    return text


def _render(text: str, smap: dict, rpr) -> list:
    """Text with sentinels -> runs, each sentinel expanded to its original element.

    An unknown sentinel (one the model invented) renders as nothing: the tool never authors
    an equation. The ``invented_sentinels`` guard rejects the rewrite before it reaches here,
    so this is a backstop, not a path.
    """
    out: list = []
    for piece in _SENTINEL_SPLIT.split(text):
        if not piece:
            continue
        if _SENTINEL_SPLIT.fullmatch(piece):
            el = smap.get(piece)
            if el is not None:
                out.extend(copy.deepcopy(e) for e in _atom_els(el))
        else:
            out.append(_text_run(piece, rpr))
    return out


def _segments(chunk: str) -> tuple[list[str], list[str]]:
    """Split on sentinels: (text segments, sentinels). ``len(texts) == len(sents) + 1``."""
    parts = _SENTINEL_SPLIT.split(chunk)
    return parts[0::2], parts[1::2]


def _redline_chunk(old_chunk: str, new_chunk: str, smap: dict,
                   author: str, ids: _Ids, rpr) -> list:
    """Redline one changed span, never touching an atom.

    An equation is re-laid as accepted content between the redlined prose around it —
    the tool cannot author an equation, so it must not claim to have deleted or inserted
    one. When the sentinel sequence is unchanged (the guarded case) the prose segments around
    each atom are redlined individually, so even a rewritten sentence keeps its numbers
    exactly where they were.
    """
    o_texts, o_sents = _segments(old_chunk)
    n_texts, n_sents = _segments(new_chunk)
    body: list = []

    if o_sents != n_sents:
        # The reviser moved, dropped, or invented an atom. The guard rejects this upstream;
        # here we simply never lose one: redline the prose, then re-lay every original atom.
        o_all, n_all = "".join(o_texts), "".join(n_texts)
        if o_all:
            body.append(_del(o_all, author, ids.next(), rpr))
        if n_all:
            body.append(_ins(n_all, author, ids.next(), rpr))
        for s in o_sents:
            if s in smap:
                body.extend(copy.deepcopy(e) for e in _atom_els(smap[s]))
        return body

    for k in range(len(o_sents) + 1):
        o = o_texts[k] if k < len(o_texts) else ""
        n = n_texts[k] if k < len(n_texts) else ""
        if o and o == n:
            body.append(_text_run(o, rpr))      # unchanged prose around an atom
        else:
            if o:
                body.append(_del(o, author, ids.next(), rpr))
            if n:
                body.append(_ins(n, author, ids.next(), rpr))
        if k < len(o_sents):
            el = smap.get(o_sents[k])
            if el is not None:
                # the atom itself: accepted, in place, never inside a w:ins/w:del. For an
                # authored span this is what keeps the reviewer's tracked insertion intact,
                # still pending, still theirs to accept.
                body.extend(copy.deepcopy(e) for e in _atom_els(el))
    return body


def _relay(p_el, body: list, consumed: list) -> None:
    """Detach what we consumed and re-lay [starts] <body> [ends] [reference runs].

    Whatever we did NOT consume keeps its place. A tracked deletion from an earlier cycle
    is the case that matters: it is not part of the text being revised (deleted text reads
    as gone), but it is still in the paragraph, and a rebuild that re-laid the new body at
    the top swept every old deletion to the paragraph's tail — severing the struck-through
    text from the prose it was struck from, which is most of what makes a redline readable.
    """
    starts = p_el.findall(qn("w:commentRangeStart"))
    ends = p_el.findall(qn("w:commentRangeEnd"))
    ref_runs = [r for r in p_el.findall(qn("w:r")) if _is_ref_run(r)]
    ppr = p_el.find(qn("w:pPr"))

    moved = {id(e) for e in (*starts, *ends, *ref_runs)}
    eaten = {id(e) for e in consumed}
    if ppr is not None:
        moved.add(id(ppr))

    middle: list = []
    placed = False
    for child in list(p_el):
        if id(child) in moved:
            continue
        if id(child) in eaten:
            if not placed:          # the body lands where its text came from
                middle.extend(body)
                placed = True
            continue
        middle.append(child)        # a prior cycle's w:del, a bookmark: keep it in place
    if not placed:
        middle.extend(body)

    for child in list(p_el):
        if child is not ppr:
            p_el.remove(child)
    insert_at = 1 if ppr is not None else 0
    for offset, el in enumerate([*starts, *middle, *ends, *ref_runs]):
        p_el.insert(insert_at + offset, el)


# ── tracked edits ─────────────────────────────────────────────────────────────

def tracked_replace(p_el, new_text: str, author: str, ids: _Ids | None = None,
                    protect_authored: bool = False) -> bool:
    """Replace a paragraph's text with one tracked deletion of the old + insertion of the new,
    preserving comment anchors and every opaque atom.

    Coarse: the whole paragraph is redlined. ``tracked_replace_sentencewise`` is what the
    redline path uses; this remains for callers that genuinely mean to replace wholesale.

    ``protect_authored`` holds the reviewer's own tracked insertions as atoms — the tool
    rewrites around them and cannot touch them.

    Returns False (a no-op) when there is no text to replace or the text is unchanged.
    """
    ids = ids or _Ids(1000)
    old_text, smap, consumed = serialize_paragraph(p_el, protect_authored)
    if not consumed or new_text.strip() == old_text.strip():
        return False
    rpr = _dominant_rpr(consumed)
    _relay(p_el, _redline_chunk(old_text, new_text, smap, author, ids, rpr), consumed)
    return True


def tracked_replace_sentencewise(p_el, new_text: str, author: str,
                                 ids: _Ids | None = None,
                                 protect_authored: bool = False) -> bool:
    """Replace a paragraph's text with SENTENCE-level tracked changes.

    Diffs old against new at sentence granularity and redlines only the sentences that
    actually changed; every unchanged sentence is re-laid as a plain (accepted) run,
    byte-for-byte, so its [@citekey] tags, its grounding, and its equations survive the
    revision untouched. Opaque atoms are never deleted or inserted — see ``_redline_chunk``.

    ``protect_authored`` holds the reviewer's own tracked insertions as atoms: a sentence
    the author wrote by hand cannot be edited, only written around. The tool's prose beside
    it stays fair game — which is the whole difference between deferring to the author and
    freezing the paragraph they happened to touch.

    Returns False (a no-op) when there is no text to replace or the text is unchanged.
    """
    ids = ids or _Ids(1000)
    old_text, smap, consumed = serialize_paragraph(p_el, protect_authored)
    if not consumed or new_text.strip() == old_text.strip():
        return False
    rpr = _dominant_rpr(consumed)

    old_units = sentence_units(old_text)
    new_units = sentence_units(new_text)
    sm = difflib.SequenceMatcher(a=old_units, b=new_units, autojunk=False)
    body: list = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for u in old_units[i1:i2]:
                body.extend(_render(u, smap, rpr))  # unchanged — accepted, atoms in place
        else:
            body.extend(_redline_chunk("".join(old_units[i1:i2]),
                                       "".join(new_units[j1:j2]),
                                       smap, author, ids, rpr))
    _relay(p_el, body, consumed)
    return True


def tracked_insert_after(p_el, text: str, author: str, ids: _Ids | None = None):
    """Insert a brand-new paragraph (wholly a tracked insertion) after ``p_el``, cloning its
    paragraph properties. For structural comments that ask to split a paragraph or add
    material."""
    ids = ids or _Ids(1000)
    new_p = OxmlElement("w:p")
    ppr = p_el.find(qn("w:pPr"))
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    rpr = _dominant_rpr(p_el.findall(qn("w:r")))
    new_p.append(_ins(text, author, ids.next(), rpr))
    p_el.addnext(new_p)
    return new_p


def tracked_heading_before(p_el, text: str, author: str, ids: _Ids | None = None,
                           style: str = "Heading2"):
    """Insert a heading paragraph (wholly a tracked insertion) before ``p_el``.

    For structural migrations — e.g. promoting a paragraph's inline lead-in label
    to a real section heading. The style id must exist in the document's styles
    part for Word to render it as a heading."""
    ids = ids or _Ids(1000)
    new_p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), style)
    ppr.append(pstyle)
    new_p.append(ppr)
    new_p.append(_ins(text, author, ids.next()))
    p_el.addprevious(new_p)
    return new_p


# ── comment reading / anchoring ───────────────────────────────────────────────

def comments_by_id(path: Path) -> dict[str, dict]:
    """Map comment id -> {author, text} from the comments part."""
    doc = Document(str(path))
    out: dict[str, dict] = {}
    for rel in doc.part.rels.values():
        if not rel.reltype.lower().endswith("/comments"):
            continue
        for c in rel.target_part._element.findall(".//" + qn("w:comment")):
            cid = c.get(qn("w:id"))
            texts = [t.text for t in c.findall(".//" + qn("w:t")) if t.text]
            out[cid] = {"author": c.get(qn("w:author"), "reviewer"),
                        "text": " ".join(texts)}
        break
    return out


def comment_spans(p_el, protect_authored: bool = False,
                  tool_authors=TOOL_AUTHORS) -> dict[str, tuple[int, int]]:
    """Character offsets ``[start, end)`` of each comment's anchored range, measured over the
    same serialized text ``paragraph_text`` returns.

    The reviewer highlights the phrase their comment is about, so this recovers WHICH
    SENTENCES a comment actually bears on. This is what lets the minimal-edit guard know that
    a comment on sentence 2 does not license rewriting sentences 1 and 3-7.

    ``protect_authored`` MUST match what the caller passed to ``serialize_paragraph``: these
    offsets are measured over that string, and an authored span is six characters as a
    sentinel and two hundred as prose. Disagree, and every comment after the first authored
    span anchors to the wrong sentence.
    """
    offset = 0
    n = 0
    opens: dict[str, int] = {}
    spans: dict[str, tuple[int, int]] = {}
    pending = False              # mid-run of adjacent authored insertions

    def _open(cid):
        if cid:
            opens[cid] = offset

    def _close(cid):
        if cid in opens:
            spans[cid] = (opens.pop(cid), offset)

    def _flush_authored():
        """Mirror serialize_paragraph: one sentinel for a coalesced authored span."""
        nonlocal n, offset, pending
        if pending:
            n += 1
            offset += len(f"⟦a:{n}⟧")
            pending = False

    for child in list(p_el):
        tag = child.tag
        if tag == qn("w:commentRangeStart"):
            _open(child.get(qn("w:id")))
        elif tag == qn("w:commentRangeEnd"):
            _close(child.get(qn("w:id")))
        elif tag == qn("w:pPr") or _is_ref_run(child):
            continue
        elif protect_authored and _is_authored_span(child, tool_authors):
            # The author's span is a sentinel in the serialized text, so it advances the
            # offset by the sentinel's width, not by its prose. A comment anchored inside
            # it (the author commenting on their own sentence) collapses onto the sentinel.
            for el in child.iter(qn("w:commentRangeStart"), qn("w:commentRangeEnd")):
                (_open if el.tag == qn("w:commentRangeStart") else _close)(el.get(qn("w:id")))
            pending = True
        elif tag == qn("w:del"):
            # Deleted text is not in the serialized paragraph, so it advances no offset —
            # and it does NOT interrupt a run of authored insertions, exactly as
            # serialize_paragraph does not. (A reviewer typically deletes a sentence and
            # types its replacement around the corpse; those insertions are ONE span.)
            # An anchor nested INSIDE the deletion must still be seen: a search of direct
            # children alone loses the comment entirely. It collapses to a point — the
            # place where the text they were talking about used to be.
            for el in child.iter(qn("w:commentRangeStart"), qn("w:commentRangeEnd")):
                (_open if el.tag == qn("w:commentRangeStart") else _close)(el.get(qn("w:id")))
        elif _is_opaque(child):
            _flush_authored()
            # Must mirror serialize_paragraph's numbering exactly. Hardcoding the width as
            # len("⟦m:0⟧") drifts one char per atom from the tenth atom onward, which
            # silently mis-anchors every comment after it in a paragraph with many atoms —
            # e.g. a Results paragraph full of inline statistics.
            n += 1
            offset += len(f"⟦{_sentinel_kind(child)}:{n}⟧")
        elif tag == qn("w:ins"):
            _flush_authored()
            # Insertions read as accepted text; anchors nested inside one are real anchors
            # (a reviewer commenting on their own inserted sentence).
            for el in child.iter():
                if el.tag == qn("w:commentRangeStart"):
                    _open(el.get(qn("w:id")))
                elif el.tag == qn("w:commentRangeEnd"):
                    _close(el.get(qn("w:id")))
                elif el.tag == qn("w:t"):
                    offset += len(el.text or "")
        elif tag == qn("w:r"):
            _flush_authored()
            offset += sum(len(t.text or "") for t in child.findall(qn("w:t")))
    _flush_authored()
    for cid, start in opens.items():  # range never closed in this paragraph
        spans[cid] = (start, offset)
    return spans


def anchors_in_deleted_text(p_el) -> set[str]:
    """Comment ids whose anchor lies wholly inside deleted text.

    The reviewer wrote "what IS tonal negotiation?" and then deleted the sentence that
    said it. The comment is still open, but its subject is gone — and a tool that hands
    such a comment to a model as a live instruction gets a confident rewrite of prose that
    no longer exists. There is nothing to revise here; there is something to SAY.
    """
    inside: set[str] = set()
    for d in p_el.findall(qn("w:del")):
        for el in d.iter(qn("w:commentRangeStart")):
            inside.add(el.get(qn("w:id")))
    live: set[str] = set()
    for el in p_el.iter(qn("w:commentRangeStart")):
        parent = el.getparent()
        if parent is not None and parent.tag != qn("w:del"):
            live.add(el.get(qn("w:id")))
    return {cid for cid in inside if cid and cid not in live}


def anchored_sentences(text: str, span: tuple[int, int]) -> set[int]:
    """Indices of the sentences a comment's character range overlaps.

    A zero-width span is not a comment about nothing: it is a comment whose text was
    deleted (see ``comment_spans``). It bears on the sentence it now sits in.
    """
    start, end = span
    if end <= start:
        end = start + 1
    out: set[int] = set()
    pos = 0
    for i, unit in enumerate(sentence_units(text)):
        if pos < end and start < pos + len(unit):
            out.add(i)
        pos += len(unit)
    return out


def is_heading_style(style_name: str) -> bool:
    """True for Word heading/title styles (so we never rewrite a heading as prose)."""
    s = (style_name or "").lower()
    return s.startswith("heading") or s == "title"


def _accepted_para_text(p_el) -> str:
    """One paragraph as it reads with every tracked change accepted.

    Insertions kept, deletions dropped — ``w:t`` lives in normal and ``<w:ins>`` runs, while
    deleted text sits in ``<w:delText>`` and is therefore skipped.
    """
    out: list[str] = []
    for r in p_el.iter(qn("w:r")):
        parent = r.getparent()
        if parent is not None and parent.tag == qn("w:del"):
            continue
        for t in r.iter(qn("w:t")):
            out.append(t.text or "")
    return "".join(out)


# ── the gate mechanics: unresolved detection, accepting, minting ──────────────
# `haarpi next` runs these when a markup task is marked done in trundlr: a clean
# markup (nothing unresolved) mints a RELEASE — tracked changes accepted,
# comment threads stripped, bare-chain filename (see haarpi.naming). The check
# is deliberately mechanical: the gate needs no LLM.

import zipfile

from lxml import etree

_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"

TOOL_AUTHORS = ("rabbitHole", "raconteur", "raster", "rayleigh", "haarpi", "the tool")


def _zip_parts(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as z:
        return {n: z.read(n) for n in z.namelist()}


def comment_threads(path: Path, tool_authors=TOOL_AUTHORS) -> dict[str, dict]:
    """Every comment, with its thread position and its resolution state.

    One reading of the comments part, so that every caller agrees about what a
    comment IS. They did not, and the disagreement was expensive: the gate honoured
    ``w15:done`` while the reviser did not, so nine settled comments were re-fired as
    fresh instructions and the tool spent a cycle rewriting text the author had
    already blessed.

    Returns ``{id: {id, author, text, done, is_tool, parent, replies}}`` where
    ``parent`` is the comment id this one replies to (None for a top-level ask) and
    ``replies`` is the ordered list of ids replying to it.

    Resolution is recorded per PARAGRAPH id in commentsExtended.xml, and Word marks a
    thread done by marking its top comment — so a reply inherits its parent's state.
    """
    parts = _zip_parts(path)
    if "word/comments.xml" not in parts:
        return {}
    croot = etree.fromstring(parts["word/comments.xml"])

    done_paras: set[str] = set()
    parent_of_para: dict[str, str] = {}
    if "word/commentsExtended.xml" in parts:
        ce = etree.fromstring(parts["word/commentsExtended.xml"])
        for cex in ce.iter(f"{{{_W15}}}commentEx"):
            pid = cex.get(f"{{{_W15}}}paraId")
            if not pid:
                continue
            if cex.get(f"{{{_W15}}}done") == "1":
                done_paras.add(pid)
            parent = cex.get(f"{{{_W15}}}paraIdParent")
            if parent:
                parent_of_para[pid] = parent

    tool_lower = {a.lower() for a in tool_authors}
    owner_of_para: dict[str, str] = {}
    out: dict[str, dict] = {}
    for c in croot.findall(qn("w:comment")):
        cid = c.get(qn("w:id"))
        para_ids = [p for p in (p.get(f"{{{_W14}}}paraId") for p in c.findall(qn("w:p"))) if p]
        for p in para_ids:
            owner_of_para[p] = cid
        author = c.get(qn("w:author")) or ""
        out[cid] = {
            "id": cid,
            "author": author,
            "text": "".join(t.text or "" for t in c.iter(qn("w:t"))),
            "is_tool": author.lower() in tool_lower,
            "done": any(p in done_paras for p in para_ids),
            "_paras": para_ids,
            "parent": None,
            "replies": [],
        }

    for cid, rec in out.items():
        for p in rec["_paras"]:
            parent_para = parent_of_para.get(p)
            if parent_para and owner_of_para.get(parent_para) not in (None, cid):
                rec["parent"] = owner_of_para[parent_para]
                break
    for cid, rec in out.items():
        if rec["parent"] in out:
            out[rec["parent"]]["replies"].append(cid)
    # a reply inherits the thread's resolution: Word marks the top comment done
    for cid, rec in out.items():
        if rec["parent"] in out and out[rec["parent"]]["done"]:
            rec["done"] = True
    for rec in out.values():
        rec.pop("_paras", None)
    return out


def open_asks(path: Path, tool_authors=TOOL_AUTHORS) -> list[dict]:
    """The reviewer's LIVE instructions: top-level, human-authored, not resolved.

    This is the only thing a revision may act on. A resolved comment is history, not
    an instruction — re-firing it makes the tool rewrite text the reviewer already
    accepted, which is how a settled "this is ok" produced a full paragraph re-cut.

    Each ask carries its thread, because the thread is part of the ask:

      ``followups`` — the reviewer's own replies. Under the comment protocol, new
        information about an unmet ask arrives as a reply rather than a second
        comment, so a tool that ignores replies ignores half of what it was told.

      ``repeat`` — True when the tool has ALREADY replied in this thread and the ask
        is still open. That is not a new task; it is a task the tool already failed.
        The signal is derived from the document, so the reviewer never has to
        hand-write "this still needs a definition" to be heard a second time.
    """
    threads = comment_threads(path, tool_authors)
    out = []
    for rec in threads.values():
        if rec["is_tool"] or rec["parent"] or rec["done"]:
            continue
        replies = [threads[r] for r in rec["replies"] if r in threads]
        out.append({
            "id": rec["id"],
            "author": rec["author"],
            "text": rec["text"],
            "followups": [r["text"] for r in replies if not r["is_tool"]],
            "prior_tool_replies": [r["text"] for r in replies if r["is_tool"]],
            "repeat": any(r["is_tool"] for r in replies),
        })
    return sorted(out, key=lambda r: int(r["id"]) if str(r["id"]).isdigit() else 0)


def unresolved_comments(path: Path, tool_authors=TOOL_AUTHORS) -> list[dict]:
    """Top-level comments not marked done — what blocks a gate. THE TOOL'S COUNT TOO.

    Replies are conversation, not asks: the ask lives in the parent, and resolving the
    thread marks the parent done.

    Tool comments used to be exempt, on the reasoning that the gate exists to adjudicate the
    tool's work and the tool cannot instruct itself. That holds for ``open_asks``, which is
    what a revision ACTS on and still excludes them. It does not hold for the gate: the
    skeleton's word plan rides on a tool comment, and a plan nobody read is a plan nobody
    approved. Resolving it is the acknowledgement — and resolved is not deleted, so the
    comment carries that state into the release.
    """
    threads = comment_threads(path, tool_authors)
    out = [{"id": r["id"], "author": r["author"], "text": r["text"]}
           for r in threads.values() if not r["parent"] and not r["done"]]
    return sorted(out, key=lambda r: int(r["id"]) if str(r["id"]).isdigit() else 0)


def pending_reviewer_changes(path: Path, tool_authors=TOOL_AUTHORS) -> int:
    """Tracked changes authored by someone other than the tool — a reviewer
    insertion/deletion is a directive the tool hasn't processed yet."""
    doc = Document(str(path))
    tool_lower = {a.lower() for a in tool_authors}
    n = 0
    for tag in ("w:ins", "w:del"):
        for el in doc.element.body.iter(qn(tag)):
            if (el.get(qn("w:author")) or "").lower() not in tool_lower:
                n += 1
    return n


def gate_check(path: Path, tool_authors=TOOL_AUTHORS) -> dict:
    """The mechanical gate: clean ⟺ every reviewer comment is resolved.

    Comments are the work order. In the default redline path the tool edits only in
    answer to a comment, so resolving every comment adjudicates every tool change —
    the comment gate already governs the tool's edits. The reviewer's OWN tracked
    edits are their final word on that span, not an instruction the tool must process;
    they are reported for context but do not gate, and the mint accepts them into the
    release along with everything else. (Was: also required reviewer_changes == 0 —
    the "accept every tracked change to advance" rule, now retired.)"""
    unresolved = unresolved_comments(path, tool_authors)
    reviewer_changes = pending_reviewer_changes(path, tool_authors)
    return {"unresolved": unresolved, "reviewer_changes": reviewer_changes,
            "clean": not unresolved}


def has_foreign_markup(p_el, tool_authors=TOOL_AUTHORS) -> bool:
    """True when the paragraph carries a live tracked change by a human.

    Human markup freezes a paragraph: acceptance is a human act, and the tool
    never rewrites — or silently accepts — text a person tracked into the
    document. A frozen paragraph gets comments, never edits."""
    tool_lower = {a.lower() for a in tool_authors}
    for tag in ("w:ins", "w:del"):
        for el in p_el.iter(qn(tag)):
            if (el.get(qn("w:author")) or "").lower() not in tool_lower:
                return True
    return False


_CE_CTYPE = ("application/vnd.openxmlformats-officedocument.wordprocessingml"
             ".commentsExtended+xml")
_CE_RELTYPE = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"


def _fresh_para_id() -> str:
    import random
    return f"{random.getrandbits(31) | 0x10000000:08X}"


def add_replies(path: Path, replies: dict[str, str],
                author: str = "raconteur", initials: str = "ra") -> int:
    """Append a threaded reply to each comment id in `replies`, in place.

    Word models a reply as a full <w:comment> whose commentsExtended entry
    carries w15:paraIdParent = the parent's paragraph id, with range markers
    shadowing the parent's anchor. `unresolved_comments` already treats
    paraIdParent entries as conversation, so a tool reply never blocks a gate —
    this is the writer that half of the machinery was waiting for.
    """
    if not replies:
        return 0
    parts = _zip_parts(path)
    if "word/comments.xml" not in parts:
        return 0
    croot = etree.fromstring(parts["word/comments.xml"])
    docroot = etree.fromstring(parts["word/document.xml"])

    comments = {c.get(qn("w:id")): c for c in croot.findall(qn("w:comment"))}
    next_id = max([int(i) for i in comments if i and i.isdigit()] + [0]) + 1

    ce_name = "word/commentsExtended.xml"
    created_ce = ce_name not in parts
    ceroot = (etree.fromstring(parts[ce_name]) if not created_ce
              else etree.fromstring(f'<w15:commentsEx xmlns:w15="{_W15}"/>'.encode()))
    known_ex = {ex.get(f"{{{_W15}}}paraId")
                for ex in ceroot.iter(f"{{{_W15}}}commentEx")}

    def _ensure_para_id(comment_el) -> str | None:
        ps = comment_el.findall(qn("w:p"))
        if not ps:
            return None
        pid = ps[-1].get(f"{{{_W14}}}paraId")
        if pid is None:
            pid = _fresh_para_id()
            ps[-1].set(f"{{{_W14}}}paraId", pid)
        return pid

    added = 0
    for parent_id, text in replies.items():
        parent = comments.get(str(parent_id))
        if parent is None or not (text or "").strip():
            continue
        pend = docroot.find(f".//{qn('w:commentRangeEnd')}[@{qn('w:id')}='{parent_id}']")
        pref = docroot.find(f".//{qn('w:commentReference')}[@{qn('w:id')}='{parent_id}']")
        if pend is None or pref is None:
            continue                        # no anchor in the body — nothing to thread onto

        parent_para = _ensure_para_id(parent)
        if parent_para and parent_para not in known_ex:
            ex = etree.SubElement(ceroot, f"{{{_W15}}}commentEx")
            ex.set(f"{{{_W15}}}paraId", parent_para)
            ex.set(f"{{{_W15}}}done", "0")
            known_ex.add(parent_para)

        nid = str(next_id)
        next_id += 1
        reply_para = _fresh_para_id()

        c = etree.SubElement(croot, qn("w:comment"))
        c.set(qn("w:id"), nid)
        c.set(qn("w:author"), author)
        c.set(qn("w:initials"), initials)
        c.set(qn("w:date"), _now())
        p = etree.SubElement(c, qn("w:p"))
        p.set(f"{{{_W14}}}paraId", reply_para)
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.set(_XML_SPACE, "preserve")
        t.text = text.strip()

        s_el = etree.Element(qn("w:commentRangeStart")); s_el.set(qn("w:id"), nid)
        e_el = etree.Element(qn("w:commentRangeEnd"));   e_el.set(qn("w:id"), nid)
        pend.addprevious(s_el)
        pend.addnext(e_el)
        ref_run = etree.Element(qn("w:r"))
        ref = etree.SubElement(ref_run, qn("w:commentReference"))
        ref.set(qn("w:id"), nid)
        pref.getparent().addnext(ref_run)

        ex = etree.SubElement(ceroot, f"{{{_W15}}}commentEx")
        ex.set(f"{{{_W15}}}paraId", reply_para)
        if parent_para:
            ex.set(f"{{{_W15}}}paraIdParent", parent_para)
        ex.set(f"{{{_W15}}}done", "0")
        added += 1

    if not added:
        return 0

    def _ser(root):
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                              standalone=True)
    parts["word/comments.xml"] = _ser(croot)
    parts["word/document.xml"] = _ser(docroot)
    parts[ce_name] = _ser(ceroot)
    if created_ce:
        ct = etree.fromstring(parts["[Content_Types].xml"])
        ns = "http://schemas.openxmlformats.org/package/2006/content-types"
        ov = etree.SubElement(ct, f"{{{ns}}}Override")
        ov.set("PartName", "/word/commentsExtended.xml")
        ov.set("ContentType", _CE_CTYPE)
        parts["[Content_Types].xml"] = _ser(ct)
        rels_name = "word/_rels/document.xml.rels"
        rr = etree.fromstring(parts[rels_name])
        rns = "http://schemas.openxmlformats.org/package/2006/relationships"
        nums = [int(rel.get("Id")[3:]) for rel in rr
                if (rel.get("Id") or "").startswith("rId") and rel.get("Id")[3:].isdigit()]
        rel = etree.SubElement(rr, f"{{{rns}}}Relationship")
        rel.set("Id", f"rId{max(nums, default=0) + 1}")
        rel.set("Type", _CE_RELTYPE)
        rel.set("Target", "commentsExtended.xml")
        parts[rels_name] = _ser(rr)

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in parts.items():
            z.writestr(n, b)
    return added


def _visible_runs(p_el) -> list:
    """Every run carrying visible text, in document order — including runs INSIDE a
    tracked insertion, which is where the author's own words live. Deleted text is not
    visible, and a comment-reference run is plumbing."""
    out = []
    for r in p_el.iter(qn("w:r")):
        parent = r.getparent()
        if parent is not None and parent.tag == qn("w:del"):
            continue
        if r.find(qn("w:commentReference")) is not None:
            continue
        if r.find(qn("w:t")) is not None:
            out.append(r)
    return out


def _split_run(run, text: str):
    """A copy of ``run`` carrying ``text`` — same formatting, same place in the redline."""
    clone = copy.deepcopy(run)
    t = clone.find(qn("w:t"))
    t.text = text
    t.set(_XML_SPACE, "preserve")
    return clone


def anchor_fragment(p_el, fragment: str, nid: str) -> bool:
    """Bracket EXACTLY ``fragment`` with comment-range markers, splitting runs to do it.

    A comment on a whole paragraph, whose text is "its", makes the reader hunt for the word
    it means. A comment ON the word says everything by where it sits.

    The markers may land INSIDE a ``w:ins`` — they must, since the author's sentences are
    tracked insertions and that is exactly the text we most often have something to say
    about. This is legal (comment-range markers are run-level elements, which a tracked
    insertion's content model admits) and it is how Word itself comments on inserted text.
    The author's characters are not touched: the run is split, and each piece keeps its
    formatting and its place in the tracked change.

    Returns False if the fragment is not in this paragraph — the caller falls back to
    anchoring the paragraph, which is worse but never wrong.
    """
    runs = _visible_runs(p_el)
    spans = [(r, r.find(qn("w:t")).text or "") for r in runs]
    full = "".join(t for _, t in spans)
    idx = full.find(fragment)
    if idx < 0:
        return False
    end = idx + len(fragment)

    first = last = None
    pos = 0
    for run, txt in spans:
        r_start, r_end = pos, pos + len(txt)
        pos = r_end
        if r_end <= idx or r_start >= end:
            continue                                   # this run is outside the fragment
        lo = max(idx, r_start) - r_start
        hi = min(end, r_end) - r_start
        parent = run.getparent()
        at = list(parent).index(run)
        pieces = []
        if txt[:lo]:
            pieces.append(_split_run(run, txt[:lo]))
        mid = _split_run(run, txt[lo:hi])
        pieces.append(mid)
        if txt[hi:]:
            pieces.append(_split_run(run, txt[hi:]))
        parent.remove(run)
        for k, el in enumerate(pieces):
            parent.insert(at + k, el)
        if first is None:
            first = mid
        last = mid

    if first is None:
        return False
    s_el = etree.Element(qn("w:commentRangeStart")); s_el.set(qn("w:id"), nid)
    e_el = etree.Element(qn("w:commentRangeEnd"));   e_el.set(qn("w:id"), nid)
    first.addprevious(s_el)
    last.addnext(e_el)
    return True


def _is_heading_p(p_el) -> bool:
    """A heading, read off the XML — the caller may not have a python-docx paragraph."""
    pPr = p_el.find(qn("w:pPr"))
    st = pPr.find(qn("w:pStyle")) if pPr is not None else None
    return bool(st is not None and (st.get(qn("w:val")) or "").lower().startswith("heading"))


def add_anchored_comments(path: Path, notes: list[tuple[str, str]],
                          author: str = "raconteur", initials: str = "ra",
                          headings_only: bool = False) -> int:
    """Write new top-level comments, each anchored to the text it is ABOUT.

    ``notes`` is a list of (where, comment_text). ``where`` is either:

    ``notes`` is a list of (fragment, comment_text). The fragment locates the paragraph, and
    the comment anchors on those characters alone where they can be found: on the word
    "it's", not the sentence containing it. A whole paragraph equal to the fragment wins over
    one merely containing it; failing both it falls back to the paragraph, which is worse but
    never wrong.

    ``headings_only`` restricts the search to heading paragraphs. A section's word plan is
    about the SECTION, and searching the whole document for "Results" put it on an Abstract
    bullet beginning "Results show that…" — the first paragraph containing the word. No
    reader could then find it: ``heading_comments`` returns only comments on headings, so
    the section had no rate and the draft would have banded it at a flat 150 without a word.

    This is the tool's voice on text it must not edit (the author's own): a disagreement, a
    correction, a proposal lands as a comment, never as a tracked change. Tool-authored
    comments never block a gate (see ``unresolved_comments``).

    Requires an existing word/comments.xml (in practice these notes only arise on
    documents the reviewer has already commented); returns 0 if there is none.
    """
    if not notes:
        return 0
    parts = _zip_parts(path)
    if "word/comments.xml" not in parts:
        return 0
    croot = etree.fromstring(parts["word/comments.xml"])
    docroot = etree.fromstring(parts["word/document.xml"])

    existing = [int(i) for c in croot.findall(qn("w:comment"))
                if (i := c.get(qn("w:id"))) and i.isdigit()]
    next_id = max(existing + [0]) + 1

    ce_name = "word/commentsExtended.xml"
    created_ce = ce_name not in parts
    ceroot = (etree.fromstring(parts[ce_name]) if not created_ce
              else etree.fromstring(f'<w15:commentsEx xmlns:w15="{_W15}"/>'.encode()))

    def _norm(s: str) -> str:
        return " ".join(s.split())

    body = docroot.find(qn("w:body"))
    paras = list(body.iter(qn("w:p"))) if body is not None else []

    added = 0
    for where, text in notes:
        if not (text or "").strip():
            continue
        fragment = where
        if not (fragment or "").strip():
            continue
        frag = _norm(fragment)
        cands = [p for p in paras if not headings_only or _is_heading_p(p)]
        texts = [_norm("".join(t.text or "" for t in p.iter(qn("w:t")))) for p in cands]
        # A WHOLE paragraph equal to the fragment wins over one merely containing it.
        # Anchoring a section's word plan by searching for "Results" put it on an Abstract
        # bullet beginning "Results show that…" — the first paragraph containing the word.
        # heading_comments then could not see it, so the section had no rate at all and the
        # draft would have banded it at a flat 150. Containment stays as the fallback: a
        # copyedit note anchors on a phrase INSIDE a paragraph and must keep doing so.
        target = next((p for p, t in zip(cands, texts) if t == frag), None)
        if target is None:
            target = next((p for p, t in zip(cands, texts) if frag in t), None)
        if target is None:
            continue

        nid = str(next_id)
        next_id += 1

        # On the words themselves where we can find them; on the paragraph where we cannot.
        if not anchor_fragment(target, fragment.strip(), nid):
            s_el = etree.Element(qn("w:commentRangeStart")); s_el.set(qn("w:id"), nid)
            e_el = etree.Element(qn("w:commentRangeEnd"));   e_el.set(qn("w:id"), nid)
            first_run = target.find(qn("w:r"))
            if first_run is not None:
                first_run.addprevious(s_el)
            else:
                target.insert(0, s_el)
            target.append(e_el)
        ref_run = etree.SubElement(target, qn("w:r"))
        ref = etree.SubElement(ref_run, qn("w:commentReference"))
        ref.set(qn("w:id"), nid)

        c = etree.SubElement(croot, qn("w:comment"))
        c.set(qn("w:id"), nid)
        c.set(qn("w:author"), author)
        c.set(qn("w:initials"), initials)
        c.set(qn("w:date"), _now())
        p = etree.SubElement(c, qn("w:p"))
        p.set(f"{{{_W14}}}paraId", _fresh_para_id())
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.set(_XML_SPACE, "preserve")
        t.text = text.strip()

        ex = etree.SubElement(ceroot, f"{{{_W15}}}commentEx")
        ex.set(f"{{{_W15}}}paraId", p.get(f"{{{_W14}}}paraId"))
        ex.set(f"{{{_W15}}}done", "0")
        added += 1

    if not added:
        return 0

    def _ser(root):
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                              standalone=True)
    parts["word/comments.xml"] = _ser(croot)
    parts["word/document.xml"] = _ser(docroot)
    parts[ce_name] = _ser(ceroot)
    if created_ce:
        ct = etree.fromstring(parts["[Content_Types].xml"])
        ns = "http://schemas.openxmlformats.org/package/2006/content-types"
        ov = etree.SubElement(ct, f"{{{ns}}}Override")
        ov.set("PartName", "/word/commentsExtended.xml")
        ov.set("ContentType", _CE_CTYPE)
        parts["[Content_Types].xml"] = _ser(ct)
        rels_name = "word/_rels/document.xml.rels"
        rr = etree.fromstring(parts[rels_name])
        rns = "http://schemas.openxmlformats.org/package/2006/relationships"
        nums = [int(rel.get("Id")[3:]) for rel in rr
                if (rel.get("Id") or "").startswith("rId") and rel.get("Id")[3:].isdigit()]
        rel = etree.SubElement(rr, f"{{{rns}}}Relationship")
        rel.set("Id", f"rId{max(nums, default=0) + 1}")
        rel.set("Type", _CE_RELTYPE)
        rel.set("Target", "commentsExtended.xml")
        parts[rels_name] = _ser(rr)

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in parts.items():
            z.writestr(n, b)
    return added


def _revisable_roots(doc):
    """Every part a revision can live in, not just the body.

    Footnotes, endnotes, headers, footers and the style definitions are separate parts, and
    ``doc.element.body`` reaches none of them. The css2026 manuscript minted "clean" with
    three insertions in footnotes.xml and a style-definition change in styles.xml, because
    both the accept and the check that verified it looked only at document.xml.
    """
    roots = [doc.element.body]
    seen = {id(doc.element.body)}
    for rel in doc.part.rels.values():
        try:
            el = rel.target_part._element
        except AttributeError:
            continue
        low = rel.reltype.lower()
        if not any(k in low for k in ("footnotes", "endnotes", "header", "footer",
                                      "styles", "comments")):
            continue
        if id(el) not in seen:
            seen.add(id(el))
            roots.append(el)
    return roots


def accept_all_changes(doc) -> dict:
    """Accept EVERY tracked change, not only the two that carry prose.

    Word records seven kinds of revision and this handled two. The css2026 manuscript
    released with 19 of them intact — moved text, paragraph and character formatting — so a
    document minted as clean opened in Word still showing tracked changes, and its own gate
    would have reported reviewer edits pending. The move markers were the harmful ones: a
    ``w:moveFrom`` left beside its ``w:moveTo`` is the same passage present twice.

    A deleted paragraph MARK (w:pPr/w:rPr/w:del) is how Word records "this paragraph ends
    here no longer" — it is what you get for deleting a heading, and it is not exotic. This
    used to strip the marker and keep the paragraph, on the reasoning that the tool never
    authors one; the AUTHOR does, every time they cut a section. The css2026 skeleton
    release carried eleven empty Heading 3 paragraphs from eleven headings the author had
    deleted, each still occupying a number in a numbered document.

    A mark-deleted paragraph whose text is also fully deleted is removed outright. One that
    still has text is left in place with the marker stripped: Word would merge it into its
    successor, and dropping surviving prose to imitate that would lose content, which is the
    worse error of the two.
    """
    counts = {"ins": 0, "del": 0, "paras": 0, "moved": 0, "format": 0, "cells": 0}
    for body in _revisable_roots(doc):
        _accept_in(body, counts)
    return counts


def _accept_in(body, counts: dict) -> None:
    """Accept every revision under one root element."""
    # FIRST, while the markers still exist: a paragraph-mark deletion lives in
    # w:pPr/w:rPr/w:del, and the sweep below removes every w:del in the body — including
    # those. Collecting after it would find nothing, which is what the first version of
    # this fix did.
    def _mark_deleted(p) -> bool:
        pPr = p.find(qn("w:pPr"))
        rPr = pPr.find(qn("w:rPr")) if pPr is not None else None
        return rPr is not None and rPr.find(qn("w:del")) is not None

    marked = [p for p in body.iter(qn("w:p")) if _mark_deleted(p)]
    for ins in list(body.iter(qn("w:ins"))):
        parent = ins.getparent()
        idx = list(parent).index(ins)
        for child in list(ins):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)
        counts["ins"] += 1
    for d in list(body.iter(qn("w:del"))):
        d.getparent().remove(d)
        counts["del"] += 1
    # MOVED TEXT is a deletion and an insertion under different names. Accepting means
    # dropping the origin and keeping the destination — leaving both, which is what handling
    # only w:ins/w:del did, puts the moved passage in the document TWICE.
    for mf in list(body.iter(qn("w:moveFrom"))):
        mf.getparent().remove(mf)
        counts["moved"] += 1
    for mt in list(body.iter(qn("w:moveTo"))):
        parent = mt.getparent()
        idx = list(parent).index(mt)
        for child in list(mt):
            parent.insert(idx, child)
            idx += 1
        parent.remove(mt)
        counts["moved"] += 1
    for tag in ("w:moveFromRangeStart", "w:moveFromRangeEnd",
                "w:moveToRangeStart", "w:moveToRangeEnd"):
        for el in list(body.iter(qn(tag))):
            el.getparent().remove(el)
    # A *PrChange records the FORMER formatting of a paragraph, run, section, table, row or
    # cell. Accepting the current formatting is simply dropping that record — there is
    # nothing to unwrap, because the live properties are already the accepted ones.
    for tag in ("w:pPrChange", "w:rPrChange", "w:sectPrChange", "w:tblPrChange",
                "w:tcPrChange", "w:trPrChange", "w:tblGridChange", "w:numberingChange"):
        for el in list(body.iter(qn(tag))):
            el.getparent().remove(el)
            counts["format"] += 1
    # Table cell insertions/deletions live in the row properties, not the body text.
    for tag in ("w:cellIns", "w:cellDel", "w:cellMerge"):
        for el in list(body.iter(qn(tag))):
            el.getparent().remove(el)
            counts["cells"] += 1
    for p in marked:
        parent = p.getparent()
        if parent is None:
            continue
        # Deletions are gone now, so what remains is what the author kept.
        if not "".join(t.text or "" for t in p.iter(qn("w:t"))).strip():
            parent.remove(p)
            counts["paras"] += 1
            continue
        # A deleted paragraph MARK with surviving text is a MERGE: the author removed the
        # break between two paragraphs. Stripping the marker and leaving the split was a
        # false choice between losing prose and leaving it broken — moving the runs into the
        # next paragraph loses nothing and is what Word does. Section 5.3 of the css2026
        # manuscript shipped as two paragraphs, the second beginning mid-sentence.
        sibs = list(parent)
        nxt = next((e for e in sibs[sibs.index(p) + 1:] if e.tag == qn("w:p")), None)
        if nxt is None:
            continue                     # nothing to merge into; leave it whole
        keep = [c for c in p if c.tag not in (qn("w:pPr"),)]
        anchor = nxt.find(qn("w:pPr"))
        at = list(nxt).index(anchor) + 1 if anchor is not None else 0
        for c in keep:
            p.remove(c)
            nxt.insert(at, c)
            at += 1
        parent.remove(p)
        counts["paras"] += 1
    for rpr in list(body.iter(qn("w:rPr"))):
        for tag in ("w:ins", "w:del"):
            m = rpr.find(qn(tag))
            if m is not None:
                rpr.remove(m)


def strip_comment_anchors(doc) -> int:
    """Remove comment range markers and reference runs from the body."""
    body = doc.element.body
    n = 0
    for tag in ("w:commentRangeStart", "w:commentRangeEnd"):
        for el in list(body.iter(qn(tag))):
            el.getparent().remove(el)
            n += 1
    for r in list(body.iter(qn("w:r"))):
        if r.find(qn("w:commentReference")) is not None:
            r.getparent().remove(r)
    return n


def _clear_comment_parts(path: Path) -> None:
    """Empty the comments parts so no orphaned threads survive the mint."""
    parts = _zip_parts(path)
    changed = False
    for name in list(parts):
        if name.startswith("word/comments") and name.endswith(".xml"):
            root = etree.fromstring(parts[name])
            for child in list(root):
                root.remove(child)
            parts[name] = etree.tostring(root, xml_declaration=True,
                                         encoding="UTF-8", standalone=True)
            changed = True
    if changed:
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
            for n, b in parts.items():
                z.writestr(n, b)


_LIST_STYLE_RE = re.compile(r"^(list\b|compact$)", re.IGNORECASE)


def _list_level(p) -> int | None:
    """How deeply this paragraph is nested in a list, or None if it is not in one.

    Two spellings, because two producers. Word and pandoc put ``w:numPr`` on the PARAGRAPH,
    with ``w:ilvl`` for depth. python-docx's built-in "List Bullet" styles put the numbering
    on the STYLE and leave the paragraph bare, so numPr alone reports a hand-built list as
    prose. Accept either; take the depth from ilvl when it is there and from the style's
    trailing digit ("List Bullet 2") when it is not.

    Headings carry numPr too — the reference document numbers them that way — so heading
    paragraphs must be classified before this is consulted, never by it.
    """
    pPr = p._p.find(qn("w:pPr"))
    numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
    if numPr is not None:
        ilvl = numPr.find(qn("w:ilvl"))
        try:
            return int(ilvl.get(qn("w:val"))) if ilvl is not None else 0
        except (TypeError, ValueError):
            return 0
    style = p.style.name if p.style is not None else ""
    if not _LIST_STYLE_RE.match(style.strip()):
        return None
    tail = style.strip().split()[-1]
    return int(tail) - 1 if tail.isdigit() and int(tail) > 0 else 0


def release_markdown(doc) -> str:
    """A plain markdown rendering of the (already accepted) document body.

    Bullets must come back as bullets. This used to emit every non-heading paragraph flat,
    so an outline the author had redlined — where each bullet is a real Word list item —
    was released as running prose, and the whole ladder below it lost the one structure it
    is built on. raconteur's draft prompt says "write ONE PARAGRAPH per outline bullet"; the
    css2026 release it said that against contained no bullets at all, only paragraphs, and
    the guards that count beats and the model that counts bullets were reading different
    documents.

    Any list becomes a ``-`` list: an outline's items are beats, never an ordinal sequence,
    and re-deriving ``1.`` from ``w:numFmt`` would only invite a numbering to drift from the
    heading numbers around it.
    """
    lines: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        if is_heading_style(style):
            level = int("".join(ch for ch in style if ch.isdigit()) or "1")
            lines.append("#" * level + " " + text)
            continue
        depth = _list_level(p)
        lines.append(text if depth is None else "  " * depth + "- " + text)
    return "\n\n".join(lines) + "\n"


_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_SIDE_PARTS = re.compile(r"^word/(comments\w*\.xml|people\.xml)$")


def set_comment_text(path: Path, updates: dict[str, str]) -> int:
    """Rewrite the text of existing comments, in place, keyed by comment id.

    Author, date, anchor and the ``w14:paraId`` that carries the resolved flag are all left
    alone: this changes what a comment SAYS, never whose it is or whether it is settled.
    """
    if not updates:
        return 0
    doc = Document(str(path))
    n = 0
    for rel in doc.part.rels.values():
        if not rel.reltype.lower().endswith("/comments"):
            continue
        root = rel.target_part._element
        for c in root.findall(".//" + qn("w:comment")):
            new = updates.get(c.get(qn("w:id")))
            if new is None:
                continue
            runs = [r for r in c.iter(qn("w:r")) if r.find(qn("w:t")) is not None]
            if not runs:
                continue
            first, rest = runs[0], runs[1:]
            for extra in first.findall(qn("w:t"))[1:]:
                first.remove(extra)
            t = first.find(qn("w:t"))
            t.set(_XML_SPACE, "preserve")
            t.text = new
            for r in rest:
                parent = r.getparent()
                if parent is not None:
                    parent.remove(r)
            n += 1
    if n:
        doc.save(str(path))
    return n


def _carry_comment_parts(src: Path, dst: Path) -> list[str]:
    """Copy the comment side-parts python-docx does not understand from src into dst.

    python-docx round-trips ``comments.xml`` and silently drops ``commentsExtended.xml`` —
    the part holding the RESOLVED flags. A release minted without it carries every comment
    as though it were still open, so a gate run on the release blocks on a conversation
    settled before the mint. Verified rather than assumed: the flags were gone and
    ``gate_check`` on the release returned False.

    Relationship ids are re-minted rather than copied. The source's rId may already mean
    something else in the saved document, and a rels file with two of the same id is a
    document Word offers to repair.
    """
    with zipfile.ZipFile(src) as zs:
        src_parts = {n: zs.read(n) for n in zs.namelist()}
    with zipfile.ZipFile(dst) as zd:
        dst_parts = {n: zd.read(n) for n in zd.namelist()}
    missing = [n for n in src_parts if _SIDE_PARTS.match(n) and n not in dst_parts]
    if not missing:
        return []

    ct_name = "[Content_Types].xml"
    ct = etree.fromstring(dst_parts[ct_name])
    have_ct = {o.get("PartName") for o in ct.findall(f"{{{_CT_NS}}}Override")}
    src_ct = etree.fromstring(src_parts[ct_name])
    src_over = {o.get("PartName"): o.get("ContentType")
                for o in src_ct.findall(f"{{{_CT_NS}}}Override")}

    rels_name = "word/_rels/document.xml.rels"
    rels = etree.fromstring(dst_parts[rels_name])
    used = {r.get("Id") for r in rels}
    src_rels = etree.fromstring(src_parts[rels_name])
    src_by_target = {r.get("Target"): r.get("Type") for r in src_rels}
    next_id = max([int(i[3:]) for i in used if i.startswith("rId") and i[3:].isdigit()]
                  or [0]) + 1

    for name in missing:
        dst_parts[name] = src_parts[name]
        part = "/" + name
        if part not in have_ct and part in src_over:
            o = etree.SubElement(ct, f"{{{_CT_NS}}}Override")
            o.set("PartName", part)
            o.set("ContentType", src_over[part])
        target = name.split("/", 1)[1]
        if target in src_by_target:
            r = etree.SubElement(rels, f"{{{_REL_NS}}}Relationship")
            r.set("Id", f"rId{next_id}"); next_id += 1
            r.set("Type", src_by_target[target])
            r.set("Target", target)
    dst_parts[ct_name] = etree.tostring(ct, xml_declaration=True, encoding="UTF-8",
                                        standalone=True)
    dst_parts[rels_name] = etree.tostring(rels, xml_declaration=True, encoding="UTF-8",
                                          standalone=True)
    with zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in dst_parts.items():
            z.writestr(n, b)
    return missing


def read_release(path: Path) -> str:
    """A release as markdown, whatever it is stored as.

    Releases are .docx and only .docx. A markdown sibling was a second copy of an approved
    contract, and a derived one: it went stale the moment its deriver was fixed, which is
    how an outline release sat carrying zero bullets for a day after release_markdown
    learned to keep them. It also could not carry the word plan at all, since a comment is
    not markdown.
    """
    if path.suffix.lower() == ".docx":
        return release_markdown(Document(str(path)))
    return path.read_text(encoding="utf-8", errors="replace")


# Every element Word uses to record a revision. Two of them carry prose; the rest carry
# moves and formatting, and were invisible to a mint that only knew w:ins and w:del.
REVISION_TAGS = (
    "w:ins", "w:del", "w:moveFrom", "w:moveTo",
    "w:moveFromRangeStart", "w:moveFromRangeEnd",
    "w:moveToRangeStart", "w:moveToRangeEnd",
    "w:pPrChange", "w:rPrChange", "w:sectPrChange", "w:tblPrChange",
    "w:tcPrChange", "w:trPrChange", "w:tblGridChange", "w:numberingChange",
    "w:cellIns", "w:cellDel", "w:cellMerge",
)


class UnacceptedRevisions(RuntimeError):
    """A release still carries tracked changes."""


def surviving_revisions(path: Path) -> dict[str, int]:
    """Revision elements left anywhere in the package, keyed ``part:tag``.

    EVERY part, because a check that reads only document.xml fails exactly as the accept
    did: the css2026 release passed this while footnotes.xml carried three insertions and
    styles.xml a style-definition change. A verification blind in the same place as the
    thing it verifies is not a verification.
    """
    import re as _re
    out: dict[str, int] = {}
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            xml = z.read(name).decode("utf-8", errors="replace")
            for tag in REVISION_TAGS:
                n = len(_re.findall(rf"<{tag}[ />]", xml))
                if n:
                    out[f"{name.split('/')[-1]}:{tag}"] = n
    return out


def _accept_in_parts(path: Path) -> int:
    """Accept revisions in parts python-docx does not model, at zip level.

    ``doc.part.rels`` reaches styles and comments but not footnotes, endnotes, headers or
    footers: those come back as generic parts with no ``_element``, so a loop over the rels
    skips them without a word. The css2026 release carried three insertions in
    footnotes.xml through a mint that reported clean.
    """
    parts = _zip_parts(path)
    touched = 0
    for name in list(parts):
        if not name.startswith("word/") or not name.endswith(".xml"):
            continue
        if name == "word/document.xml":
            continue                       # already done through python-docx
        raw = parts[name]
        if not any(f"<{t}".encode() in raw for t in REVISION_TAGS):
            continue
        root = etree.fromstring(raw)
        _accept_in(root, {"ins": 0, "del": 0, "paras": 0,
                          "moved": 0, "format": 0, "cells": 0})
        parts[name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                                     standalone=True)
        touched += 1
    if touched:
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
            for n, b in parts.items():
                z.writestr(n, b)
    return touched


def mint_release(src: Path, dst: Path, md_sibling: bool = True, post=None) -> dict:
    """Consolidate a gate-passed markup into a release: accept every tracked
    change, strip the comment threads, write the bare-name docx (and its .md
    sibling — what downstream LLM consumers actually read).

    The [@citekeys] stay as they are, here and everywhere (author's call, 2026-07-14): the
    key names the exact bibliography entry, and "(Bowling et al. 2018)" does not — three
    sources can render to the same author-year. The reviewer reads the key and knows what
    was cited; so does the next stage.
    """
    doc = Document(str(src))
    counts = accept_all_changes(doc)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    # Comments SURVIVE. The gate guarantees every one of them is resolved, so what the
    # release carries is a settled conversation and the acknowledgement that closed it —
    # and the next rung reads its predecessor's word plan off exactly that. Erasing them
    # was a coherence pair with strip_comment_anchors (anchors gone, so the threads had to
    # go too or dangle), not a policy; nothing in the pipeline ever read a release's
    # comments, and now something does.
    anchors = _carry_comment_parts(src, dst)
    _accept_in_parts(dst)
    # VERIFY, do not assume. The mint reported clean while 19 revisions rode through it,
    # and nothing noticed until the author opened the file in Word. A release that still
    # carries a tracked change is not a release.
    left = surviving_revisions(dst)
    if left:
        raise UnacceptedRevisions(
            f"{dst.name} still carries tracked changes after accepting: "
            + ", ".join(f"{k}={v}" for k, v in left.items())
            + ". Refusing to mint a release that is not clean.")
    # A rung may need the release reconciled with what the author actually approved. The
    # skeleton does: its word-plan comments were written against the structure as generated,
    # and the author has since moved subsections around. Runs AFTER the side-parts are
    # carried, so the resolved flags are in place before anything rewrites a comment.
    if post is not None:
        post(dst)
    # A markdown sibling is a SECOND copy of an approved contract, and a derived one: it
    # goes stale the moment its deriver is fixed, and it cannot carry the word plan at all,
    # since a comment is not markdown. Rungs are moved off it one at a time — a rung whose
    # consumers still read markdown must keep getting it, or they fall back in silence to
    # the tool's own pre-redline draft. Consumers read a release through ``read_release``.
    md_path = None
    if md_sibling:
        md_path = dst.with_suffix(".md")
        md_path.write_text(release_markdown(Document(str(dst))), encoding="utf-8")
    return {**counts, "anchors": anchors, "docx": dst, "md": md_path}
