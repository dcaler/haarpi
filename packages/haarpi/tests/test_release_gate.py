"""The mechanical gate + release minting, against real OOXML.

Documents are built with python-docx (comments via Document.add_comment,
tracked changes via the engine's own tracked_replace), and resolved flags are
injected into commentsExtended.xml the way Word writes them."""

import zipfile

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from haarpi import redline

W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
W15 = "http://schemas.microsoft.com/office/word/2012/wordml"


def _make_markup(path, comment_author="DCR", resolved=False, tracked=False):
    """A doc with one paragraph, one comment, optionally a tool tracked change,
    optionally the comment marked done in commentsExtended."""
    doc = Document()
    p = doc.add_paragraph("One sentence. Two sentence.")
    doc.add_comment(runs=[p.runs[0]], text="more on X", author=comment_author)
    if tracked:
        ids = redline.ids_for(doc)
        redline.tracked_replace(p._p, "One rewritten. Two sentence.", "rabbitHole", ids)
    doc.save(str(path))

    # give the comment paragraphs w14:paraId + write commentsExtended, as Word does
    with zipfile.ZipFile(path) as z:
        parts = {n: z.read(n) for n in z.namelist()}
    croot = etree.fromstring(parts["word/comments.xml"])
    pids = []
    for i, cp in enumerate(croot.iter(qn("w:p"))):
        pid = f"0000000{i+1}"
        cp.set(f"{{{W14}}}paraId", pid)
        pids.append(pid)
    parts["word/comments.xml"] = etree.tostring(croot, xml_declaration=True,
                                                encoding="UTF-8", standalone=True)
    ce = etree.Element(f"{{{W15}}}commentsEx", nsmap={"w15": W15})
    for pid in pids:
        cex = etree.SubElement(ce, f"{{{W15}}}commentEx")
        cex.set(f"{{{W15}}}paraId", pid)
        cex.set(f"{{{W15}}}done", "1" if resolved else "0")
    parts["word/commentsExtended.xml"] = etree.tostring(
        ce, xml_declaration=True, encoding="UTF-8", standalone=True)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in parts.items():
            z.writestr(n, b)
    return path


def test_unresolved_comment_blocks_the_gate(tmp_path):
    p = _make_markup(tmp_path / "m.docx", resolved=False)
    check = redline.gate_check(p)
    assert not check["clean"]
    assert check["unresolved"][0]["text"] == "more on X"
    assert check["unresolved"][0]["author"] == "DCR"


def test_resolved_comment_passes_the_gate(tmp_path):
    p = _make_markup(tmp_path / "m.docx", resolved=True)
    assert redline.gate_check(p)["clean"]


def test_a_tool_comment_blocks_until_it_is_acknowledged(tmp_path):
    """Tool comments used to be exempt, on the reasoning that the gate adjudicates the
    tool's work and the tool cannot instruct itself. That still holds for ``open_asks`` —
    what a revision ACTS on — and it does not hold for the gate: the skeleton's word plan
    rides on a tool comment, and a plan nobody read is a plan nobody approved."""
    p = _make_markup(tmp_path / "m.docx", comment_author="rabbitHole", resolved=False)
    assert not redline.gate_check(p)["clean"]


def test_a_resolved_tool_comment_passes(tmp_path):
    """Resolving is the acknowledgement. Resolved is not deleted — it rides into the
    release still carrying that state."""
    p = _make_markup(tmp_path / "m.docx", comment_author="rabbitHole", resolved=True)
    assert redline.gate_check(p)["clean"]


def test_a_tool_comment_is_still_not_an_instruction(tmp_path):
    """open_asks is what a revision acts on, and must keep excluding them — otherwise the
    tool answers its own word plan as though the author had asked for something."""
    p = _make_markup(tmp_path / "m.docx", comment_author="rabbitHole", resolved=False)
    assert redline.open_asks(p) == []
    assert redline.unresolved_comments(p)


def test_reviewer_tracked_change_alone_does_not_block(tmp_path):
    # A reviewer's own edit is their final word on that span, not a comment the tool
    # must answer. With no unresolved comment, the gate is clean and reports the edit
    # for context; the mint accepts it into the release.
    doc = Document()
    para = doc.add_paragraph("Original text here.")
    ids = redline.ids_for(doc)
    redline.tracked_replace(para._p, "Reviewer rewrote this.", "DCR", ids)
    doc.save(str(tmp_path / "m.docx"))
    check = redline.gate_check(tmp_path / "m.docx")
    assert check["clean"] and check["reviewer_changes"] > 0


def test_an_open_comment_still_blocks_with_tracked_changes_present(tmp_path):
    # Dropping the tracked-change conjunct must not weaken the comment gate: an
    # unresolved comment blocks even when the doc also carries tracked changes.
    p = _make_markup(tmp_path / "m.docx", resolved=False, tracked=True)
    assert not redline.gate_check(p)["clean"]


def test_tool_tracked_changes_do_not_block(tmp_path):
    p = _make_markup(tmp_path / "m.docx", resolved=True, tracked=True)
    assert redline.gate_check(p)["clean"]


def test_mint_release_accepts_and_writes_md(tmp_path):
    src = _make_markup(tmp_path / "m.docx", resolved=True, tracked=True)
    dst = tmp_path / "out" / "260715_myproj_litreview.docx"
    result = redline.mint_release(src, dst)
    assert dst.exists() and result["md"].exists()

    out = Document(str(dst))
    text = "\n".join(par.text for par in out.paragraphs)
    assert "One rewritten." in text            # insertion accepted
    assert "One sentence." not in text         # deletion applied
    body = out.element.body
    assert not list(body.iter(qn("w:ins"))) and not list(body.iter(qn("w:del")))
    assert "One rewritten." in result["md"].read_text()


def test_the_release_keeps_the_settled_conversation(tmp_path):
    """The gate guarantees every comment is resolved before a mint, so what survives is a
    settled conversation and the acknowledgement that closed it. The next rung reads its
    predecessor's word plan off exactly this."""
    src = _make_markup(tmp_path / "m.docx", resolved=True, tracked=True)
    dst = tmp_path / "out" / "260715_myproj_litreview.docx"
    redline.mint_release(src, dst)
    body = Document(str(dst)).element.body
    assert list(body.iter(qn("w:commentRangeStart"))), "anchors must survive"
    threads = redline.comment_threads(dst)
    assert [r["text"] for r in threads.values()] == ["more on X"]
    assert all(r["done"] for r in threads.values()), "the resolved flag must survive"


def test_the_resolved_flag_needs_a_part_python_docx_drops(tmp_path):
    """commentsExtended.xml holds the done flags and python-docx does not understand it, so
    a plain save loses it and every preserved comment reads as still open — the release then
    fails its own gate. Verified, not assumed."""
    src = _make_markup(tmp_path / "m.docx", resolved=True, tracked=True)
    dst = tmp_path / "out" / "260715_myproj_litreview.docx"
    result = redline.mint_release(src, dst, md_sibling=False)
    assert "word/commentsExtended.xml" in result["anchors"]
    assert redline.gate_check(dst)["clean"]


def test_minted_release_is_itself_gate_clean(tmp_path):
    src = _make_markup(tmp_path / "m.docx", resolved=True, tracked=True)
    dst = tmp_path / "260715_myproj_litreview.docx"
    redline.mint_release(src, dst, md_sibling=False)
    assert redline.gate_check(dst)["clean"]


def test_release_markdown_renders_headings(tmp_path):
    doc = Document()
    doc.add_heading("Results", level=2)
    doc.add_paragraph("Body text.")
    md = redline.release_markdown(doc)
    assert "## Results" in md and "Body text." in md


def test_release_markdown_keeps_bullets_as_bullets():
    """The release is what the next stage READS, and for an outline the bullets are the
    whole contract: raconteur's draft prompt says "write ONE PARAGRAPH per outline bullet".
    Emitting every non-heading paragraph flat released the css2026 outline as running prose
    with zero bullet markers — the author's approved beats, handed to the drafter as a
    summary to infer them from."""
    doc = Document()
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("Introduce the simulation.", style="List Bullet")
    doc.add_paragraph("Define the metric.", style="List Bullet")
    doc.add_paragraph("Not a list item.")
    md = redline.release_markdown(doc)
    assert "- Introduce the simulation." in md
    assert "- Define the metric." in md
    assert "\nNot a list item." in md and "- Not a list item." not in md


def test_a_nested_list_item_keeps_its_depth():
    doc = Document()
    doc.add_paragraph("top", style="List Bullet")
    doc.add_paragraph("under", style="List Bullet 2")
    md = redline.release_markdown(doc)
    assert "- top" in md
    assert "  - under" in md


def test_a_numbered_heading_is_not_mistaken_for_a_list_item():
    """Headings carry w:numPr too — the reference document numbers them that way. Classify
    the heading first or every section title comes back as a bullet."""
    doc = Document()
    doc.add_paragraph("Results", style="Heading 2")
    p = doc.paragraphs[-1]
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); numPr.append(ilvl)
    pPr.append(numPr)
    assert redline.release_markdown(doc).strip() == "## Results"


# ── one release, one contract ────────────────────────────────────────────────
# The release is two files: the .docx the author reads and the .md the next stage reads.
# Deriving the second from the first was supposed to make divergence impossible; derivation
# is exactly where structure is lost, so it guaranteed nothing.

def _skeleton_with_deleted_headings(path):
    """What a reviewer produces by deleting a heading in Word: the run text goes AND the
    paragraph mark is marked deleted."""
    doc = Document()
    doc.add_paragraph("Kept Heading", style="Heading 2")
    doc.add_paragraph("A bullet the author kept.", style="List Bullet")
    for text in ("Cut One", "Cut Two"):
        p = doc.add_paragraph(text, style="Heading 3")
        ids = redline.ids_for(doc)
        redline.tracked_replace(p._p, "", "DCR", ids)
        pPr = p._p.get_or_add_pPr()
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr"); pPr.insert(0, rPr)
        d = OxmlElement("w:del")
        d.set(qn("w:id"), "9001"); d.set(qn("w:author"), "DCR")
        d.set(qn("w:date"), "2026-07-21T00:00:00Z")
        rPr.insert(0, d)
    doc.save(str(path))
    return path


def test_a_deleted_heading_does_not_survive_as_an_empty_paragraph(tmp_path):
    """The css2026 skeleton released eleven empty Heading 3s from eleven headings the
    author had cut — each still holding a number in a numbered document. accept_all_changes
    stripped the mark-deletion marker and kept the paragraph, on the reasoning that the tool
    never authors one. The author does, every time they cut a section."""
    src = _skeleton_with_deleted_headings(tmp_path / "m.docx")
    doc = Document(str(src))
    counts = redline.accept_all_changes(doc)
    assert counts["paras"] == 2
    assert [p.text for p in doc.paragraphs if not p.text.strip()] == []
    assert [p.text for p in doc.paragraphs] == ["Kept Heading", "A bullet the author kept."]


def test_the_mark_is_collected_before_the_deletion_sweep(tmp_path):
    """A paragraph-mark deletion is a w:del, so the sweep that drops every w:del in the body
    removes it too. Collect after that and there is nothing left to find — which is what the
    first version of this fix did, silently removing zero paragraphs."""
    import inspect
    src = inspect.getsource(redline._accept_in)
    assert src.index("marked = [") < src.index('for d in list(body.iter(qn("w:del")))')


def test_no_queue_is_decided_once_for_every_branch():
    """run_next queues in three places — the rung, the stage advance, and rework. A flag
    that silenced one while another still fired would be worse than no flag: you would mint
    a skeleton believing nothing was queued and find a rework chain waiting."""
    import inspect
    from haarpi import planner
    src = inspect.getsource(planner.run_next)
    assert "queueing = bool(m.trundlr_project_id) and not no_queue" in src
    assert src.count("if queueing:") == 2
    assert "if not queueing:" in src
    assert "if m.trundlr_project_id:" not in src, "no branch may decide for itself"


def test_the_gate_record_is_still_written_without_queuing():
    """The ladder's own state must stay truthful — only the scheduler is left alone."""
    import inspect
    from haarpi import planner
    src = inspect.getsource(planner.run_next)
    rung = src.split("if deliverable:")[1].split("archived = _archive_chain")[0]
    assert "project.record_plan(root, {" in rung
    assert rung.index("project.record_plan") < rung.index("if queueing:")


def test_the_skipped_note_does_not_blame_a_missing_project_id():
    import inspect
    from haarpi import planner
    src = inspect.getsource(planner.run_next)
    assert 'why = ("--no-queue" if m.trundlr_project_id else "no project id")' in src


# ── accepting means accepting everything, everywhere ─────────────────────────
# Word records seven kinds of revision across half a dozen parts. This handled two kinds in
# one part, and the check that verified it was blind in the same place — so a release minted
# "clean" opened in Word still showing tracked changes.

def _revision_doc(path, body_extra=""):
    from docx import Document
    doc = Document()
    doc.add_paragraph("First half of a sentence")
    doc.add_paragraph("and its second half.")
    doc.save(str(path))
    return path


def test_a_deleted_paragraph_mark_merges_rather_than_splitting(tmp_path):
    """The defect that reached the page: section 5.3 shipped as two paragraphs, the second
    beginning mid-sentence. The old code stripped the marker and left the split, arguing
    that merging would lose prose — it does not. Moving the runs is what Word does."""
    from docx import Document
    src = tmp_path / "m.docx"
    doc = Document()
    a = doc.add_paragraph("Rendering a model as sound is an old move,")
    doc.add_paragraph("and listening reveals what a plot cannot.")
    pPr = a._p.get_or_add_pPr()
    rPr = OxmlElement("w:rPr"); pPr.insert(0, rPr)
    d = OxmlElement("w:del")
    d.set(qn("w:id"), "77"); d.set(qn("w:author"), "DCR")
    d.set(qn("w:date"), "2026-07-22T00:00:00Z")
    rPr.insert(0, d)
    doc.save(str(src))

    out = Document(str(src))
    redline.accept_all_changes(out)
    texts = [p.text for p in out.paragraphs if p.text.strip()]
    assert len(texts) == 1, "the two halves must become one paragraph"
    assert texts[0].startswith("Rendering a model as sound")
    assert "listening reveals" in texts[0]


def test_moved_text_is_not_left_in_both_places(tmp_path):
    """w:moveFrom is the deletion half of a move. Left beside its w:moveTo, the passage is
    in the document twice — 85 words of the css2026 release were."""
    import inspect
    src = inspect.getsource(redline._accept_in)
    assert 'qn("w:moveFrom")' in src and 'qn("w:moveTo")' in src
    assert "moveFromRangeStart" in src


def test_every_revision_kind_is_handled():
    import inspect
    src = inspect.getsource(redline._accept_in)
    for tag in ("w:pPrChange", "w:rPrChange", "w:sectPrChange", "w:tblPrChange",
                "w:tcPrChange", "w:trPrChange", "w:cellIns", "w:cellDel"):
        assert tag in src, tag


def test_the_check_reads_every_part_not_just_the_body():
    """It passed a release whose footnotes.xml carried three insertions and whose
    styles.xml carried a style-definition change. A verification blind in the same place as
    the thing it verifies is not a verification."""
    import inspect
    src = inspect.getsource(redline.surviving_revisions)
    assert 'z.namelist()' in src and 'name.startswith("word/")' in src
    assert 'z.read("word/document.xml")' not in src


def test_parts_python_docx_cannot_reach_are_accepted_at_zip_level():
    """doc.part.rels reaches styles and comments but not footnotes, endnotes, headers or
    footers — those come back with no _element and a rels loop skips them silently."""
    import inspect
    assert "footnotes" in inspect.getsource(redline._accept_in_parts)
    assert "_accept_in_parts(dst)" in inspect.getsource(redline.mint_release)


def test_a_release_that_is_not_clean_is_refused(tmp_path, monkeypatch):
    """The mint reported clean while 19 revisions rode through it."""
    src = _make_markup(tmp_path / "m.docx", resolved=True, tracked=True)
    monkeypatch.setattr(redline, "_accept_in", lambda *a, **k: None)
    with pytest.raises(redline.UnacceptedRevisions) as e:
        redline.mint_release(src, tmp_path / "out" / "260722_p_paper.docx", md_sibling=False)
    assert "still carries tracked changes" in str(e.value)
