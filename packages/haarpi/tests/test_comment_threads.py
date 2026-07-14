"""What a comment IS — resolution, threads, and the escalation signal.

Every caller must agree about this, and they did not: the gate honoured w15:done while
the reviser did not, so nine settled comments were re-fired as fresh instructions and a
resolved "this sentence is ok" provoked a full re-cut of the paragraph it had blessed.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from haarpi import redline as rl
from haarpi.testing import write_commented_docx

HUMAN = "D. Cale Reeves"


def _docx_with_comments(tmp_path: Path, comments: list[dict]) -> Path:
    return write_commented_docx(
        tmp_path / f"doc{len(comments)}.docx",
        ["The quick brown fox. It jumps the lazy dog."],
        comments,
    )


@pytest.fixture()
def doc(tmp_path) -> Path:
    return _docx_with_comments(tmp_path, [
        {"cid": "1", "author": HUMAN, "text": "define this"},
        {"cid": "2", "author": HUMAN, "text": "this is fine", "done": True},
        {"cid": "3", "author": "raconteur", "text": "We defined it.", "parent": "1"},
        {"cid": "4", "author": HUMAN, "text": "still not defined", "parent": "1"},
    ])


def test_a_resolved_comment_is_not_an_instruction(doc):
    ids = {a["id"] for a in rl.open_asks(doc)}
    assert "2" not in ids, "a comment marked done must never be re-fired"


def test_a_reply_is_not_a_separate_ask(doc):
    ids = {a["id"] for a in rl.open_asks(doc)}
    assert ids == {"1"}, "replies belong to their parent ask, not beside it"


def test_the_thread_travels_with_the_ask(doc):
    """New information about an unmet ask arrives as a reply. A brief that drops it
    drops half of what the reviewer said."""
    ask = rl.open_asks(doc)[0]
    assert ask["followups"] == ["still not defined"]


def test_an_answered_ask_left_open_is_a_repeat(doc):
    """The reviewer should never have to hand-write 'this STILL needs a definition' to
    be heard twice — the document already says so."""
    ask = rl.open_asks(doc)[0]
    assert ask["repeat"] is True
    assert ask["prior_tool_replies"] == ["We defined it."]


def test_an_unanswered_ask_is_not_a_repeat(tmp_path):
    d = _docx_with_comments(tmp_path, [
        {"cid": "1", "author": HUMAN, "text": "define this"},
    ])
    assert rl.open_asks(d)[0]["repeat"] is False


def test_resolving_the_parent_resolves_the_thread(tmp_path):
    d = _docx_with_comments(tmp_path, [
        {"cid": "1", "author": HUMAN, "text": "ask", "done": True},
        {"cid": "2", "author": HUMAN, "text": "more", "parent": "1"},
    ])
    assert rl.open_asks(d) == []
    assert rl.comment_threads(d)["2"]["done"] is True


def test_the_gate_and_the_reviser_now_read_the_same_comments(doc):
    assert ({c["id"] for c in rl.unresolved_comments(doc)}
            == {a["id"] for a in rl.open_asks(doc)})


def test_a_comment_on_deleted_text_is_still_seen(tmp_path):
    """A reviewer who comments on a phrase and then deletes it leaves the anchor inside
    the w:del. A direct-children search loses the comment silently — the worst way for
    an instruction to die."""
    from haarpi.redline import _Ids, tracked_replace

    path = tmp_path / "d.docx"
    d = Document()
    p = d.add_paragraph("Alpha sentence here. Beta sentence here.")
    p_el = p._p
    p_el.insert(0, p_el.makeelement(qn("w:commentRangeStart"), {qn("w:id"): "7"}))
    p_el.append(p_el.makeelement(qn("w:commentRangeEnd"), {qn("w:id"): "7"}))
    # the tool deletes the commented text: the anchor ends up inside <w:del>
    tracked_replace(p_el, "Gamma sentence entirely.", "raconteur", _Ids(500))
    d.save(str(path))

    p_el = Document(str(path)).paragraphs[0]._p
    nested = [s.get(qn("w:id")) for s in p_el.iter(qn("w:commentRangeStart"))]
    assert "7" in nested
    spans = rl.comment_spans(p_el)
    assert "7" in spans, "a comment anchored inside deleted text must still have a span"
    text = rl.paragraph_text(p_el)
    assert rl.anchored_sentences(text, spans["7"]), "and must bear on some sentence"
