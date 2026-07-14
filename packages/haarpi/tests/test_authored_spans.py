"""Deference is owed to the text a person wrote — not to the paragraph it sits in.

The tool used to freeze any paragraph carrying a human tracked change. That protected the
author's sentences and, in the same stroke, froze the tool's own prose beside them: a
surviving machine-written sentence in the Gap beat became untouchable because the author
had edited elsewhere in the paragraph. Wrong unit.

An authored span is an ATOM: content the tool preserves but never authors — the identical
contract an equation already has. The tool writes AROUND it.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from haarpi import redline as rl

TOOL, HUMAN = "raconteur", "D. Cale Reeves"


def _para_with_authored_span(tmp_path):
    """A paragraph the tool wrote, into which the author typed a sentence of their own."""
    path = tmp_path / "p.docx"
    d = Document()
    p = d.add_paragraph()
    p.add_run("Tool wrote this first sentence. ")
    p._p.append(rl._ins("The author typed this one by hand. ", HUMAN, 1))
    p.add_run("Tool wrote this last one.")
    d.save(str(path))
    return Document(str(path)).paragraphs[0]._p, path


def test_an_authored_span_serializes_as_an_atom(tmp_path):
    p_el, _ = _para_with_authored_span(tmp_path)
    text = rl.paragraph_text(p_el, protect_authored=True)
    assert "⟦a:1⟧" in text
    assert "The author typed this one" not in text, "the model must not be handed it as prose"


def test_without_protection_provenance_is_erased(tmp_path):
    """The old behaviour, pinned: this is why freezing had to be paragraph-wide."""
    p_el, _ = _para_with_authored_span(tmp_path)
    text = rl.paragraph_text(p_el)
    assert "The author typed this one by hand." in text
    assert "⟦a:" not in text


def test_the_legend_carries_the_authors_exact_words(tmp_path):
    """Readable, so the tool can write around it coherently. Immutable, all the same."""
    p_el, _ = _para_with_authored_span(tmp_path)
    assert rl.authored_atoms(p_el) == {"⟦a:1⟧": "The author typed this one by hand. "}


def test_word_split_runs_coalesce_into_one_span(tmp_path):
    """Word splits one typed sentence across several w:ins elements. What the author wrote
    is ONE span, not five."""
    path = tmp_path / "split.docx"
    d = Document()
    p = d.add_paragraph()
    p.add_run("Tool prose. ")
    for frag in ("The author ", "typed this ", "across three runs. "):
        p._p.append(rl._ins(frag, HUMAN, 1))
    d.save(str(path))
    p_el = Document(str(path)).paragraphs[0]._p
    atoms = rl.authored_atoms(p_el)
    assert list(atoms) == ["⟦a:1⟧"]
    assert atoms["⟦a:1⟧"] == "The author typed this across three runs. "


def test_the_tool_rewrites_its_own_prose_around_the_authors(tmp_path):
    """The whole point: the tool's sentences beside the author's are still fair game."""
    p_el, path = _para_with_authored_span(tmp_path)
    ok = rl.tracked_replace_sentencewise(
        p_el,
        "Tool rewrote its opener. ⟦a:1⟧Tool rewrote its closer.",
        TOOL, rl._Ids(500), protect_authored=True)
    assert ok

    tool_ins = [e for e in p_el.iter(qn("w:ins"))
                if e.get(qn("w:author")) == TOOL]
    human_ins = [e for e in p_el.iter(qn("w:ins"))
                 if e.get(qn("w:author")) == HUMAN]
    assert tool_ins, "the tool's own prose was re-cut"
    assert human_ins, "the author's insertion is still there"

    # still THEIRS: their tracked insertion, still pending, not nested inside a tool edit
    kept = "".join(t.text or "" for e in human_ins for t in e.iter(qn("w:t")))
    assert kept == "The author typed this one by hand. "
    assert all(e.getparent().tag == qn("w:p") for e in human_ins)


def test_the_authors_deletion_stays_dead(tmp_path):
    """Dead is dead: text the author deleted is never resurrected by a re-cut."""
    path = tmp_path / "del.docx"
    d = Document()
    p = d.add_paragraph()
    p.add_run("Kept sentence. ")
    p._p.append(rl._del("The author deleted this. ", HUMAN, 1))
    d.save(str(path))
    p_el = Document(str(path)).paragraphs[0]._p

    assert "deleted this" not in rl.paragraph_text(p_el, protect_authored=True)
    rl.tracked_replace_sentencewise(p_el, "Rewritten sentence.", TOOL, rl._Ids(500),
                                    protect_authored=True)
    accepted = rl._accepted_para_text(p_el)
    assert "The author deleted this." not in accepted


def test_a_prior_deletion_keeps_its_place_in_the_paragraph(tmp_path):
    """A rebuild that re-laid the new body at the top swept every old deletion to the
    paragraph's tail, severing the struck-through text from the prose it was struck from."""
    path = tmp_path / "order.docx"
    d = Document()
    p = d.add_paragraph()
    p._p.append(rl._del("Deleted opener. ", HUMAN, 1))
    p.add_run("Surviving tool prose here.")
    d.save(str(path))
    p_el = Document(str(path)).paragraphs[0]._p

    rl.tracked_replace_sentencewise(p_el, "Rewritten tool prose here.", TOOL,
                                    rl._Ids(500), protect_authored=True)
    tags = [el.tag.split("}")[1] for el in p_el
            if el.tag.split("}")[1] in ("del", "ins", "r")]
    assert tags.index("del") == 0, f"the deletion must stay where it was, got {tags}"
