"""A correction should POINT, not restate.

raconteur found three real typos in the author's own sentences — "though" for "thought",
"has dates" for "dates", "it's" for "its" — and reported each one by handing back the whole
sentence it appeared in, fifty-five words of the author's own prose with the change buried
somewhere inside. The author read it, correctly, as the tool vomiting his text back at him.

The comment belongs ON the word.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from haarpi import redline as rl
from haarpi.testing import write_commented_docx
from haarpi.text import copyedit_notes

HUMAN = "D. Cale Reeves"


# ── the diff: what to say ────────────────────────────────────────────────────

def test_a_correction_is_the_words_that_change():
    notes = copyedit_notes(
        "The Schelling model has dates to 1971 and encourages new though patterns.",
        "The Schelling model dates to 1971 and encourages new thought patterns.")
    assert notes == [("has", "(cut)"), ("though", "thought")]


def test_an_apostrophe_is_reported_as_an_apostrophe():
    assert copyedit_notes("by re-expressing it’s visual output",
                          "by re-expressing its visual output") == [("it’s", "its")]


def test_an_identical_suggestion_says_nothing():
    assert copyedit_notes("No error here.", "No error here.") == []


def test_an_insertion_hangs_off_the_words_that_follow_it():
    notes = copyedit_notes("the model dates 1971", "the model dates to 1971")
    assert notes == [("1971", 'insert: "to"')]


# ── the anchor: where to say it ──────────────────────────────────────────────

def _doc_with_authored_sentence(tmp_path) -> Path:
    """A paragraph of tool prose into which the author typed a sentence, tracked."""
    path = tmp_path / "d.docx"
    d = Document()
    p = d.add_paragraph()
    p.add_run("Motivation — ")
    p._p.append(rl._ins("However, by re-expressing it’s visual output in audio.",
                        HUMAN, 1))
    d.save(str(path))
    return path


def test_the_anchor_covers_the_offending_word_and_nothing_else(tmp_path):
    doc = _doc_with_authored_sentence(tmp_path)
    p_el = Document(str(doc)).paragraphs[0]._p

    assert rl.anchor_fragment(p_el, "it’s", "900")

    lo, hi = rl.comment_spans(p_el)["900"]
    assert rl.paragraph_text(p_el)[lo:hi] == "it’s"


def test_anchoring_inside_a_tracked_insertion_leaves_it_the_authors(tmp_path):
    """The author's sentences ARE tracked insertions — that is where the tool most often has
    something to say. Splitting the run must not touch a character of it, nor cost them the
    revision record."""
    doc = _doc_with_authored_sentence(tmp_path)
    d = Document(str(doc))
    p_el = d.paragraphs[0]._p
    before = rl._accepted_para_text(p_el)

    rl.anchor_fragment(p_el, "it’s", "900")

    assert rl._accepted_para_text(p_el) == before, "not one character moved"
    assert list(rl.authored_atoms(p_el)) == ["⟦a:1⟧"], "still ONE authored span"
    ins = [e for e in p_el.iter(qn("w:ins"))
           if e.get(qn("w:author")) == HUMAN]
    assert ins, "still the author's tracked insertion, still theirs to accept"


def test_a_fragment_spanning_two_runs_still_anchors(tmp_path):
    """Word splits a typed sentence across runs at arbitrary points."""
    path = tmp_path / "split.docx"
    d = Document()
    p = d.add_paragraph()
    p._p.append(rl._ins("the blind-mon", HUMAN, 1))
    p._p.append(rl._ins("key baseline holds", HUMAN, 2))
    d.save(str(path))

    p_el = Document(str(path)).paragraphs[0]._p
    assert rl.anchor_fragment(p_el, "blind-monkey", "901")
    lo, hi = rl.comment_spans(p_el)["901"]
    assert rl.paragraph_text(p_el)[lo:hi] == "blind-monkey"


def test_a_fragment_that_is_not_there_does_not_anchor(tmp_path):
    p_el = Document(str(_doc_with_authored_sentence(tmp_path))).paragraphs[0]._p
    assert rl.anchor_fragment(p_el, "nonexistent words", "902") is False


def test_the_delivered_comment_sits_on_the_word(tmp_path):
    """End to end: the author sees the word highlighted and the correction beside it."""
    doc = write_commented_docx(
        tmp_path / "rev.docx",
        ["Motivation — However, by re-expressing it’s visual output in audio."],
        [{"cid": "1", "author": HUMAN, "text": "an existing comment"}])

    n = rl.add_anchored_comments(doc, [("it’s", "its")], author="raconteur")
    assert n == 1

    p_el = Document(str(doc)).paragraphs[0]._p
    cm = rl.comments_by_id(doc)
    text = rl.paragraph_text(p_el)
    tool = [cid for cid, c in cm.items()
            if (c.get("author") or "").lower() == "raconteur"]
    lo, hi = rl.comment_spans(p_el)[tool[0]]

    assert text[lo:hi] == "it’s", "highlighted: the word"
    assert cm[tool[0]]["text"] == "its", "margin: the correction, and nothing else"


def test_an_unfindable_fragment_falls_back_to_the_paragraph(tmp_path):
    """Worse, but never wrong."""
    doc = write_commented_docx(
        tmp_path / "rev.docx", ["A sentence about nothing in particular."],
        [{"cid": "1", "author": HUMAN, "text": "an existing comment"}])
    assert rl.add_anchored_comments(doc, [("A sentence", "a whole-paragraph note")]) == 1
